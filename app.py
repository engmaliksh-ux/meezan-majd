import os
import secrets
from flask import Flask, render_template, request, redirect, url_for, session, flash, jsonify
from database import get_connection, init_db, generate_org_code
from translations import TRANSLATIONS
from functools import wraps
from werkzeug.security import generate_password_hash, check_password_hash
from email_service import generate_verification_code, send_org_verification, send_staff_notification
from datetime import datetime, timedelta

app = Flask(__name__)

# ══════════════════════════════════════════
# 1. Secret Key من البيئة (لا تضعها في الكود أبداً)
#    في PythonAnywhere: Web → Environment variables → SECRET_KEY
#    قيمة عشوائية مثال: python3 -c "import secrets; print(secrets.token_hex(32))"
# ══════════════════════════════════════════
app.secret_key = os.environ.get("SECRET_KEY") or "3c23c600bdf76f847bcd15584b92ad49d63764935b05e1a4a06d5c75e9e03aaa"

# ══════════════════════════════════════════
# 2. إعدادات Session الآمنة
# ══════════════════════════════════════════
app.config["SESSION_COOKIE_HTTPONLY"]  = True   # لا يصل JavaScript للـ cookie
app.config["SESSION_COOKIE_SAMESITE"]  = "Lax"  # حماية CSRF جزئية
app.config["SESSION_COOKIE_SECURE"]    = True   # HTTPS فقط
app.config["PERMANENT_SESSION_LIFETIME"] = timedelta(hours=8)  # تنتهي بعد 8 ساعات

UPLOAD_FOLDER = os.path.join("static", "uploads")
ALLOWED_EXTENSIONS = {"png", "jpg", "jpeg", "webp"}
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER
app.config["MAX_CONTENT_LENGTH"] = 3 * 1024 * 1024  # 3MB حد رفع الملفات

# ══════════════════════════════════════════
# 3. Brute Force Protection — تتبع محاولات تسجيل الدخول
# ══════════════════════════════════════════
_login_attempts: dict = {}   # {ip: {"count": N, "blocked_until": datetime}}
MAX_LOGIN_ATTEMPTS = 5
BLOCK_MINUTES = 15

def _get_ip():
    # نستخدم remote_addr الحقيقي من الـ proxy — لا نثق بـ X-Forwarded-For القابل للتلاعب
    forwarded = request.headers.get("X-Forwarded-For")
    if forwarded:
        # آخر IP في السلسلة هو الـ proxy الموثوق
        return forwarded.split(",")[-1].strip()
    return request.remote_addr or "unknown"

def _is_blocked(ip):
    rec = _login_attempts.get(ip)
    if not rec:
        return False
    if rec.get("blocked_until") and datetime.now() < rec["blocked_until"]:
        return True
    if rec.get("blocked_until") and datetime.now() >= rec["blocked_until"]:
        _login_attempts.pop(ip, None)
    return False

def _record_failed(ip):
    rec = _login_attempts.setdefault(ip, {"count": 0, "blocked_until": None})
    rec["count"] += 1
    if rec["count"] >= MAX_LOGIN_ATTEMPTS:
        rec["blocked_until"] = datetime.now() + timedelta(minutes=BLOCK_MINUTES)

def _clear_attempts(ip):
    _login_attempts.pop(ip, None)

# ══════════════════════════════════════════
# 4. CSRF Token
# ══════════════════════════════════════════
def generate_csrf():
    if "_csrf" not in session:
        session["_csrf"] = secrets.token_hex(24)
    return session["_csrf"]

def validate_csrf():
    token = session.get("_csrf")
    form_token = request.form.get("_csrf_token") or request.headers.get("X-CSRF-Token")
    if not token or not form_token or not secrets.compare_digest(token, form_token):
        return False
    return True

app.jinja_env.globals["csrf_token"] = generate_csrf
# مساعدات Jinja2 للقوالب
app.jinja_env.globals["enumerate"] = enumerate
app.jinja_env.filters["enumerate"] = enumerate   # للاستخدام كـ |enumerate في القوالب
from datetime import datetime as _DT
app.jinja_env.globals["now"] = _DT.now

# ══════════════════════════════════════════
# 5. Security Headers لكل الردود
# ══════════════════════════════════════════
@app.after_request
def add_security_headers(resp):
    resp.headers["X-Frame-Options"]        = "DENY"
    resp.headers["X-Content-Type-Options"] = "nosniff"
    resp.headers["X-XSS-Protection"]       = "1; mode=block"
    resp.headers["Referrer-Policy"]        = "strict-origin-when-cross-origin"
    resp.headers["Content-Security-Policy"] = (
        "default-src 'self' https://cdn.jsdelivr.net https://fonts.googleapis.com "
        "https://fonts.gstatic.com https://cdnjs.cloudflare.com; "
        "img-src 'self' data:; style-src 'self' 'unsafe-inline' "
        "https://cdn.jsdelivr.net https://fonts.googleapis.com "
        "https://cdnjs.cloudflare.com; "
        "script-src 'self' 'unsafe-inline' 'unsafe-eval' https://cdn.jsdelivr.net https://cdnjs.cloudflare.com;"
    )
    return resp

# ══════════════════════════════════════════
# 6. Session Timeout — التحقق عند كل طلب
# ══════════════════════════════════════════
@app.before_request
def check_session_timeout():
    if "user_id" not in session:
        return
    last = session.get("_last_active")
    if last:
        idle = (datetime.now() - datetime.fromisoformat(last)).total_seconds()
        if idle > 8 * 3600:
            session.clear()
            flash("انتهت جلستك، يرجى تسجيل الدخول مجدداً", "warning")
            return redirect(url_for("login"))
    session["_last_active"] = datetime.now().isoformat()

init_db()
from database import migrate_db; migrate_db()

# ══════════════════════════════════════════
# Language support — AR / TR
# ══════════════════════════════════════════
@app.route("/set_lang/<lang>")
def set_lang(lang):
    if lang in ('ar', 'tr', 'en'):
        session['lang'] = lang
    return redirect(request.referrer or url_for('dashboard'))

@app.context_processor
def inject_lang():
    lang = session.get('lang', 'ar')
    return dict(t=TRANSLATIONS[lang], lang=lang)

# ══════════════════════════════════════════
# Helpers
# ══════════════════════════════════════════

# فحص الامتداد + المحتوى الفعلي للصورة (magic bytes)
_IMG_MAGIC = {
    b"\xff\xd8\xff": "jpeg",
    b"\x89PNG":      "png",
    b"RIFF":         "webp",
    b"GIF8":         "gif",
}

def allowed_file(f):
    if "." not in f:
        return False
    return f.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS

def allowed_file_content(file_storage):
    """يفحص magic bytes الفعلية للملف، ليس فقط الامتداد"""
    header = file_storage.read(12)
    file_storage.seek(0)
    for magic, ftype in _IMG_MAGIC.items():
        if header.startswith(magic):
            return ftype in ALLOWED_EXTENSIONS
    return False


def _csrf_check():
    """يرفض POST بدون CSRF token صحيح"""
    if request.method == "POST" and not validate_csrf():
        flash("طلب غير صالح (CSRF). أعد المحاولة.", "danger")
        return redirect(request.referrer or url_for("dashboard"))


def login_required(f):
    @wraps(f)
    def dec(*a, **kw):
        if "user_id" not in session:
            flash("يجب تسجيل الدخول أولاً", "warning")
            return redirect(url_for("login"))
        err = _csrf_check()
        if err: return err
        return f(*a, **kw)
    return dec


def admin_required(f):
    @wraps(f)
    def dec(*a, **kw):
        if "user_id" not in session:
            return redirect(url_for("login"))
        if session.get("role") != "admin":
            flash("هذه الصفحة للمدير فقط", "danger")
            return redirect(url_for("dashboard"))
        err = _csrf_check()
        if err: return err
        return f(*a, **kw)
    return dec


def invoices_required(f):
    @wraps(f)
    def dec(*a, **kw):
        if "user_id" not in session:
            return redirect(url_for("login"))
        if session.get("role") not in ("admin", "accountant"):
            flash("ليس لديك صلاحية الوصول", "danger")
            return redirect(url_for("dashboard"))
        err = _csrf_check()
        if err: return err
        return f(*a, **kw)
    return dec


def data_entry_required(f):
    @wraps(f)
    def dec(*a, **kw):
        if "user_id" not in session:
            return redirect(url_for("login"))
        if session.get("role") not in ("admin", "data_entry"):
            flash("ليس لديك صلاحية الوصول", "danger")
            return redirect(url_for("dashboard"))
        err = _csrf_check()
        if err: return err
        return f(*a, **kw)
    return dec


def observer_required(f):
    """المراقب العام: يرى فقط — لا يعدّل"""
    @wraps(f)
    def dec(*a, **kw):
        if "user_id" not in session:
            return redirect(url_for("login"))
        if session.get("role") not in ("admin", "accountant", "data_entry", "observer"):
            flash("ليس لديك صلاحية الوصول", "danger")
            return redirect(url_for("dashboard"))
        return f(*a, **kw)
    return dec


def invoices_view_required(f):
    """عرض صفحات المشتريات/المخزن: أدمن + محاسب + مراقب عام (المراقب لا يستطيع الحفظ)"""
    @wraps(f)
    def dec(*a, **kw):
        if "user_id" not in session:
            return redirect(url_for("login"))
        if session.get("role") not in ("admin", "accountant", "observer"):
            flash("ليس لديك صلاحية الوصول", "danger")
            return redirect(url_for("dashboard"))
        return f(*a, **kw)
    return dec


def data_entry_view_required(f):
    """عرض صفحة المستفيدين: أدمن + مدخل بيانات + مراقب عام (المراقب لا يستطيع الحفظ)"""
    @wraps(f)
    def dec(*a, **kw):
        if "user_id" not in session:
            return redirect(url_for("login"))
        if session.get("role") not in ("admin", "data_entry", "observer"):
            flash("ليس لديك صلاحية الوصول", "danger")
            return redirect(url_for("dashboard"))
        return f(*a, **kw)
    return dec


def programs_view_required(f):
    """عرض صفحات البرامج والمشاريع: أدمن + محاسب + مراقب عام"""
    @wraps(f)
    def dec(*a, **kw):
        if "user_id" not in session:
            return redirect(url_for("login"))
        if session.get("role") not in ("admin", "accountant", "observer"):
            flash("ليس لديك صلاحية الوصول", "danger")
            return redirect(url_for("dashboard"))
        return f(*a, **kw)
    return dec


def set_session(user, org):
    session["user_id"]   = user["id"]
    session["org_id"]    = user["org_id"]
    session["org_name"]  = org["name"]
    session["org_code"]  = org["org_code"]
    session["user"]      = user["username"]
    session["full_name"] = user["full_name"] or user["username"]
    session["role"]      = user["role"]
    session["photo"]     = user["photo"] or ""


def row_to_dict(row, fallback_keys=None):
    """تحويل sqlite3.Row لـ dict بأمان"""
    try:
        return dict(row)
    except Exception:
        if fallback_keys and row:
            return {k: (row[i] if i < len(row) else None)
                    for i, k in enumerate(fallback_keys)}
        return {}


def resequence(org_id):
    """إعادة الترقيم التسلسلي بعد الحذف"""
    conn = get_connection()
    c = conn.cursor()
    c.execute("SELECT id FROM beneficiaries WHERE org_id=? ORDER BY id ASC", (org_id,))
    for i, row in enumerate(c.fetchall(), 1):
        c.execute("UPDATE beneficiaries SET seq_num=? WHERE id=?", (i, row["id"]))
    conn.commit()
    conn.close()


# last_seen يُحدَّث عند تسجيل الدخول فقط لتجنب database lock


# ══════════════════════════════════════════
# الصفحة الرئيسية
# ══════════════════════════════════════════
@app.route("/show-error")
def show_error():
    import traceback, subprocess
    try:
        log = subprocess.run(["tail", "-50",
            "/var/log/malikmohs.pythonanywhere.com.error.log"],
            capture_output=True, text=True).stdout
    except Exception as e:
        log = str(e)
    return f"<pre style='direction:ltr;font-size:12px'>{log}</pre>"


@app.route("/diag")
def diag_route():
    import traceback, sys
    try:
        conn = get_connection()
        conn.close()
        db_ok = "OK"
    except Exception as e:
        db_ok = str(e)
    return f"<pre>DB:{db_ok}\nPY:{sys.version}\n</pre>"


@app.route("/")
def home():
    total_benef = total_camps = total_orgs = 0
    try:
        conn = get_connection()
        c = conn.cursor()
        try:
            c.execute("SELECT COUNT(*) FROM beneficiaries WHERE beneficiary_type != 'camp'")
            total_benef = c.fetchone()[0]
            c.execute("SELECT COUNT(*) FROM beneficiaries WHERE beneficiary_type = 'camp'")
            total_camps = c.fetchone()[0]
        except Exception:
            pass
        try:
            c.execute("SELECT COUNT(*) FROM organizations WHERE is_active=1")
            total_orgs = c.fetchone()[0]
        except Exception:
            pass
        conn.close()
    except Exception:
        pass
    return render_template("index.html",
        total_benef=total_benef,
        total_camps=total_camps,
        total_orgs=total_orgs)


# ══════════════════════════════════════════
# صفحة الاختيار
# ══════════════════════════════════════════
@app.route("/register")
def register():
    if "user_id" in session:
        return redirect(url_for("dashboard"))
    return render_template("register_choice.html")


# ══════════════════════════════════════════
# تسجيل مؤسسة جديدة
# ══════════════════════════════════════════
@app.route("/register/org", methods=["GET", "POST"])
def register_org():
    if "user_id" in session:
        return redirect(url_for("dashboard"))

    error = None
    if request.method == "POST":
        org_name   = request.form.get("org_name", "").strip()
        email      = request.form.get("email", "").strip().lower()
        donor_name = request.form.get("donor_name", "").strip()
        full_name  = request.form.get("full_name", "").strip()
        id_number  = request.form.get("id_number", "").strip()
        username   = request.form.get("username", "").strip()
        password   = request.form.get("password", "")
        password2  = request.form.get("password2", "")

        if not all([org_name, email, full_name, id_number, username, password]):
            error = "يرجى تعبئة جميع الحقول الإلزامية"
        elif len(username) < 3:
            error = "اسم المستخدم يجب أن يكون 3 أحرف على الأقل"
        elif len(password) < 8:
            error = "كلمة المرور يجب أن تكون 8 أحرف على الأقل"
        elif password != password2:
            error = "كلمتا المرور غير متطابقتين"
        elif "@" not in email:
            error = "البريد الإلكتروني غير صحيح"
        else:
            conn = get_connection()
            c = conn.cursor()
            c.execute("SELECT id FROM organizations WHERE email=?", (email,))
            if c.fetchone():
                error = "البريد الإلكتروني مسجّل لمؤسسة أخرى"
            else:
                c.execute("SELECT id FROM users WHERE username=?", (username,))
                if c.fetchone():
                    error = "اسم المستخدم مأخوذ، اختر اسماً آخر"
            conn.close()

        if not error:
            conn = get_connection()
            c = conn.cursor()
            # رمز المؤسسة فريد
            while True:
                org_code = generate_org_code()
                c.execute("SELECT id FROM organizations WHERE org_code=?", (org_code,))
                if not c.fetchone():
                    break
            # كود التحقق
            verify_code = generate_verification_code()
            expires_at  = (datetime.now() + timedelta(minutes=3)).strftime("%Y-%m-%d %H:%M:%S")
            c.execute("DELETE FROM verification_codes WHERE email=? AND purpose='org_register'", (email,))
            c.execute("INSERT INTO verification_codes (email, code, purpose, expires_at) VALUES (?,?,?,?)",
                      (email, verify_code, "org_register", expires_at))
            conn.commit()
            conn.close()

            session["pending_org"] = {
                "org_name": org_name, "email": email, "donor_name": donor_name,
                "full_name": full_name, "id_number": id_number,
                "username": username, "password": password, "org_code": org_code,
            }

            sent = send_org_verification(email, org_name, full_name, verify_code, org_code, username)
            if sent:
                flash(f"📧 تم إرسال كود التحقق إلى {email}", "success")
            else:
                flash("⚠️ تحقق من مجلد Spam في بريدك", "warning")
            return redirect(url_for("verify_org_email"))

    return render_template("register_org.html", error=error)


# ══════════════════════════════════════════
# التحقق من كود الإيميل
# ══════════════════════════════════════════
@app.route("/register/verify", methods=["GET", "POST"])
def verify_org_email():
    if "user_id" in session:
        return redirect(url_for("dashboard"))

    pending = session.get("pending_org")
    if not pending:
        flash("انتهت الجلسة، يرجى إعادة التسجيل", "warning")
        return redirect(url_for("register_org"))

    error = None
    if request.method == "POST":
        entered_code = request.form.get("code", "").strip()
        conn = get_connection()
        c = conn.cursor()
        c.execute(
            "SELECT * FROM verification_codes WHERE email=? AND purpose='org_register' AND used=0 ORDER BY id DESC LIMIT 1",
            (pending["email"],)
        )
        row = c.fetchone()

        if not row:
            error = "الكود غير موجود، أعد التسجيل"
        elif datetime.now() > datetime.strptime(row["expires_at"], "%Y-%m-%d %H:%M:%S"):
            error = "انتهت صلاحية الكود (3 دقائق)، اضغط إعادة الإرسال"
            c.execute("DELETE FROM verification_codes WHERE id=?", (row["id"],))
            conn.commit()
        elif entered_code != row["code"]:
            error = "الكود غير صحيح"
        else:
            c.execute("UPDATE verification_codes SET used=1 WHERE id=?", (row["id"],))
            c.execute("INSERT INTO organizations (name, email, donor_name, org_code) VALUES (?,?,?,?)",
                      (pending["org_name"], pending["email"], pending["donor_name"], pending["org_code"]))
            org_id = c.lastrowid
            hashed = generate_password_hash(pending["password"])
            c.execute(
                "INSERT INTO users (org_id, username, password, full_name, id_number, email, role, status) VALUES (?,?,?,?,?,?,?,?)",
                (org_id, pending["username"], hashed, pending["full_name"],
                 pending["id_number"], pending["email"], "admin", "approved")
            )
            conn.commit()
            conn.close()
            org_code = pending["org_code"]
            org_name = pending["org_name"]
            full_name = pending["full_name"]
            username = pending["username"]
            session.pop("pending_org", None)
            return render_template("register_success.html",
                org_name=org_name, org_code=org_code,
                full_name=full_name, username=username)
        conn.close()

    return render_template("verify_email.html", email=pending["email"], error=error)


@app.route("/register/resend")
def resend_verification():
    pending = session.get("pending_org")
    if not pending:
        return redirect(url_for("register_org"))
    verify_code = generate_verification_code()
    expires_at  = (datetime.now() + timedelta(minutes=3)).strftime("%Y-%m-%d %H:%M:%S")
    conn = get_connection()
    c = conn.cursor()
    c.execute("DELETE FROM verification_codes WHERE email=? AND purpose='org_register'", (pending["email"],))
    c.execute("INSERT INTO verification_codes (email, code, purpose, expires_at) VALUES (?,?,?,?)",
              (pending["email"], verify_code, "org_register", expires_at))
    conn.commit()
    conn.close()
    send_org_verification(pending["email"], pending["org_name"], pending["full_name"],
                          verify_code, pending["org_code"], pending["username"])
    flash(f"📧 تم إعادة إرسال الكود إلى {pending['email']}", "success")
    return redirect(url_for("verify_org_email"))


# ══════════════════════════════════════════
# تسجيل موظف
# ══════════════════════════════════════════
@app.route("/register/staff", methods=["GET", "POST"])
def register_staff():
    if "user_id" in session:
        return redirect(url_for("dashboard"))

    error = None
    if request.method == "POST":
        org_code  = request.form.get("org_code", "").strip().upper()
        role      = request.form.get("role", "")
        full_name = request.form.get("full_name", "").strip()
        id_number = request.form.get("id_number", "").strip()
        username  = request.form.get("username", "").strip()
        password  = request.form.get("password", "")
        password2 = request.form.get("password2", "")

        if role not in ("accountant", "data_entry", "observer"):
            error = "يرجى اختيار نوع الحساب"
        elif not all([org_code, full_name, id_number, username, password]):
            error = "يرجى تعبئة جميع الحقول الإلزامية"
        elif len(username) < 3:
            error = "اسم المستخدم 3 أحرف على الأقل"
        elif len(password) < 8:
            error = "كلمة المرور 8 أحرف على الأقل"
        elif password != password2:
            error = "كلمتا المرور غير متطابقتين"
        else:
            conn = get_connection()
            c = conn.cursor()
            c.execute("SELECT * FROM organizations WHERE org_code=? AND is_active=1", (org_code,))
            org = c.fetchone()
            if not org:
                error = "رمز المؤسسة غير صحيح"
                conn.close()
            else:
                c.execute("SELECT id FROM users WHERE username=?", (username,))
                if c.fetchone():
                    error = "اسم المستخدم مأخوذ"
                    conn.close()
                else:
                    photo_filename = None
                    photo_file = request.files.get("photo")
                    if photo_file and photo_file.filename and allowed_file(photo_file.filename) and allowed_file_content(photo_file):
                        ext = photo_file.filename.rsplit(".", 1)[1].lower()
                        photo_filename = f"user_{username}.{ext}"
                        photo_file.save(os.path.join(app.config["UPLOAD_FOLDER"], photo_filename))

                    hashed = generate_password_hash(password)
                    c.execute(
                        "INSERT INTO users (org_id, username, password, full_name, id_number, photo, role, status) VALUES (?,?,?,?,?,?,?,?)",
                        (org["id"], username, hashed, full_name, id_number, photo_filename, role, "pending")
                    )
                    conn.commit()
                    # إشعار الأدمن
                    c.execute("SELECT email, full_name FROM users WHERE org_id=? AND role='admin' LIMIT 1", (org["id"],))
                    admin = c.fetchone()
                    conn.close()
                    if admin and admin["email"]:
                        send_staff_notification(admin["email"], admin["full_name"],
                                                full_name, role, org["name"])
                    role_label = "محاسب" if role == "accountant" else "مدخل بيانات"
                    return render_template("register_pending.html",
                                           role_label=role_label, full_name=full_name)

    return render_template("register_staff.html", error=error)


# ══════════════════════════════════════════
# تسجيل الدخول / الخروج
# ══════════════════════════════════════════
@app.route("/login", methods=["GET", "POST"])
def login():
    if "user_id" in session:
        return redirect(url_for("dashboard"))

    ip = _get_ip()
    error = None

    if request.method == "POST":
        # حماية Brute Force
        if _is_blocked(ip):
            error = f"تم تجميد IP بسبب محاولات كثيرة. انتظر {BLOCK_MINUTES} دقيقة وحاول مجدداً."
            return render_template("login.html", error=error)

        org_code = request.form.get("org_code", "").strip().upper()
        username = request.form.get("username", "").strip()
        password = request.form.get("password", "")

        if not org_code or not username or not password:
            error = "يرجى تعبئة جميع الحقول"
        else:
            conn = get_connection()
            c = conn.cursor()
            c.execute("SELECT * FROM organizations WHERE org_code=? AND is_active=1", (org_code,))
            org = c.fetchone()
            if not org:
                _record_failed(ip)
                error = "رمز المؤسسة أو بيانات الدخول غير صحيحة"
                conn.close()
            else:
                c.execute("SELECT * FROM users WHERE username=? AND org_id=?", (username, org["id"]))
                user = c.fetchone()
                if not user or not check_password_hash(user["password"], password):
                    _record_failed(ip)
                    remaining = MAX_LOGIN_ATTEMPTS - _login_attempts.get(ip, {}).get("count", 0)
                    error = f"بيانات الدخول غير صحيحة. محاولات متبقية: {max(0, remaining)}"
                    conn.close()
                elif user["status"] == "pending":
                    error = "حسابك قيد المراجعة، انتظر موافقة المدير"
                    conn.close()
                elif user["status"] == "rejected":
                    error = "تم رفض طلبك، تواصل مع مدير المؤسسة"
                    conn.close()
                else:
                    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                    c.execute("UPDATE users SET last_seen=? WHERE id=?", (now, user["id"]))
                    conn.commit()
                    conn.close()
                    _clear_attempts(ip)   # نجاح: امسح سجل المحاولات
                    session.clear()
                    session.permanent = True
                    set_session(user, org)
                    session["_last_active"] = datetime.now().isoformat()
                    return redirect(url_for("dashboard"))

    return render_template("login.html", error=error)


# ══════════════════════════════════════════
# الملف الشخصي
# ══════════════════════════════════════════
@app.route("/profile", methods=["GET", "POST"])
@login_required
def profile():
    user_id = session["user_id"]
    org_id  = session["org_id"]
    error   = None
    success = False

    conn = get_connection()
    c    = conn.cursor()

    if request.method == "POST":
        full_name = request.form.get("full_name", "").strip()
        phone     = request.form.get("phone", "").strip()
        photo_file = request.files.get("photo")

        if not full_name:
            error = "الاسم الكامل مطلوب"
        else:
            photo_filename = None
            if photo_file and photo_file.filename and allowed_file(photo_file.filename) and allowed_file_content(photo_file):
                import uuid as _uuid2
                ext = photo_file.filename.rsplit(".", 1)[1].lower()
                photo_filename = f"user_{_uuid2.uuid4().hex}.{ext}"
                photo_file.save(os.path.join(app.config["UPLOAD_FOLDER"], photo_filename))

            if photo_filename:
                c.execute("UPDATE users SET full_name=?, phone=?, photo=? WHERE id=? AND org_id=?",
                          (full_name, phone, photo_filename, user_id, org_id))
            else:
                c.execute("UPDATE users SET full_name=?, phone=? WHERE id=? AND org_id=?",
                          (full_name, phone, user_id, org_id))
            conn.commit()
            session["full_name"] = full_name
            success = True

    c.execute("SELECT full_name, username, role, phone, photo, email, created_at FROM users WHERE id=?", (user_id,))
    user = dict(c.fetchone())
    conn.close()
    return render_template("profile.html", user=user, error=error, success=success)


# ══════════════════════════════════════════
# تغيير كلمة المرور (لكل المستخدمين)
# ══════════════════════════════════════════
@app.route("/change_password", methods=["GET", "POST"])
@login_required
def change_password():
    error = None
    success = False
    if request.method == "POST":
        current  = request.form.get("current_password", "")
        new_pass = request.form.get("new_password", "").strip()
        confirm  = request.form.get("confirm_password", "").strip()

        conn = get_connection()
        c = conn.cursor()
        try:
            c.execute("SELECT password FROM users WHERE id=?", (session["user_id"],))
            user = c.fetchone()
            if not user:
                error = "لم يتم العثور على حسابك"
            elif not check_password_hash(user["password"], current):
                error = "كلمة المرور الحالية غير صحيحة"
            elif len(new_pass) < 8:
                error = "كلمة المرور الجديدة يجب أن تكون 8 أحرف على الأقل"
            elif new_pass != confirm:
                error = "كلمتا المرور الجديدتان غير متطابقتين"
            else:
                hashed = generate_password_hash(new_pass)
                c.execute("UPDATE users SET password=? WHERE id=?",
                          (hashed, session["user_id"]))
                conn.commit()
                success = True
        except Exception as e:
            error = f"حدث خطأ: {str(e)}"
        finally:
            conn.close()

    return render_template("change_password.html", error=error, success=success)


@app.route("/delete_account", methods=["POST"])
@login_required
def delete_account():
    """حذف الحساب — الأدمن يحذف المؤسسة كاملة مع جميع بياناتها"""
    if session.get("role") != "admin":
        return jsonify({"ok": False, "err": "forbidden"}), 403

    password = request.form.get("password", "")
    org_id   = session["org_id"]
    user_id  = session["user_id"]

    conn = get_connection()
    c    = conn.cursor()
    try:
        # تحقق من كلمة السر
        c.execute("SELECT password FROM users WHERE id=?", (user_id,))
        row = c.fetchone()
        if not row or not check_password_hash(row["password"], password):
            conn.close()
            return jsonify({"ok": False, "err": "wrong_password"}), 400

        # ── حذف جميع بيانات المؤسسة ──────────────────────────
        # message_reads لمستخدمي المؤسسة
        c.execute("SELECT id FROM users WHERE org_id=?", (org_id,))
        user_ids = [r["id"] for r in c.fetchall()]
        if user_ids:
            placeholders = ",".join("?" * len(user_ids))
            c.execute(f"DELETE FROM message_reads WHERE user_id IN ({placeholders})", user_ids)
            c.execute(f"DELETE FROM org_message_reads WHERE user_id IN ({placeholders})", user_ids)

        # بنود الفواتير (لا تحتوي org_id مباشرة)
        c.execute("SELECT id FROM incoming_invoices WHERE org_id=?", (org_id,))
        inv_ids = [r["id"] for r in c.fetchall()]
        if inv_ids:
            placeholders = ",".join("?" * len(inv_ids))
            c.execute(f"DELETE FROM incoming_invoice_items WHERE invoice_id IN ({placeholders})", inv_ids)

        c.execute("SELECT id FROM outgoing_invoices WHERE org_id=?", (org_id,))
        out_ids = [r["id"] for r in c.fetchall()]
        if out_ids:
            placeholders = ",".join("?" * len(out_ids))
            c.execute(f"DELETE FROM outgoing_invoice_items WHERE invoice_id IN ({placeholders})", out_ids)

        # الجداول ذات org_id مباشرة
        for table in [
            "program_records", "beneficiaries", "programs", "workers",
            "products", "stock_batches",
            "incoming_invoices", "outgoing_invoices",
            "messages",
        ]:
            c.execute(f"DELETE FROM {table} WHERE org_id=?", (org_id,))

        # رسائل شبكة التواصل (from أو to)
        c.execute("DELETE FROM org_messages WHERE from_org_id=? OR to_org_id=?", (org_id, org_id))

        # طلبات مشاركة المستفيدين
        c.execute("DELETE FROM beneficiary_share_requests WHERE requester_org_id=? OR owner_org_id=?", (org_id, org_id))

        # المستخدمون ثم المؤسسة
        c.execute("DELETE FROM users WHERE org_id=?", (org_id,))
        c.execute("DELETE FROM organizations WHERE id=?", (org_id,))

        conn.commit()
    except Exception as e:
        conn.rollback()
        conn.close()
        return jsonify({"ok": False, "err": str(e)}), 500
    finally:
        try:
            conn.close()
        except Exception:
            pass

    # تسجيل الخروج
    session.clear()
    return jsonify({"ok": True})


# ══════════════════════════════════════════
# صفحة المراقب العام
# ══════════════════════════════════════════
@app.route("/overview")
@observer_required
def overview():
    org_id = session["org_id"]
    conn = get_connection()
    c = conn.cursor()

    # المخزون
    c.execute("""
        SELECT p.name, p.unit,
            COALESCE((SELECT SUM(quantity_remaining)
                      FROM stock_batches WHERE product_id=p.id AND org_id=?),0) as qty
        FROM products p WHERE p.org_id=? ORDER BY qty DESC
    """, (org_id, org_id))
    inventory = [dict(r) for r in c.fetchall()]

    # آخر 20 فاتورة صرف
    c.execute("""
        SELECT inv.invoice_number, inv.invoice_date, inv.beneficiary,
               inv.created_by, COUNT(oi.id) as items,
               COALESCE(SUM(oi.total_price),0) as total
        FROM outgoing_invoices inv
        LEFT JOIN outgoing_invoice_items oi ON oi.invoice_id=inv.id
        WHERE inv.org_id=?
        GROUP BY inv.id ORDER BY inv.invoice_date DESC LIMIT 20
    """, (org_id,))
    recent_outgoing = [dict(r) for r in c.fetchall()]

    # البرامج
    c.execute("SELECT * FROM programs WHERE org_id=? AND is_active=1 ORDER BY name", (org_id,))
    programs = [dict(r) for r in c.fetchall()]

    # إحصاءات
    c.execute("SELECT COUNT(*) FROM beneficiaries WHERE org_id=?", (org_id,))
    benef_count = c.fetchone()[0]
    c.execute("SELECT COUNT(*) FROM outgoing_invoices WHERE org_id=?", (org_id,))
    out_count = c.fetchone()[0]

    conn.close()
    return render_template("overview.html",
        inventory=inventory, recent_outgoing=recent_outgoing,
        programs=programs, benef_count=benef_count, out_count=out_count)


@app.route("/logout")
def logout():
    # تصفير last_seen عند تسجيل الخروج
    if "user_id" in session:
        try:
            conn = get_connection()
            c = conn.cursor()
            c.execute("UPDATE users SET last_seen=NULL WHERE id=?", (session["user_id"],))
            conn.commit()
            conn.close()
        except Exception:
            pass
    session.clear()
    return redirect(url_for("login"))


# ══════════════════════════════════════════
# لوحة التحكم
# ══════════════════════════════════════════
@app.route("/dashboard")
@login_required
def dashboard():
    org_id = session["org_id"]
    conn = get_connection()
    c = conn.cursor()
    c.execute("SELECT COUNT(*) FROM products WHERE org_id=?", (org_id,))
    products_count = c.fetchone()[0]
    c.execute("SELECT COALESCE(SUM(quantity_remaining),0) FROM stock_batches WHERE org_id=?", (org_id,))
    total_stock = c.fetchone()[0]
    c.execute("SELECT COUNT(*) FROM beneficiaries WHERE org_id=?", (org_id,))
    beneficiaries_count = c.fetchone()[0]
    c.execute("SELECT COUNT(*) FROM incoming_invoices WHERE org_id=?", (org_id,))
    incoming_count = c.fetchone()[0]
    c.execute("SELECT COUNT(*) FROM outgoing_invoices WHERE org_id=?", (org_id,))
    outgoing_count = c.fetchone()[0]
    pending_users = []
    if session.get("role") == "admin":
        c.execute(
            "SELECT id, full_name, role, id_number, photo, created_at FROM users WHERE org_id=? AND status='pending' ORDER BY id DESC",
            (org_id,)
        )
        pending_users = [dict(r) for r in c.fetchall()]

    # تنبيهات المخزون المنخفض/النافد
    c.execute("""
        SELECT p.id, p.name, p.unit,
               COALESCE(SUM(sb.quantity_remaining), 0) AS qty
        FROM products p
        LEFT JOIN stock_batches sb ON sb.product_id = p.id AND sb.org_id = p.org_id
        WHERE p.org_id = ?
        GROUP BY p.id
        HAVING qty <= 10
        ORDER BY qty ASC
    """, (org_id,))
    low_stock = [dict(r) for r in c.fetchall()]

    conn.close()
    return render_template("dashboard.html",
        products_count=products_count, total_stock=total_stock,
        beneficiaries_count=beneficiaries_count,
        incoming_count=incoming_count, outgoing_count=outgoing_count,
        pending_users=pending_users,
        low_stock=low_stock)



# ══════════════════════════════════════════
# بحث لوحة التحكم (صنف / مشروع)
# ══════════════════════════════════════════
@app.route("/api/dashboard_search")
@login_required
def dashboard_search():
    if not session.get("role"):
        return jsonify({"error": "غير مصرح"})
    org_id = session["org_id"]
    q      = request.args.get("q", "").strip()
    stype  = request.args.get("type", "product")  # product | project
    if not q:
        return jsonify({"results": []})

    conn = get_connection()
    c    = conn.cursor()
    out  = {}

    if stype == "product":
        # ── بحث الصنف ──
        c.execute(
            "SELECT id, name, unit FROM products WHERE org_id=? AND name LIKE ? ORDER BY name LIMIT 10",
            (org_id, f"%{q}%")
        )
        products = [dict(r) for r in c.fetchall()]
        items = []
        for p in products:
            pid = p["id"]
            # المخزون الحالي
            c.execute("SELECT COALESCE(SUM(quantity_remaining),0) FROM stock_batches WHERE org_id=? AND product_id=?", (org_id, pid))
            stock = c.fetchone()[0]
            # إجمالي المشتريات
            c.execute("""
                SELECT COALESCE(SUM(ii.quantity),0) AS total_in,
                       COUNT(DISTINCT ii.invoice_id) AS inv_count
                FROM incoming_invoice_items ii
                JOIN incoming_invoices inv ON inv.id=ii.invoice_id AND inv.org_id=?
                WHERE ii.product_id=?
            """, (org_id, pid))
            r = c.fetchone()
            total_in = r["total_in"]; inv_count = r["inv_count"]
            # إجمالي الصرف
            c.execute("""
                SELECT COALESCE(SUM(oi.quantity),0) AS total_out,
                       COUNT(DISTINCT oi.invoice_id) AS out_count
                FROM outgoing_invoice_items oi
                JOIN outgoing_invoices outv ON outv.id=oi.invoice_id AND outv.org_id=?
                WHERE oi.product_id=?
            """, (org_id, pid))
            r2 = c.fetchone()
            total_out = r2["total_out"]; out_count = r2["out_count"]
            # آخر 5 حركات إدخال
            c.execute("""
                SELECT inv.supplier, inv.invoice_date, ii.quantity, ii.unit_price, ii.total_price
                FROM incoming_invoice_items ii
                JOIN incoming_invoices inv ON inv.id=ii.invoice_id AND inv.org_id=?
                WHERE ii.product_id=?
                ORDER BY inv.invoice_date DESC, inv.id DESC LIMIT 5
            """, (org_id, pid))
            in_rows = [dict(r) for r in c.fetchall()]
            # آخر 5 حركات صرف
            c.execute("""
                SELECT outv.beneficiary, outv.invoice_date, oi.quantity, oi.unit_price,
                       prog.name AS project_name
                FROM outgoing_invoice_items oi
                JOIN outgoing_invoices outv ON outv.id=oi.invoice_id AND outv.org_id=?
                LEFT JOIN programs prog ON prog.id=outv.program_id
                WHERE oi.product_id=?
                ORDER BY outv.invoice_date DESC, outv.id DESC LIMIT 5
            """, (org_id, pid))
            out_rows = [dict(r) for r in c.fetchall()]
            items.append({
                "name": p["name"], "unit": p["unit"] or "",
                "stock": stock, "total_in": total_in, "inv_count": inv_count,
                "total_out": total_out, "out_count": out_count,
                "in_rows": in_rows, "out_rows": out_rows
            })
        out = {"type": "product", "items": items}

    elif stype == "project":
        # ── بحث المشروع ──
        c.execute(
            "SELECT id, name, is_active FROM programs WHERE org_id=? AND name LIKE ? ORDER BY name LIMIT 10",
            (org_id, f"%{q}%")
        )
        programs = [dict(r) for r in c.fetchall()]
        items = []
        for prog in programs:
            pid = prog["id"]
            # فواتير الصرف المرتبطة بالمشروع
            c.execute("""
                SELECT outv.invoice_date, outv.beneficiary, outv.grand_total, outv.notes
                FROM outgoing_invoices outv
                WHERE outv.org_id=? AND outv.program_id=?
                ORDER BY outv.invoice_date DESC LIMIT 10
            """, (org_id, pid))
            inv_rows = [dict(r) for r in c.fetchall()]
            # إجمالي الصرف
            c.execute("""
                SELECT COALESCE(SUM(grand_total),0) AS total, COUNT(*) AS cnt
                FROM outgoing_invoices WHERE org_id=? AND program_id=?
            """, (org_id, pid))
            r = c.fetchone()
            total_dist = r["total"]; dist_count = r["cnt"]
            # سجلات البرنامج (program_records)
            c.execute("""
                SELECT pr.benefit_date, b.full_name, pr.benefit_type, pr.quantity, pr.notes
                FROM program_records pr
                JOIN beneficiaries b ON b.id=pr.beneficiary_id
                WHERE pr.org_id=? AND pr.program_id=?
                ORDER BY pr.benefit_date DESC LIMIT 10
            """, (org_id, pid))
            prog_rows = [dict(r) for r in c.fetchall()]
            c.execute("SELECT COUNT(*) FROM program_records WHERE org_id=? AND program_id=?", (org_id, pid))
            prog_count = c.fetchone()[0]
            items.append({
                "name": prog["name"], "is_active": prog["is_active"],
                "total_dist": total_dist, "dist_count": dist_count,
                "inv_rows": inv_rows,
                "prog_count": prog_count, "prog_rows": prog_rows
            })
        out = {"type": "project", "items": items}

    conn.close()
    return jsonify(out)

# ══════════════════════════════════════════
# تنبيهات المخزون المنخفض (API)
# ══════════════════════════════════════════
@app.route("/api/low_stock")
@login_required
def api_low_stock():
    """إرجاع قائمة الأصناف ذات المخزون المنخفض أو النافد بصيغة JSON"""
    if session.get("role") not in ("admin", "accountant", "observer"):
        return jsonify([])
    org_id = session["org_id"]
    conn = get_connection()
    c = conn.cursor()
    c.execute("""
        SELECT p.id, p.name, p.unit,
               COALESCE(SUM(sb.quantity_remaining), 0) AS qty
        FROM products p
        LEFT JOIN stock_batches sb ON sb.product_id = p.id AND sb.org_id = p.org_id
        WHERE p.org_id = ?
        GROUP BY p.id
        HAVING qty <= 10
        ORDER BY qty ASC
    """, (org_id,))
    items = [{"id": r["id"], "name": r["name"], "unit": r["unit"] or "", "qty": r["qty"]} for r in c.fetchall()]
    conn.close()
    return jsonify(items)


# ══════════════════════════════════════════
# موافقة / رفض الموظفين
# ══════════════════════════════════════════
@app.route("/approve_user/<int:id>")
@admin_required
def approve_user(id):
    conn = get_connection()
    c = conn.cursor()
    c.execute("UPDATE users SET status='approved' WHERE id=? AND org_id=?", (id, session["org_id"]))
    conn.commit()
    conn.close()
    flash("✅ تمت الموافقة على الحساب", "success")
    return redirect(url_for("dashboard"))


@app.route("/reject_user/<int:id>")
@admin_required
def reject_user(id):
    conn = get_connection()
    c = conn.cursor()
    c.execute("UPDATE users SET status='rejected' WHERE id=? AND org_id=?", (id, session["org_id"]))
    conn.commit()
    conn.close()
    flash("تم رفض الحساب", "warning")
    return redirect(url_for("dashboard"))


# ══════════════════════════════════════════
# إدارة المستخدمين
# ══════════════════════════════════════════
@app.route("/users")
@admin_required
def users_list():
    conn = get_connection()
    c = conn.cursor()
    # جلب الأعمدة الأساسية فقط (last_seen قد لا يكون موجوداً في قاعدة بيانات قديمة)
    c.execute(
        "SELECT id, username, full_name, id_number, role, status, photo, created_at FROM users WHERE org_id=? ORDER BY id",
        (session["org_id"],)
    )
    rows = c.fetchall()

    # محاولة جلب last_seen بشكل منفصل آمن
    last_seen_map = {}
    try:
        c.execute("SELECT id, last_seen FROM users WHERE org_id=?", (session["org_id"],))
        for r in c.fetchall():
            last_seen_map[r["id"]] = r["last_seen"]
    except Exception:
        pass  # العمود غير موجود - لا مشكلة

    conn.close()

    def seconds_since(last_seen_str):
        if not last_seen_str:
            return 999999
        try:
            ls = datetime.strptime(last_seen_str, "%Y-%m-%d %H:%M:%S")
            return int((datetime.now() - ls).total_seconds())
        except Exception:
            return 999999

    users = []
    for r in rows:
        last_seen_val = last_seen_map.get(r["id"], "")
        d = {
            "id":        r["id"],
            "username":  r["username"],
            "full_name": r["full_name"] or "",
            "id_number": r["id_number"] or "",
            "role":      r["role"],
            "status":    r["status"],
            "photo":     r["photo"] or "",
            "created_at":r["created_at"] or "",
            "last_seen": last_seen_val or "",
            "last_seen_seconds": seconds_since(last_seen_val),
        }
        users.append(d)

    return render_template("users.html", users=users)


@app.route("/add_user", methods=["GET", "POST"])
@admin_required
def add_user():
    error = None
    if request.method == "POST":
        if not validate_csrf():
            flash("طلب غير صالح (CSRF). أعد المحاولة.", "danger")
            return redirect(url_for("add_user"))
        full_name = request.form.get("full_name", "").strip()
        id_number = request.form.get("id_number", "").strip()
        username  = request.form.get("username", "").strip()
        password  = request.form.get("password", "")
        role      = request.form.get("role", "")

        if not role or role not in ("accountant", "data_entry", "observer"):
            error = "يرجى اختيار نوع الحساب (محاسب، مدخل بيانات، أو مراقب عام)"
        elif not all([full_name, id_number, username, password]):
            error = "يرجى تعبئة جميع الحقول"
        elif len(password) < 8:
            error = "كلمة المرور 8 أحرف على الأقل"
        elif id_number and (not id_number.isdigit() or len(id_number) != 9):
            error = "رقم الهوية يجب أن يكون 9 أرقام"
        else:
            conn = get_connection()
            c = conn.cursor()
            c.execute("SELECT id FROM users WHERE username=?", (username,))
            if c.fetchone():
                error = "اسم المستخدم مأخوذ"
                conn.close()
            else:
                hashed = generate_password_hash(password)
                c.execute(
                    "INSERT INTO users (org_id, username, password, full_name, id_number, role, status, created_by) VALUES (?,?,?,?,?,?,?,?)",
                    (session["org_id"], username, hashed, full_name, id_number, role, "approved", session["user_id"])
                )
                conn.commit()
                conn.close()
                flash(f"✅ تم إضافة '{full_name}' بنجاح", "success")
                return redirect(url_for("users_list"))

    return render_template("add_user.html", error=error)


@app.route("/edit_user/<int:id>", methods=["GET", "POST"])
@admin_required
def edit_user(id):
    if id == session["user_id"]:
        flash("لا يمكنك تعديل حسابك من هنا", "warning")
        return redirect(url_for("users_list"))
    if request.method == "POST" and not validate_csrf():
        flash("طلب غير صالح (CSRF). أعد المحاولة.", "danger")
        return redirect(url_for("edit_user", id=id))

    conn = get_connection()
    c = conn.cursor()

    if request.method == "POST":
        full_name    = request.form.get("full_name", "").strip()
        role         = request.form.get("role", "")
        status       = request.form.get("status", "approved")
        new_password = request.form.get("new_password", "").strip()

        if new_password:
            if len(new_password) < 8:
                flash("كلمة المرور 8 أحرف على الأقل", "danger")
                c.execute("SELECT * FROM users WHERE id=? AND org_id=?", (id, session["org_id"]))
                user = dict(c.fetchone())
                conn.close()
                return render_template("edit_user.html", user=user)
            hashed = generate_password_hash(new_password)
            c.execute("UPDATE users SET full_name=?,role=?,status=?,password=? WHERE id=? AND org_id=?",
                      (full_name, role, status, hashed, id, session["org_id"]))
        else:
            c.execute("UPDATE users SET full_name=?,role=?,status=? WHERE id=? AND org_id=?",
                      (full_name, role, status, id, session["org_id"]))
        conn.commit()
        conn.close()
        flash("✅ تم تحديث بيانات المستخدم", "success")
        return redirect(url_for("users_list"))

    c.execute("SELECT * FROM users WHERE id=? AND org_id=?", (id, session["org_id"]))
    row = c.fetchone()
    conn.close()
    if not row:
        flash("المستخدم غير موجود", "danger")
        return redirect(url_for("users_list"))
    return render_template("edit_user.html", user=dict(row))


@app.route("/delete_user/<int:id>")
@admin_required
def delete_user(id):
    if id == session["user_id"]:
        flash("لا يمكنك حذف حسابك الخاص", "danger")
        return redirect(url_for("users_list"))
    conn = get_connection()
    c = conn.cursor()
    c.execute("DELETE FROM users WHERE id=? AND org_id=?", (id, session["org_id"]))
    conn.commit()
    conn.close()
    flash("تم حذف المستخدم", "warning")
    return redirect(url_for("users_list"))


# ══════════════════════════════════════════
# الأصناف
# ══════════════════════════════════════════
@app.route("/products")
@observer_required
def products():
    org_id = session["org_id"]
    conn = get_connection()
    c = conn.cursor()
    c.execute("""
        SELECT p.id, p.name, p.unit,
            COALESCE((SELECT SUM(quantity_remaining) FROM stock_batches WHERE product_id=p.id AND org_id=?),0),
            COALESCE((SELECT unit_price FROM stock_batches WHERE product_id=p.id ORDER BY id ASC  LIMIT 1),0),
            COALESCE((SELECT unit_price FROM stock_batches WHERE product_id=p.id ORDER BY id DESC LIMIT 1),0),
            p.last_modified
        FROM products p WHERE p.org_id=? ORDER BY p.id DESC
    """, (org_id, org_id))
    data = c.fetchall()
    conn.close()
    return render_template("products.html", products=data)


@app.route("/add_product", methods=["GET", "POST"])
@invoices_required
def add_product():
    if request.method == "POST":
        name = request.form.get("name", "").strip()
        unit = request.form.get("unit", "").strip()
        if not name or not unit:
            flash("يرجى تعبئة جميع الحقول", "danger")
            return render_template("add_product.html")
        conn = get_connection()
        c = conn.cursor()
        c.execute("INSERT INTO products (org_id, name, unit) VALUES (?,?,?)",
                  (session["org_id"], name, unit))
        conn.commit()
        conn.close()
        flash("✅ تمت إضافة الصنف", "success")
        return redirect(url_for("products"))
    return render_template("add_product.html")


@app.route("/edit_product/<int:id>", methods=["GET", "POST"])
@invoices_required
def edit_product(id):
    conn = get_connection()
    c = conn.cursor()
    if request.method == "POST":
        name = request.form.get("name", "").strip()
        unit = request.form.get("unit", "").strip()
        if not name or not unit:
            flash("يرجى تعبئة جميع الحقول", "danger")
        else:
            now = datetime.now().strftime("%Y/%m/%d")
            c.execute("UPDATE products SET name=?,unit=?,last_modified=? WHERE id=? AND org_id=?",
                      (name, unit, now, id, session["org_id"]))
            conn.commit()
            flash("✅ تم تعديل الصنف", "success")
            conn.close()
            return redirect(url_for("products"))
    c.execute("SELECT * FROM products WHERE id=? AND org_id=?", (id, session["org_id"]))
    product = c.fetchone()
    conn.close()
    if not product:
        flash("الصنف غير موجود", "danger")
        return redirect(url_for("products"))
    return render_template("edit_product.html", product=product)


@app.route("/delete_product/<int:id>")
@invoices_required
def delete_product(id):
    conn = get_connection()
    c = conn.cursor()
    c.execute("DELETE FROM stock_batches WHERE product_id=? AND org_id=?", (id, session["org_id"]))
    c.execute("DELETE FROM products WHERE id=? AND org_id=?", (id, session["org_id"]))
    conn.commit()
    conn.close()
    flash("تم حذف الصنف", "warning")
    return redirect(url_for("products"))


# ══════════════════════════════════════════
# المستفيدون
# ══════════════════════════════════════════
BENEF_KEYS = ["id", "org_id", "seq_num", "full_name", "gender", "phone",
              "address", "address2", "camp_name", "id_number", "family_size",
              "marital_status", "children_count", "wife_pregnant", "wife_nursing",
              "has_orphans", "orphans_count", "notes", "created_at", "beneficiary_type",
              "camp_manager_name", "camp_coordinator", "camp_coordinator_phone",
              "camp_address", "camp_family_count",
              "wife_name", "personal_photo", "death_cert_image",
              "guardianship_image", "guardian_whatsapp", "guardian_name", "guardian_id_number"]


@app.route("/beneficiaries")
@data_entry_view_required
def beneficiaries():
    org_id = session["org_id"]
    conn = get_connection()
    c = conn.cursor()
    # جلب المستفيدين عبر جدول الربط
    try:
        c.execute("""
            SELECT DISTINCT b.* FROM beneficiaries b
            LEFT JOIN org_beneficiary_links obl ON b.id = obl.beneficiary_id AND obl.org_id = ?
            WHERE b.org_id = ? OR obl.org_id = ?
            ORDER BY b.id ASC
        """, (org_id, org_id, org_id))
    except Exception:
        # fallback: الجدول الجديد لم يُنشأ بعد
        c.execute("SELECT * FROM beneficiaries WHERE org_id=? ORDER BY id ASC", (org_id,))
    rows = c.fetchall()

    # جلب أفراد الأسرة لكل المستفيدين
    benef_ids = tuple(r["id"] for r in rows)
    if benef_ids:
        placeholders = ",".join("?" * len(benef_ids))
        c.execute(f"SELECT * FROM beneficiary_family_members WHERE beneficiary_id IN ({placeholders}) ORDER BY member_type, id",
                  benef_ids)
    else:
        c.execute("SELECT * FROM beneficiary_family_members WHERE 1=0")
    fam_rows = c.fetchall()
    conn.close()

    # بناء قاموس أفراد الأسرة مرتبة بـ beneficiary_id
    family_map = {}
    for fr in fam_rows:
        bid = fr["beneficiary_id"]
        if bid not in family_map:
            family_map[bid] = []
        family_map[bid].append(dict(fr))

    data = []
    for i, row in enumerate(rows, 1):
        d = row_to_dict(row, BENEF_KEYS)
        d["seq_num"] = i
        for k, v in [("address2",""), ("camp_name",""), ("camp_manager_name",""),
                     ("camp_coordinator",""), ("camp_coordinator_phone",""),
                     ("camp_address",""), ("camp_family_count",""),
                     ("marital_status","married"), ("gender","male"),
                     ("children_count",0), ("wife_pregnant",0), ("wife_nursing",0),
                     ("has_orphans",0), ("orphans_count",0), ("beneficiary_type","person"),
                     ("wife_name",""), ("personal_photo",""), ("death_cert_image",""),
                     ("guardianship_image",""), ("guardian_whatsapp",""),
                     ("guardian_name",""), ("guardian_id_number","")]:
            d.setdefault(k, v)
        d["family_members"] = family_map.get(d["id"], [])
        data.append(d)

    # بناء الهيكل الهرمي: مخيمات → أشخاص تحتها
    camps = [d for d in data if d.get("beneficiary_type") == "camp"]
    persons = [d for d in data if d.get("beneficiary_type") != "camp"]

    # ربط الأشخاص بمخيماتهم
    camp_names = {c["camp_name"].strip().lower(): c for c in camps if c["camp_name"]}
    linked_ids = set()
    for camp in camps:
        camp_key = camp["camp_name"].strip().lower() if camp["camp_name"] else ""
        camp["persons"] = []
        for p in persons:
            if p.get("camp_name","").strip().lower() == camp_key and camp_key:
                camp["persons"].append(p)
                linked_ids.add(p["id"])

    unlinked = [p for p in persons if p["id"] not in linked_ids]

    # ── قوائم التصنيف ──
    from datetime import date as _date
    def _calc_age(bd_str):
        if not bd_str: return None
        try:
            bd = __import__('datetime').datetime.strptime(bd_str, '%Y-%m-%d').date()
            today = _date.today()
            return today.year - bd.year - ((today.month, today.day) < (bd.month, bd.day))
        except: return None

    orphan_list = []
    pregnant_list = []
    nursing_list = []
    persons_sorted = sorted(persons, key=lambda x: x.get('full_name',''))
    for p in persons:
        if p.get('wife_pregnant'): pregnant_list.append(p)
        if p.get('wife_nursing'):  nursing_list.append(p)
        if p.get('has_orphans'):
            for fm in p.get('family_members', []):
                if fm.get('is_orphan'):
                    orphan_list.append({
                        'name':           fm.get('full_name',''),
                        'birth_date':     fm.get('birth_date',''),
                        'age':            _calc_age(fm.get('birth_date','')),
                        'guardian_name':  p.get('guardian_name','') or p.get('full_name',''),
                        'guardian_phone': p.get('guardian_whatsapp','') or p.get('phone',''),
                        'camp_name':      p.get('camp_name',''),
                        'parent_name':    p.get('full_name',''),
                    })
    orphan_list.sort(key=lambda x: x.get('name',''))
    pregnant_list.sort(key=lambda x: x.get('full_name',''))
    nursing_list.sort(key=lambda x: x.get('full_name',''))

    # قائمة المرضى من beneficiary_health
    try:
        conn2 = get_connection()
        c2 = conn2.cursor()
        c2.execute("""
            SELECT bh.beneficiary_id, bh.condition_type, bh.notes, bh.report_path,
                   b.full_name, b.camp_name, b.phone
            FROM beneficiary_health bh
            JOIN beneficiaries b ON bh.beneficiary_id = b.id
            WHERE b.org_id = ?
            ORDER BY b.full_name
        """, (org_id,))
        sick_list = [dict(r) for r in c2.fetchall()]
        conn2.close()
    except Exception:
        sick_list = []

    return render_template("beneficiaries.html",
                           beneficiaries=data, camps=camps,
                           unlinked_persons=unlinked, persons=persons,
                           persons_sorted=persons_sorted,
                           orphan_list=orphan_list,
                           pregnant_list=pregnant_list,
                           nursing_list=nursing_list,
                           sick_list=sick_list)


@app.route("/add_beneficiary", methods=["GET", "POST"])
@data_entry_required
def add_beneficiary():
    if request.method == "POST":
        full_name      = request.form.get("full_name", "").strip()
        gender         = request.form.get("gender", "male")
        phone          = request.form.get("phone", "").strip()
        address        = request.form.get("address", "").strip()
        address2       = request.form.get("address2", "").strip()
        beneficiary_type = request.form.get("beneficiary_type", "person").strip()
        # للمخيم: اسم المخيم يُخزن في full_name وcamp_name معاً؛ للشخص: linked_camp هو اسم المخيم
        if beneficiary_type == 'camp':
            camp_name = full_name  # نفس الاسم في العمودين حتى يبقى الربط صحيحاً
        else:
            camp_name = request.form.get("linked_camp", "").strip()
        camp_manager_name     = request.form.get("camp_manager_name", "").strip()
        camp_coordinator      = request.form.get("camp_coordinator", "").strip()
        camp_coordinator_phone= request.form.get("camp_coordinator_phone", "").strip()
        camp_address          = request.form.get("camp_address", "").strip()
        camp_family_count     = request.form.get("camp_family_count", "").strip() or None
        id_number      = request.form.get("id_number", "").strip()
        family_size    = request.form.get("family_size", "1").strip()
        marital_status = request.form.get("marital_status", "married").strip()
        children_count = request.form.get("children_count", "0").strip()
        wife_pregnant  = 1 if request.form.get("wife_pregnant") else 0
        wife_nursing   = 1 if request.form.get("wife_nursing") else 0
        has_orphans    = 1 if request.form.get("has_orphans") else 0
        orphans_count    = request.form.get("orphans_count", "0").strip()
        notes            = request.form.get("notes", "").strip()

        errors = []
        if beneficiary_type == 'person' and len(full_name.split()) < 2:
            errors.append("الاسم يجب أن يكون رباعياً (كلمتان على الأقل)")
        elif beneficiary_type == 'camp' and not full_name.strip():
            errors.append("يرجى إدخال اسم المخيم")
        if id_number and (not id_number.isdigit() or len(id_number) != 9):
            errors.append("رقم الهوية يجب أن يكون 9 أرقام فقط")
        if phone and (not phone.isdigit() or len(phone) != 10 or not phone.startswith("05")):
            errors.append("رقم الجوال يجب أن يكون 10 أرقام ويبدأ بـ 05")

        if errors:
            for e in errors:
                flash(e, "danger")
            return render_template("add_beneficiary.html", form_data=request.form)

        try:
            family_size    = int(family_size)
            children_count = int(children_count)
            orphans_count  = int(orphans_count) if has_orphans else 0
        except ValueError:
            family_size = 1; children_count = 0; orphans_count = 0

        conn = get_connection()
        c = conn.cursor()

        # ── المنطق الجديد: رقم الهوية محور الربط المركزي ──
        org_id = session["org_id"]

        # 1. رقم الهوية إلزامي للأشخاص
        if beneficiary_type == 'person' and not id_number:
            conn.close()
            flash("⚠️ رقم الهوية إلزامي — لا يمكن إضافة مستفيد بدون هوية", "danger")
            return render_template("add_beneficiary.html", form_data=request.form,
                                   prefill_camp=request.form.get("linked_camp",""))

        # 2. هل رقم الهوية موجود في النظام؟
        new_benef_id = None
        if id_number and beneficiary_type == 'person':
            c.execute("SELECT id, full_name, org_id FROM beneficiaries WHERE TRIM(id_number)=TRIM(?)", (id_number,))
            existing = c.fetchone()
            if existing:
                # تحقق: هل هو مرتبط بهذه المؤسسة مسبقاً؟
                c.execute("SELECT id FROM org_beneficiary_links WHERE org_id=? AND beneficiary_id=?",
                          (org_id, existing["id"]))
                already = c.fetchone()
                if already:
                    conn.close()
                    flash(f"⚠️ رقم الهوية {id_number} مسجل مسبقاً في مؤسستك للمستفيد: {existing['full_name']}", "danger")
                    return redirect(url_for("add_beneficiary"))
                else:
                    # ربط المستفيد الموجود بهذه المؤسسة تلقائياً
                    try:
                        c.execute("INSERT OR IGNORE INTO org_beneficiary_links (org_id, beneficiary_id) VALUES (?,?)",
                                  (org_id, existing["id"]))
                        conn.commit()
                    except Exception:
                        pass
                    conn.close()
                    flash(f"✅ المستفيد {existing['full_name']} موجود في النظام — تم ربطه بمؤسستك تلقائياً", "success")
                    return redirect(url_for("beneficiaries"))

        c.execute("SELECT COUNT(*) FROM beneficiaries WHERE org_id=?", (org_id,))
        next_seq = c.fetchone()[0] + 1

        # حقول شخص
        wife_name           = request.form.get("wife_name", "").strip()
        guardian_name       = request.form.get("guardian_name", "").strip()
        guardian_id_number  = request.form.get("guardian_id_number", "").strip()
        guardian_whatsapp   = request.form.get("guardian_whatsapp", "").strip()

        def save_upload(field, prefix):
            f2 = request.files.get(field)
            if f2 and f2.filename:
                ext = f2.filename.rsplit(".", 1)[-1].lower()
                fname = f"{prefix}_{org_id}_{next_seq}.{ext}"
                f2.save(os.path.join(app.config["UPLOAD_FOLDER"], fname))
                return fname
            return None

        personal_photo      = save_upload("personal_photo", "bphoto")
        death_cert_image    = save_upload("death_cert_image", "bdeath")
        guardianship_image  = save_upload("guardianship_image", "bguard")

        try:
            c.execute(
                """INSERT INTO beneficiaries
                   (org_id, seq_num, full_name, gender, phone, address, address2, camp_name,
                    id_number, family_size, marital_status, children_count,
                    wife_pregnant, wife_nursing, has_orphans, orphans_count, notes, beneficiary_type,
                    camp_manager_name, camp_coordinator, camp_coordinator_phone, camp_address, camp_family_count,
                    wife_name, personal_photo, death_cert_image, guardianship_image,
                    guardian_whatsapp, guardian_name, guardian_id_number)
                   VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
                (org_id, next_seq, full_name, gender, phone, address, address2,
                 camp_name, id_number, family_size, marital_status, children_count,
                 wife_pregnant, wife_nursing, has_orphans, orphans_count, notes, beneficiary_type,
                 camp_manager_name, camp_coordinator, camp_coordinator_phone, camp_address, camp_family_count,
                 wife_name, personal_photo, death_cert_image, guardianship_image,
                 guardian_whatsapp, guardian_name, guardian_id_number)
            )
            new_benef_id = c.lastrowid
            # ربط المستفيد الجديد بالمؤسسة
            try:
                c.execute("INSERT OR IGNORE INTO org_beneficiary_links (org_id, beneficiary_id) VALUES (?,?)",
                          (org_id, new_benef_id))
            except Exception:
                pass
        except Exception as ex:
            conn.close()
            if "UNIQUE" in str(ex):
                flash(f"⚠️ رقم الهوية {id_number} مسجل مسبقاً في النظام لمستفيد آخر", "danger")
            else:
                flash(f"خطأ: {ex}", "danger")
            return render_template("add_beneficiary.html", form_data=request.form,
                                   prefill_camp=request.form.get("linked_camp",""))

        # حفظ أفراد الأسرة (أطفال)
        child_names  = request.form.getlist("child_name[]")
        child_dates  = request.form.getlist("child_birth_date[]")
        child_orphan = request.form.getlist("child_is_orphan[]")
        child_files  = request.files.getlist("child_birth_cert[]")
        for idx, cname in enumerate(child_names):
            cname = cname.strip()
            if not cname:
                continue
            cdate   = child_dates[idx] if idx < len(child_dates) else ""
            corphan = 1 if str(idx) in child_orphan or (idx < len(child_orphan) and child_orphan[idx]) else 0
            cimg    = None
            if idx < len(child_files) and child_files[idx].filename:
                ext2 = child_files[idx].filename.rsplit(".", 1)[-1].lower()
                fn2  = f"child_{session['org_id']}_{new_benef_id}_{idx}.{ext2}"
                child_files[idx].save(os.path.join(app.config["UPLOAD_FOLDER"], fn2))
                cimg = fn2
            c.execute(
                "INSERT INTO beneficiary_family_members (beneficiary_id, org_id, member_type, full_name, birth_date, is_orphan, birth_cert_img) VALUES (?,?,?,?,?,?,?)",
                (new_benef_id, session["org_id"], "child", cname, cdate or None, corphan, cimg)
            )

        conn.commit()
        conn.close()
        flash("✅ تمت إضافة المستفيد بنجاح", "success")
        return redirect(url_for("beneficiaries"))

    prefill_camp = request.args.get("linked_camp", "")
    return render_template("add_beneficiary.html", form_data={}, prefill_camp=prefill_camp)



# ══════════════════════════════════════════
# أفراد أسرة المستفيد
# ══════════════════════════════════════════
@app.route("/beneficiary/<int:bid>/family/add", methods=["POST"])
@data_entry_required
def family_member_add(bid):
    if not validate_csrf(): return redirect(url_for("beneficiaries"))
    org_id = session["org_id"]
    conn = get_connection(); c = conn.cursor()
    c.execute("SELECT id FROM beneficiaries WHERE id=? AND org_id=?", (bid, org_id))
    if not c.fetchone():
        conn.close(); return redirect(url_for("beneficiaries"))
    mtype     = request.form.get("member_type","child")
    full_name = request.form.get("full_name","").strip()
    birth_date= request.form.get("birth_date","").strip()
    is_orphan = 1 if request.form.get("is_orphan") else 0
    cimg = None
    f2 = request.files.get("birth_cert_img")
    if f2 and f2.filename:
        ext = f2.filename.rsplit(".",1)[-1].lower()
        fn  = f"fam_{org_id}_{bid}_{mtype[:1]}.{ext}"
        f2.save(os.path.join(app.config["UPLOAD_FOLDER"], fn)); cimg = fn
    if full_name:
        c.execute("INSERT INTO beneficiary_family_members (beneficiary_id,org_id,member_type,full_name,birth_date,is_orphan,birth_cert_img) VALUES (?,?,?,?,?,?,?)",
                  (bid, org_id, mtype, full_name, birth_date or None, is_orphan, cimg))
        conn.commit()
    conn.close()
    return redirect(url_for("beneficiaries"))

@app.route("/beneficiary/family/<int:fid>/delete", methods=["POST"])
@data_entry_required
def family_member_delete(fid):
    if not validate_csrf(): return redirect(url_for("beneficiaries"))
    org_id = session["org_id"]
    conn = get_connection(); c = conn.cursor()
    c.execute("DELETE FROM beneficiary_family_members WHERE id=? AND org_id=?", (fid, org_id))
    conn.commit(); conn.close()
    return redirect(url_for("beneficiaries"))

@app.route("/edit_beneficiary/<int:id>", methods=["GET", "POST"])
@data_entry_required
def edit_beneficiary(id):
    org_id = session["org_id"]
    conn = get_connection()
    c = conn.cursor()

    if request.method == "POST":
        full_name      = request.form.get("full_name", "").strip()
        gender         = request.form.get("gender", "male")
        phone          = request.form.get("phone", "").strip()
        address        = request.form.get("address", "").strip()
        address2       = request.form.get("address2", "").strip()
        beneficiary_type = request.form.get("beneficiary_type", "person").strip()
        # للمخيم: اسم المخيم يُخزن في full_name وcamp_name معاً؛ للشخص: linked_camp هو اسم المخيم
        if beneficiary_type == 'camp':
            camp_name = full_name  # نفس الاسم في العمودين حتى يبقى الربط صحيحاً
        else:
            camp_name = request.form.get("linked_camp", "").strip()
        camp_manager_name     = request.form.get("camp_manager_name", "").strip()
        camp_coordinator      = request.form.get("camp_coordinator", "").strip()
        camp_coordinator_phone= request.form.get("camp_coordinator_phone", "").strip()
        camp_address          = request.form.get("camp_address", "").strip()
        camp_family_count     = request.form.get("camp_family_count", "").strip() or None
        id_number      = request.form.get("id_number", "").strip()
        family_size    = request.form.get("family_size", "1").strip()
        marital_status = request.form.get("marital_status", "married").strip()
        children_count = request.form.get("children_count", "0").strip()
        wife_pregnant  = 1 if request.form.get("wife_pregnant") else 0
        wife_nursing   = 1 if request.form.get("wife_nursing") else 0
        has_orphans    = 1 if request.form.get("has_orphans") else 0
        orphans_count    = request.form.get("orphans_count", "0").strip()
        notes            = request.form.get("notes", "").strip()

        errors = []
        if beneficiary_type == 'person' and len(full_name.split()) < 2:
            errors.append("الاسم يجب أن يكون رباعياً (كلمتان على الأقل)")
        elif beneficiary_type == 'camp' and not full_name.strip():
            errors.append("يرجى إدخال اسم المخيم")
        if beneficiary_type == 'person' and not id_number:
            errors.append("رقم الهوية إلزامي للمستفيد")
        if id_number and (not id_number.isdigit() or len(id_number) != 9):
            errors.append("رقم الهوية يجب أن يكون 9 أرقام فقط")
        if phone and (not phone.isdigit() or len(phone) != 10 or not phone.startswith("05")):
            errors.append("رقم الجوال يجب أن يكون 10 أرقام ويبدأ بـ 05")
        # فحص التكرار: لا يوجد مستفيد آخر بنفس الهوية
        if id_number and not errors:
            c.execute("SELECT id FROM beneficiaries WHERE TRIM(id_number)=TRIM(?) AND id!=?", (id_number, id))
            dup = c.fetchone()
            if dup:
                errors.append(f"رقم الهوية {id_number} مسجل لمستفيد آخر في النظام")

        if errors:
            for e in errors:
                flash(e, "danger")
            conn.close()
            # أعِد عرض البيانات المُدخلة (لا تمسحها بإعادة تحميل DB)
            b = {
                "id": id, "full_name": full_name, "gender": gender,
                "phone": phone, "address": address, "address2": address2,
                "camp_name": camp_name, "id_number": id_number,
                "family_size": family_size, "marital_status": marital_status,
                "children_count": children_count, "wife_pregnant": wife_pregnant,
                "wife_nursing": wife_nursing, "has_orphans": has_orphans,
                "orphans_count": orphans_count, "notes": notes,
                "beneficiary_type": beneficiary_type,
            }
            return render_template("edit_beneficiary.html", b=b)

        try:
            family_size    = int(family_size)
            children_count = int(children_count)
            orphans_count  = int(orphans_count) if has_orphans else 0
        except ValueError:
            family_size = 1; children_count = 0; orphans_count = 0

        c.execute(
            """UPDATE beneficiaries
               SET full_name=?,gender=?,phone=?,address=?,address2=?,camp_name=?,
                   id_number=?,family_size=?,marital_status=?,children_count=?,
                   wife_pregnant=?,wife_nursing=?,has_orphans=?,orphans_count=?,notes=?,
                   beneficiary_type=?,camp_manager_name=?,camp_coordinator=?,
                   camp_coordinator_phone=?,camp_address=?,camp_family_count=?
               WHERE id=? AND (
                   org_id=? OR
                   EXISTS (SELECT 1 FROM org_beneficiary_links obl WHERE obl.beneficiary_id=beneficiaries.id AND obl.org_id=?)
               )""",
            (full_name, gender, phone, address, address2, camp_name,
             id_number, family_size, marital_status, children_count,
             wife_pregnant, wife_nursing, has_orphans, orphans_count, notes,
             beneficiary_type, camp_manager_name, camp_coordinator,
             camp_coordinator_phone, camp_address, camp_family_count, id, org_id, org_id)
        )
        conn.commit()
        conn.close()
        flash("✅ تم تعديل بيانات المستفيد", "success")
        return redirect(url_for("beneficiaries"))

    c.execute("""
        SELECT b.* FROM beneficiaries b
        WHERE b.id=? AND (
            b.org_id=? OR
            EXISTS (SELECT 1 FROM org_beneficiary_links obl WHERE obl.beneficiary_id=b.id AND obl.org_id=?)
        )
    """, (id, org_id, org_id))
    row = c.fetchone()
    conn.close()
    if not row:
        flash("المستفيد غير موجود أو غير مرتبط بمؤسستك", "danger")
        return redirect(url_for("beneficiaries"))
    b = row_to_dict(row, BENEF_KEYS)
    b.setdefault("address2", "")
    b.setdefault("camp_name", "")
    b.setdefault("marital_status", "married")
    b.setdefault("gender", "male")
    b.setdefault("children_count", 0)
    b.setdefault("wife_pregnant", 0)
    b.setdefault("wife_nursing", 0)
    b.setdefault("has_orphans", 0)
    b.setdefault("orphans_count", 0)
    b.setdefault("beneficiary_type", "person")
    return render_template("edit_beneficiary.html", b=b)


@app.route("/delete_beneficiary/<int:id>")
@data_entry_required
def delete_beneficiary(id):
    org_id = session["org_id"]
    conn = get_connection()
    c = conn.cursor()

    # تحقق أن المستفيد مرتبط بهذه المؤسسة
    c.execute("SELECT b.full_name, b.org_id FROM beneficiaries b WHERE b.id=?", (id,))
    row = c.fetchone()
    if not row:
        conn.close()
        flash("المستفيد غير موجود", "danger")
        return redirect(url_for("beneficiaries"))

    full_name = row["full_name"]

    # احذف الربط مع هذه المؤسسة
    try:
        c.execute("DELETE FROM org_beneficiary_links WHERE org_id=? AND beneficiary_id=?", (org_id, id))
        c.execute("SELECT COUNT(*) FROM org_beneficiary_links WHERE beneficiary_id=?", (id,))
        links_remaining = c.fetchone()[0]
    except Exception:
        # الجدول لم يُنشأ بعد — احذف مباشرة
        links_remaining = 0

    if links_remaining == 0 and row["org_id"] == org_id:
        c.execute("DELETE FROM beneficiary_family_members WHERE beneficiary_id=?", (id,))
        c.execute("DELETE FROM beneficiaries WHERE id=?", (id,))

    conn.commit()
    conn.close()
    resequence(org_id)
    flash(f"تم إزالة المستفيد '{full_name}' من مؤسستك", "warning")
    return redirect(url_for("beneficiaries"))


# ══════════════════════════════════════════
# المشتريات (فواتير الإدخال)
# ══════════════════════════════════════════

@app.route("/add_incoming_invoice", methods=["GET", "POST"])
@invoices_required
def add_incoming_invoice():
    """صفحة إضافة فاتورة مشتريات جديدة"""
    if session.get("role") == "observer":
        flash("مراقب عام: لا تملك صلاحية الإضافة", "danger")
        return redirect(url_for("incoming_invoice"))

    org_id = session["org_id"]
    conn = get_connection()
    c = conn.cursor()

    if request.method == "POST":
        supplier                 = request.form.get("supplier", "").strip()
        invoice_date             = request.form.get("invoice_date", "").strip()
        purchase_invoice_number  = request.form.get("purchase_invoice_number", "").strip()
        invoice_name             = request.form.get("invoice_name", "").strip()
        product_names   = request.form.getlist("product_name[]")
        product_units   = request.form.getlist("product_unit[]")
        quantities      = request.form.getlist("quantity[]")
        unit_prices     = request.form.getlist("unit_price[]")
        purchase_dates  = request.form.getlist("purchase_date[]")

        valid_items = [
            (nm.strip(), un.strip(), qty, price, pdate)
            for nm, un, qty, price, pdate in zip(product_names, product_units, quantities, unit_prices, purchase_dates)
            if nm.strip() and qty and price and pdate
        ]

        if not supplier:
            flash("يرجى تعبئة اسم المورد", "danger")
        elif not valid_items:
            flash("يرجى إضافة صنف واحد على الأقل ببيانات كاملة", "danger")
        else:
            c.execute("SELECT COALESCE(MAX(seq_num),0) FROM incoming_invoices WHERE org_id=?", (org_id,))
            next_seq = c.fetchone()[0] + 1

            invoice_image = None
            img_file = request.files.get("invoice_image")
            if img_file and img_file.filename and allowed_file(img_file.filename) and allowed_file_content(img_file):
                ext = img_file.filename.rsplit(".", 1)[1].lower()
                stamp = datetime.now().strftime("%Y%m%d%H%M%S%f")
                invoice_image = f"invoice_{org_id}_{stamp}.{ext}"
                img_file.save(os.path.join(app.config["UPLOAD_FOLDER"], invoice_image))

            c.execute(
                "INSERT INTO incoming_invoices (org_id, seq_num, invoice_number, invoice_date, supplier, created_by, invoice_image, purchase_invoice_number, invoice_name) VALUES (?,?,?,?,?,?,?,?,?)",
                (org_id, next_seq, str(next_seq), invoice_date or datetime.now().strftime("%Y-%m-%d"),
                 supplier, session["user"], invoice_image,
                 purchase_invoice_number or None, invoice_name or None)
            )
            invoice_id = c.lastrowid

            for pname, punit, qty, price, pdate in valid_items:
                try:
                    qty_f, price_f = float(qty), float(price)
                    c.execute("SELECT id FROM products WHERE org_id=? AND LOWER(name)=LOWER(?)", (org_id, pname))
                    row = c.fetchone()
                    if row:
                        pid = row["id"]
                        if punit:
                            c.execute("UPDATE products SET unit=?, last_modified=? WHERE id=?", (punit, pdate, pid))
                    else:
                        unit = punit or "وحدة"
                        c.execute("INSERT INTO products (org_id, name, unit, last_modified) VALUES (?,?,?,?)",
                                  (org_id, pname, unit, pdate))
                        pid = c.lastrowid
                    c.execute(
                        "INSERT INTO incoming_invoice_items (invoice_id, product_id, quantity, unit_price, total_price, purchase_date) VALUES (?,?,?,?,?,?)",
                        (invoice_id, pid, qty_f, price_f, qty_f * price_f, pdate)
                    )
                    c.execute(
                        "INSERT INTO stock_batches (org_id, product_id, quantity_remaining, unit_price, entry_date, invoice_id) VALUES (?,?,?,?,?,?)",
                        (org_id, pid, qty_f, price_f, pdate, invoice_id)
                    )
                    c.execute("UPDATE products SET last_modified=? WHERE id=?", (pdate, pid))
                except ValueError:
                    pass

            conn.commit()
            conn.close()
            flash("✅ تم حفظ فاتورة المشتريات", "success")
            return redirect(url_for("incoming_invoices_list"))

    c.execute("SELECT name, unit FROM products WHERE org_id=? ORDER BY name", (org_id,))
    existing_products = [dict(r) for r in c.fetchall()]
    today_str = datetime.now().strftime("%Y-%m-%d")
    conn.close()
    return render_template("add_incoming_invoice.html",
                           existing_products=existing_products, today_str=today_str)


@app.route("/incoming_invoice", methods=["GET"])
@invoices_view_required
def incoming_invoice():
    """كشف مختصر بكل الفواتير — hover ينقل لصفحة التفاصيل"""
    org_id = session["org_id"]
    conn = get_connection()
    c = conn.cursor()

    c.execute("""
        SELECT id, seq_num, invoice_number, invoice_date, supplier,
               is_closed, is_paid, grand_total, created_at
        FROM incoming_invoices WHERE org_id=? ORDER BY COALESCE(seq_num,id) ASC
    """, (org_id,))
    all_inv_rows = c.fetchall()

    summary = []
    total_purchases = 0.0
    for inv in all_inv_rows:
        inv_d = dict(inv)
        c.execute("SELECT COALESCE(SUM(total_price),0) FROM incoming_invoice_items WHERE invoice_id=?", (inv["id"],))
        inv_total = c.fetchone()[0]
        inv_d["total"] = inv_total
        inv_d["display_num"] = inv_d.get("seq_num") or inv_d.get("invoice_number") or inv_d["id"]
        total_purchases += inv_total
        summary.append(inv_d)

    conn.close()
    return render_template("incoming_invoices.html",
        summary=summary,
        total_purchases=total_purchases)


@app.route("/incoming_invoice/<int:id>/add_items", methods=["POST"])
@invoices_required
def add_invoice_items(id):
    org_id = session["org_id"]
    conn = get_connection()
    c = conn.cursor()

    c.execute("SELECT is_closed FROM incoming_invoices WHERE id=? AND org_id=?", (id, org_id))
    inv = c.fetchone()
    if not inv:
        conn.close()
        flash("الفاتورة غير موجودة", "danger")
        return redirect(url_for("incoming_invoice"))
    if inv["is_closed"]:
        conn.close()
        flash("الفاتورة مغلقة، لا يمكن إضافة أصناف جديدة", "danger")
        return redirect(url_for("incoming_invoice"))

    product_names  = request.form.getlist("product_name[]")
    product_units  = request.form.getlist("product_unit[]")
    quantities     = request.form.getlist("quantity[]")
    unit_prices    = request.form.getlist("unit_price[]")
    purchase_dates = request.form.getlist("purchase_date[]")

    added = 0
    for pname, punit, qty, price, pdate in zip(product_names, product_units, quantities, unit_prices, purchase_dates):
        pname = pname.strip()
        if pname and qty and price and pdate:
            try:
                qty_f, price_f = float(qty), float(price)
                c.execute("SELECT id FROM products WHERE org_id=? AND LOWER(name)=LOWER(?)", (org_id, pname))
                row = c.fetchone()
                if row:
                    pid = row["id"]
                    if punit.strip():
                        c.execute("UPDATE products SET unit=?, last_modified=? WHERE id=?", (punit.strip(), pdate, pid))
                else:
                    unit = punit.strip() or "وحدة"
                    c.execute("INSERT INTO products (org_id, name, unit, last_modified) VALUES (?,?,?,?)",
                              (org_id, pname, unit, pdate))
                    pid = c.lastrowid
                c.execute(
                    "INSERT INTO incoming_invoice_items (invoice_id, product_id, quantity, unit_price, total_price, purchase_date) VALUES (?,?,?,?,?,?)",
                    (id, pid, qty_f, price_f, qty_f * price_f, pdate)
                )
                c.execute(
                    "INSERT INTO stock_batches (org_id, product_id, quantity_remaining, unit_price, entry_date, invoice_id) VALUES (?,?,?,?,?,?)",
                    (org_id, pid, qty_f, price_f, pdate, id)
                )
                c.execute("UPDATE products SET last_modified=? WHERE id=?", (pdate, pid))
                added += 1
            except ValueError:
                pass

    conn.commit()
    conn.close()
    if added:
        flash(f"✅ تم إضافة {added} صنف للفاتورة", "success")
    else:
        flash("يرجى تعبئة بيانات الصنف كاملة (الكمية، السعر، التاريخ)", "danger")
    return redirect(url_for("incoming_invoice"))


@app.route("/incoming_invoice/<int:id>/close")
@invoices_required
def close_invoice(id):
    org_id = session["org_id"]
    conn = get_connection()
    c = conn.cursor()

    c.execute("SELECT is_closed FROM incoming_invoices WHERE id=? AND org_id=?", (id, org_id))
    inv = c.fetchone()
    if not inv:
        conn.close()
        flash("الفاتورة غير موجودة", "danger")
        return redirect(url_for("incoming_invoice"))
    if inv["is_closed"]:
        conn.close()
        flash("الفاتورة مغلقة مسبقاً", "warning")
        return redirect(url_for("incoming_invoice"))

    c.execute("SELECT COALESCE(SUM(total_price),0) FROM incoming_invoice_items WHERE invoice_id=?", (id,))
    total = c.fetchone()[0]
    c.execute("UPDATE incoming_invoices SET is_closed=1, grand_total=? WHERE id=? AND org_id=?", (total, id, org_id))
    conn.commit()
    conn.close()
    flash(f"🔒 تم إغلاق الفاتورة نهائياً — الإجمالي: {total:.2f}", "success")
    return redirect(url_for("incoming_invoice"))


@app.route("/incoming_invoice/<int:id>/upload_image", methods=["POST"])
@invoices_required
def upload_invoice_image(id):
    org_id = session["org_id"]
    conn = get_connection()
    c = conn.cursor()

    c.execute("SELECT id FROM incoming_invoices WHERE id=? AND org_id=?", (id, org_id))
    if not c.fetchone():
        conn.close()
        flash("الفاتورة غير موجودة", "danger")
        return redirect(url_for("incoming_invoice"))

    img_file = request.files.get("invoice_image")
    if img_file and img_file.filename and allowed_file(img_file.filename) and allowed_file_content(img_file):
        ext = img_file.filename.rsplit(".", 1)[1].lower()
        stamp = datetime.now().strftime("%Y%m%d%H%M%S%f")
        filename = f"invoice_{org_id}_{stamp}.{ext}"
        img_file.save(os.path.join(app.config["UPLOAD_FOLDER"], filename))
        c.execute("UPDATE incoming_invoices SET invoice_image=? WHERE id=? AND org_id=?", (filename, id, org_id))
        conn.commit()
        flash("✅ تم رفع صورة الفاتورة", "success")
    else:
        flash("يرجى اختيار صورة بصيغة png / jpg / jpeg / webp", "danger")

    conn.close()
    return redirect(url_for("incoming_invoice"))


@app.route("/incoming_invoice/<int:id>/upload/<img_type>", methods=["POST"])
@invoices_required
def upload_invoice_image_typed(id, img_type):
    """رفع صور الفاتورة الثلاث: invoice / receipt / attachment"""
    if img_type not in ("invoice", "receipt", "attachment"):
        flash("نوع الصورة غير صالح", "danger")
        return redirect(url_for("incoming_invoice"))

    org_id = session["org_id"]
    conn = get_connection()
    c = conn.cursor()
    c.execute("SELECT id, is_closed FROM incoming_invoices WHERE id=? AND org_id=?", (id, org_id))
    inv = c.fetchone()
    if not inv:
        conn.close()
        flash("الفاتورة غير موجودة", "danger")
        return redirect(url_for("incoming_invoice"))

    field = {"invoice": "invoice_image", "receipt": "receipt_image", "attachment": "attachment_image"}[img_type]
    img_file = request.files.get("img_file")
    if img_file and img_file.filename and allowed_file(img_file.filename) and allowed_file_content(img_file):
        ext = img_file.filename.rsplit(".", 1)[1].lower()
        stamp = datetime.now().strftime("%Y%m%d%H%M%S%f")
        filename = f"{img_type}_{org_id}_{stamp}.{ext}"
        img_file.save(os.path.join(app.config["UPLOAD_FOLDER"], filename))
        # حالة الدفع تُحدَّد تلقائياً من وجود صورة سند القبض
        if img_type == "receipt":
            c.execute(f"UPDATE incoming_invoices SET {field}=?, is_paid=1 WHERE id=? AND org_id=?", (filename, id, org_id))
        else:
            c.execute(f"UPDATE incoming_invoices SET {field}=? WHERE id=? AND org_id=?", (filename, id, org_id))
        conn.commit()
        flash("✅ تم رفع الصورة", "success")
    else:
        flash("يرجى اختيار صورة بصيغة png / jpg / jpeg / webp", "danger")

    conn.close()
    return redirect(url_for("incoming_invoice") + f"#inv-{id}")


@app.route("/incoming_invoice/<int:id>/delete")
@invoices_required
def delete_incoming_invoice(id):
    """حذف فاتورة وإعادة ترقيم الفواتير التالية"""
    if session.get("role") not in ("admin", "accountant"):
        flash("ليس لديك صلاحية حذف الفواتير", "danger")
        return redirect(url_for("incoming_invoice"))

    org_id = session["org_id"]
    conn = get_connection()
    c = conn.cursor()

    c.execute("SELECT id, seq_num FROM incoming_invoices WHERE id=? AND org_id=?", (id, org_id))
    inv = c.fetchone()
    if not inv:
        conn.close()
        flash("الفاتورة غير موجودة", "danger")
        return redirect(url_for("incoming_invoice"))

    deleted_seq = inv["seq_num"]

    # حذف دفعات المخزون المرتبطة بهذه الفاتورة
    c.execute("DELETE FROM stock_batches WHERE invoice_id=? AND org_id=?", (id, org_id))
    # حذف أصناف الفاتورة
    c.execute("DELETE FROM incoming_invoice_items WHERE invoice_id=?", (id,))
    # حذف الفاتورة
    c.execute("DELETE FROM incoming_invoices WHERE id=? AND org_id=?", (id, org_id))

    # إعادة ترقيم الفواتير التالية
    if deleted_seq:
        c.execute("""
            UPDATE incoming_invoices
            SET seq_num = seq_num - 1,
                invoice_number = CAST((seq_num - 1) AS TEXT)
            WHERE org_id=? AND seq_num > ?
        """, (org_id, deleted_seq))

    conn.commit()
    conn.close()
    flash("🗑️ تم حذف الفاتورة وإعادة الترقيم", "warning")
    return redirect(url_for("incoming_invoice"))


@app.route("/incoming_invoice/<int:id>/update", methods=["POST"])
@invoices_required
def update_invoice(id):
    """تعديل فاتورة مفتوحة (أدمن فقط): المورد، التاريخ، الأصناف كاملة"""
    if session.get("role") != "admin":
        flash("هذه الميزة للأدمن فقط", "danger")
        return redirect(url_for("incoming_invoices_list"))

    org_id = session["org_id"]
    conn = get_connection()
    c = conn.cursor()

    c.execute("SELECT is_closed FROM incoming_invoices WHERE id=? AND org_id=?", (id, org_id))
    inv = c.fetchone()
    if not inv or inv["is_closed"]:
        conn.close()
        flash("لا يمكن تعديل فاتورة مغلقة", "danger")
        return redirect(url_for("incoming_invoices_list"))

    # تعديل رأس الفاتورة
    supplier     = request.form.get("supplier", "").strip()
    invoice_date = request.form.get("invoice_date", "").strip()
    if supplier:
        c.execute("UPDATE incoming_invoices SET supplier=?, invoice_date=? WHERE id=? AND org_id=?",
                  (supplier, invoice_date, id, org_id))

    # تعديل الأصناف (اسم + سعر + كمية + تاريخ)
    item_ids    = request.form.getlist("item_id[]")
    prod_names  = request.form.getlist("prod_name[]")
    unit_prices = request.form.getlist("unit_price[]")
    quantities  = request.form.getlist("quantity[]")
    prod_dates  = request.form.getlist("prod_date[]")

    for item_id, pname, uprice, qty_str, pdate in zip(item_ids, prod_names, unit_prices, quantities, prod_dates):
        try:
            uprice_f = float(uprice)
            new_qty  = float(qty_str)
            if new_qty <= 0:
                continue
            c.execute("SELECT product_id, quantity FROM incoming_invoice_items WHERE id=? AND invoice_id=?",
                      (item_id, id))
            row = c.fetchone()
            if not row:
                continue
            pid, old_qty = row["product_id"], row["quantity"]
            qty_diff = new_qty - old_qty   # موجب = زيادة، سالب = نقصان

            # تحديث اسم المنتج
            if pname.strip():
                c.execute("UPDATE products SET name=?, last_modified=? WHERE id=? AND org_id=?",
                          (pname.strip(), pdate, pid, org_id))

            # تحديث الصنف في الفاتورة
            c.execute("""
                UPDATE incoming_invoice_items
                SET quantity=?, unit_price=?, total_price=?, purchase_date=?
                WHERE id=?
            """, (new_qty, uprice_f, new_qty * uprice_f, pdate, item_id))

            # تحديث stock_batch: الكمية المتبقية + السعر + التاريخ
            c.execute("""
                UPDATE stock_batches
                SET quantity_remaining = quantity_remaining + ?,
                    unit_price = ?,
                    entry_date = ?
                WHERE product_id=? AND invoice_id=? AND org_id=?
            """, (qty_diff, uprice_f, pdate, pid, id, org_id))

        except (ValueError, TypeError):
            pass

    conn.commit()
    conn.close()
    flash("✅ تم حفظ التعديلات", "success")
    return redirect(url_for("incoming_invoices_list") + f"#inv-{id}")


@app.route("/incoming_invoice/item/<int:item_id>/delete")
@invoices_required
def delete_incoming_item(item_id):
    """حذف صنف من فاتورة مشتريات مفتوحة مع إزالته من المخزن (أدمن فقط)"""
    if session.get("role") != "admin":
        flash("هذه الميزة للأدمن فقط", "danger")
        return redirect(url_for("incoming_invoices_list"))
    org_id = session["org_id"]
    conn = get_connection()
    c = conn.cursor()

    c.execute("""
        SELECT ii.id, ii.invoice_id, ii.product_id, ii.quantity
        FROM incoming_invoice_items ii
        JOIN incoming_invoices inv ON inv.id = ii.invoice_id
        WHERE ii.id=? AND inv.org_id=? AND COALESCE(inv.is_closed,0)=0
    """, (item_id, org_id))
    item = c.fetchone()
    if not item:
        conn.close()
        flash("الصنف غير موجود أو الفاتورة مغلقة", "danger")
        return redirect(url_for("incoming_invoices_list"))

    invoice_id = item["invoice_id"]
    product_id = item["product_id"]
    qty        = item["quantity"]

    # إزالة الكمية من المخزن (تخفيض quantity_remaining في الـ batch المرتبط بهذه الفاتورة)
    c.execute("""
        UPDATE stock_batches
        SET quantity_remaining = MAX(0, quantity_remaining - ?)
        WHERE product_id=? AND invoice_id=? AND org_id=?
    """, (qty, product_id, invoice_id, org_id))

    c.execute("DELETE FROM incoming_invoice_items WHERE id=?", (item_id,))
    conn.commit()
    conn.close()
    flash("✅ تم حذف الصنف من الفاتورة والمخزن", "success")
    return redirect(url_for("incoming_invoices_list") + f"#inv-{invoice_id}")


@app.route("/incoming_invoices_list")
@invoices_view_required
def incoming_invoices_list():
    """صفحة الفواتير المضافة — كل الفواتير (مفتوحة ومغلقة) مع فلتر"""
    org_id = session["org_id"]
    conn = get_connection()
    c = conn.cursor()

    c.execute("""
        SELECT id, seq_num, invoice_number, invoice_date, supplier,
               is_closed, is_paid, invoice_image, receipt_image, attachment_image, created_at,
               purchase_invoice_number, invoice_name, notes
        FROM incoming_invoices
        WHERE org_id=?
        ORDER BY COALESCE(seq_num,id) ASC
    """, (org_id,))
    inv_rows = c.fetchall()

    invoices = []
    total_all = 0.0
    for inv in inv_rows:
        inv_d = dict(inv)
        c.execute("""
            SELECT ii.id, ii.quantity, ii.unit_price, ii.total_price, ii.purchase_date,
                   p.name AS product_name, p.unit AS product_unit, p.id AS product_id
            FROM incoming_invoice_items ii
            JOIN products p ON p.id = ii.product_id
            WHERE ii.invoice_id=? ORDER BY ii.id
        """, (inv["id"],))
        items = [dict(r) for r in c.fetchall()]
        inv_total = sum(it["total_price"] for it in items)
        inv_d["lines"] = items
        inv_d["total"] = inv_total
        inv_d["display_num"] = inv_d.get("seq_num") or inv_d.get("invoice_number") or inv_d["id"]
        total_all += inv_total
        invoices.append(inv_d)

    today_str = datetime.now().strftime("%Y-%m-%d")
    # بنود المصروفات التشغيلية
    c.execute("""
        SELECT * FROM expense_items WHERE org_id=? ORDER BY expense_date DESC, id DESC
    """, (org_id,))
    expense_items = [dict(r) for r in c.fetchall()]
    expense_total = sum(e['amount'] or 0 for e in expense_items)
    conn.close()
    return render_template("incoming_invoices_list.html",
                           invoices=invoices, total_all=total_all, today_str=today_str,
                           expense_items=expense_items, expense_total=expense_total)


@app.route("/incoming_invoice/<int:id>/save_notes", methods=["POST"])
@invoices_required
def save_incoming_invoice_notes(id):
    """حفظ ملاحظات فاتورة الشراء — يعمل للمفتوحة والمغلقة"""
    if session.get("role") == "observer":
        flash("ليس لديك صلاحية", "danger")
        return redirect(url_for("incoming_invoices_list") + f"#inv-{id}")
    if not validate_csrf():
        flash("خطأ في التحقق", "danger")
        return redirect(url_for("incoming_invoices_list") + f"#inv-{id}")
    org_id = session["org_id"]
    notes = request.form.get("notes", "").strip()
    conn = get_connection()
    c = conn.cursor()
    c.execute("SELECT id FROM incoming_invoices WHERE id=? AND org_id=?", (id, org_id))
    if not c.fetchone():
        conn.close()
        flash("الفاتورة غير موجودة", "danger")
        return redirect(url_for("incoming_invoices_list"))
    c.execute("UPDATE incoming_invoices SET notes=? WHERE id=? AND org_id=?", (notes, id, org_id))
    conn.commit()
    conn.close()
    flash("✅ تم حفظ الملاحظة", "success")
    return redirect(url_for("incoming_invoices_list") + f"#inv-{id}")


# ══════════════════════════════════════════
# البرامج والمشاريع (أدمن فقط)
# ══════════════════════════════════════════
DEFAULT_PROGRAMS = [
    "كفالات أيتام",
    "طرود غذائية",
    "تكية",
    "مياه حلوة",
    "وجبات ساخنة",
]

@app.route("/programs")
@programs_view_required
def programs_list():
    org_id = session["org_id"]
    conn = get_connection()
    c = conn.cursor()
    # إضافة البرامج الافتراضية إن لم تكن موجودة
    c.execute("SELECT COUNT(*) FROM programs WHERE org_id=?", (org_id,))
    if c.fetchone()[0] == 0:
        for p in DEFAULT_PROGRAMS:
            c.execute("INSERT INTO programs (org_id, name) VALUES (?,?)", (org_id, p))
        conn.commit()
    c.execute("SELECT * FROM programs WHERE org_id=? ORDER BY id", (org_id,))
    programs = [dict(r) for r in c.fetchall()]

    # إحصاءات الفئات المستحقة من بيانات المستفيدين
    c.execute("""
        SELECT COALESCE(SUM(orphans_count),0) as total_orphans,
               COUNT(*) as families_with_orphans
        FROM beneficiaries WHERE org_id=? AND has_orphans=1
    """, (org_id,))
    orphan_stats = dict(c.fetchone())

    c.execute("""
        SELECT COUNT(*) FROM beneficiaries
        WHERE org_id=? AND (marital_status='widowed' OR family_size=1)
    """, (org_id,))
    widows_count = c.fetchone()[0]

    c.execute("SELECT COUNT(*) FROM beneficiaries WHERE org_id=? AND wife_pregnant=1", (org_id,))
    pregnant_count = c.fetchone()[0]

    c.execute("SELECT COUNT(*) FROM beneficiaries WHERE org_id=? AND wife_nursing=1", (org_id,))
    nursing_count = c.fetchone()[0]

    conn.close()
    return render_template("programs.html", programs=programs,
        orphan_stats=orphan_stats, widows_count=widows_count,
        pregnant_count=pregnant_count, nursing_count=nursing_count)


@app.route("/add_program", methods=["POST"])
@admin_required
def add_program():
    name = request.form.get("name", "").strip()
    if name:
        conn = get_connection()
        c = conn.cursor()
        c.execute("INSERT INTO programs (org_id, name) VALUES (?,?)",
                  (session["org_id"], name))
        conn.commit()
        conn.close()
        flash(f"✅ تمت إضافة البرنامج '{name}'", "success")
    return redirect(url_for("programs_list"))


@app.route("/toggle_program/<int:id>")
@admin_required
def toggle_program(id):
    conn = get_connection()
    c = conn.cursor()
    c.execute("UPDATE programs SET is_active = 1 - is_active WHERE id=? AND org_id=?",
              (id, session["org_id"]))
    conn.commit()
    conn.close()
    return redirect(url_for("programs_list"))


@app.route("/delete_program/<int:id>")
@admin_required
def delete_program(id):
    conn = get_connection()
    c = conn.cursor()
    c.execute("DELETE FROM programs WHERE id=? AND org_id=?", (id, session["org_id"]))
    conn.commit()
    conn.close()
    flash("تم حذف البرنامج", "warning")
    return redirect(url_for("programs_list"))


# ══════════════════════════════════════════
# تفاصيل البرنامج وسجلات الاستفادة
# ══════════════════════════════════════════
@app.route("/program/<int:id>")
@programs_view_required
def program_detail(id):
    org_id = session["org_id"]
    conn = get_connection()
    c = conn.cursor()

    c.execute("SELECT * FROM programs WHERE id=? AND org_id=?", (id, org_id))
    program = c.fetchone()
    if not program:
        conn.close()
        flash("البرنامج غير موجود", "danger")
        return redirect(url_for("programs_list"))
    program = dict(program)

    c.execute("""
        SELECT pr.id, pr.benefit_date, pr.benefit_type, pr.quantity, pr.notes, pr.created_by,
               pr.ben_confirmed, pr.ben_confirmed_at,
               b.seq_num, b.full_name, b.id_number
        FROM program_records pr
        JOIN beneficiaries b ON b.id = pr.beneficiary_id
        WHERE pr.org_id=? AND pr.program_id=?
        ORDER BY pr.benefit_date DESC, pr.id DESC
    """, (org_id, id))
    records = [dict(r) for r in c.fetchall()]

    c.execute("""
        SELECT DISTINCT b.id, b.seq_num, b.full_name
        FROM beneficiaries b
        WHERE b.org_id=?
           OR EXISTS (SELECT 1 FROM org_beneficiary_links obl WHERE obl.beneficiary_id=b.id AND obl.org_id=?)
        ORDER BY b.full_name
    """, (org_id, org_id))
    beneficiaries = [dict(r) for r in c.fetchall()]

    conn.close()
    return render_template("program_detail.html",
        program=program, records=records, beneficiaries=beneficiaries)


@app.route("/program/<int:id>/add_record", methods=["POST"])
@admin_required
def add_program_record(id):
    org_id = session["org_id"]

    beneficiary_id = request.form.get("beneficiary_id", "").strip()
    benefit_date   = request.form.get("benefit_date", "").strip()
    benefit_type   = request.form.get("benefit_type", "").strip()
    quantity       = request.form.get("quantity", "").strip()
    notes          = request.form.get("notes", "").strip()

    if not beneficiary_id or not benefit_date:
        flash("يرجى اختيار المستفيد وتاريخ الاستفادة", "danger")
        return redirect(url_for("program_detail", id=id))

    conn = get_connection()
    c = conn.cursor()

    # التأكد أن البرنامج والمستفيد ينتميان لنفس المؤسسة
    c.execute("SELECT id FROM programs WHERE id=? AND org_id=?", (id, org_id))
    if not c.fetchone():
        conn.close()
        flash("البرنامج غير موجود", "danger")
        return redirect(url_for("programs_list"))

    c.execute("""
        SELECT id FROM beneficiaries WHERE id=? AND (
            org_id=? OR
            EXISTS (SELECT 1 FROM org_beneficiary_links obl WHERE obl.beneficiary_id=beneficiaries.id AND obl.org_id=?)
        )
    """, (beneficiary_id, org_id, org_id))
    if not c.fetchone():
        conn.close()
        flash("المستفيد غير موجود أو غير مرتبط بمؤسستك", "danger")
        return redirect(url_for("program_detail", id=id))

    c.execute("""
        INSERT INTO program_records
            (org_id, program_id, beneficiary_id, benefit_date, benefit_type, quantity, notes, created_by, ben_confirmed)
        VALUES (?,?,?,?,?,?,?,?,0)
    """, (org_id, id, beneficiary_id, benefit_date, benefit_type, quantity, notes, session["user"]))
    conn.commit()
    conn.close()
    flash("✅ تم إضافة سجل الاستفادة", "success")
    return redirect(url_for("program_detail", id=id))


@app.route("/program/<int:id>/delete_record/<int:record_id>")
@admin_required
def delete_program_record(id, record_id):
    org_id = session["org_id"]
    conn = get_connection()
    c = conn.cursor()
    c.execute("DELETE FROM program_records WHERE id=? AND org_id=? AND program_id=?",
              (record_id, org_id, id))
    conn.commit()
    conn.close()
    flash("تم حذف السجل", "warning")
    return redirect(url_for("program_detail", id=id))


# ══════════════════════════════════════════
# فواتير الصرف
# ══════════════════════════════════════════
@app.route("/outgoing_invoice", methods=["GET", "POST"])
@invoices_view_required
def outgoing_invoice():
    org_id = session["org_id"]
    conn = get_connection()
    c = conn.cursor()

    if request.method == "POST" and session.get("role") == "observer":
        flash("مراقب عام: لا تملك صلاحية الإضافة أو التعديل", "danger")
        return redirect(url_for("outgoing_invoice"))

    def get_inventory():
        c.execute("SELECT id, name, unit FROM products WHERE org_id=? ORDER BY name", (org_id,))
        prods = c.fetchall()
        inv = []
        total_val = 0.0
        for p in prods:
            c.execute("SELECT quantity_remaining, unit_price FROM stock_batches WHERE product_id=? AND org_id=? ORDER BY id", (p["id"], org_id))
            batches = c.fetchall()
            total_qty = sum(b["quantity_remaining"] for b in batches)
            if total_qty > 0:
                avg_price = sum(b["quantity_remaining"] * b["unit_price"] for b in batches) / total_qty
            elif batches:
                avg_price = batches[-1]["unit_price"]
            else:
                avg_price = 0.0
            value = total_qty * avg_price
            total_val += value
            inv.append({"id": p["id"], "name": p["name"], "unit": p["unit"],
                         "total_qty": total_qty, "avg_price": avg_price, "value": value})
        return inv, total_val

    def get_render_data():
        c.execute("SELECT id, name FROM products WHERE org_id=? ORDER BY name", (org_id,))
        pl = c.fetchall()
        c.execute("""
            SELECT DISTINCT b.id, b.full_name FROM beneficiaries b
            WHERE b.org_id=?
               OR EXISTS (SELECT 1 FROM org_beneficiary_links obl WHERE obl.beneficiary_id=b.id AND obl.org_id=?)
            ORDER BY b.full_name
        """, (org_id, org_id))
        bl = c.fetchall()
        c.execute("SELECT * FROM programs WHERE org_id=? AND is_active=1 ORDER BY name", (org_id,))
        rows = c.fetchall()
        progs = [dict(r) for r in rows]
        if not progs:
            for p in DEFAULT_PROGRAMS:
                c.execute("INSERT INTO programs (org_id, name) VALUES (?,?)", (org_id, p))
            conn.commit()
            c.execute("SELECT * FROM programs WHERE org_id=? AND is_active=1 ORDER BY name", (org_id,))
            progs = [dict(r) for r in c.fetchall()]
        return pl, bl, progs

    if request.method == "POST":
        invoice_number = request.form.get("invoice_number", "").strip()
        invoice_date   = request.form.get("invoice_date", "").strip()
        program_id     = request.form.get("program_id", "").strip()
        beneficiary    = request.form.get("beneficiary", "").strip()
        notes          = request.form.get("notes", "").strip()
        product_ids    = request.form.getlist("product_id[]")
        product_names  = request.form.getlist("product_name[]")
        quantities     = request.form.getlist("quantity[]")

        # البرنامج هو الجهة المستفيدة الرئيسية
        disbursed_to = ""
        if program_id:
            c.execute("SELECT name FROM programs WHERE id=? AND org_id=?", (program_id, org_id))
            prog = c.fetchone()
            if prog:
                disbursed_to = prog["name"]
        # لو في مستفيد محدد نضيفه للاسم
        if beneficiary:
            disbursed_to = f"{disbursed_to} — {beneficiary}" if disbursed_to else beneficiary

        if not invoice_number or not invoice_date or not disbursed_to:
            flash("يرجى تعبئة رقم الفاتورة والتاريخ والبرنامج", "danger")
            pl, bl, progs = get_render_data()
            inventory, total_value = get_inventory()
            conn.close()
            return render_template("outgoing_invoices.html",
                products=pl, beneficiaries=bl, programs=progs,
                inventory=inventory, total_value=total_value)

        c.execute(
            "INSERT INTO outgoing_invoices (org_id, invoice_number, invoice_date, beneficiary, notes, created_by) VALUES (?,?,?,?,?,?)",
            (org_id, invoice_number, invoice_date, disbursed_to, notes, session["user"])
        )
        invoice_id = c.lastrowid
        errors = []
        # ضمان نفس طول القوائم الثلاث بتعبئة الناقص بـ ''
        max_len = max(len(product_ids), len(product_names), len(quantities))
        product_ids   += [''] * (max_len - len(product_ids))
        product_names += [''] * (max_len - len(product_names))
        quantities    += [''] * (max_len - len(quantities))

        for pid, pname, qty in zip(product_ids, product_names, quantities):
            # إذا product_id فارغ ولكن في اسم، ابحث عنه بالاسم
            if not pid and pname:
                c.execute("SELECT id FROM products WHERE name=? AND org_id=?", (pname.strip(), org_id))
                row = c.fetchone()
                if row:
                    pid = str(row["id"])
            if pid and qty:
                try:
                    qty_needed = float(qty)
                    c.execute("SELECT id, quantity_remaining, unit_price FROM stock_batches WHERE product_id=? AND org_id=? AND quantity_remaining>0 ORDER BY id ASC",
                              (pid, org_id))
                    batches = c.fetchall()
                    total_available = sum(b["quantity_remaining"] for b in batches)
                    if total_available < qty_needed:
                        c.execute("SELECT name FROM products WHERE id=?", (pid,))
                        pname = c.fetchone()["name"]
                        errors.append(f"الكمية المطلوبة من '{pname}' ({qty_needed}) أكبر من المتاح ({total_available:.2f})")
                        continue
                    remaining, weighted = qty_needed, 0
                    for batch in batches:
                        if remaining <= 0:
                            break
                        take = min(remaining, batch["quantity_remaining"])
                        weighted += take * batch["unit_price"]
                        c.execute("UPDATE stock_batches SET quantity_remaining=? WHERE id=?",
                                  (batch["quantity_remaining"] - take, batch["id"]))
                        remaining -= take
                    c.execute("INSERT INTO outgoing_invoice_items (invoice_id, product_id, quantity, unit_price, total_price) VALUES (?,?,?,?,?)",
                              (invoice_id, pid, qty_needed, weighted/qty_needed, weighted))
                except ValueError:
                    pass
        conn.commit()
        conn.close()
        for err in errors:
            flash(err, "warning")
        if not errors:
            flash("✅ تم حفظ فاتورة الصرف وتم خصم الكمية من المخزن تلقائياً", "success")
            return redirect(url_for("outgoing_invoices_list"))
        return redirect(url_for("outgoing_invoice"))

    pl, bl, progs = get_render_data()
    inventory, total_value = get_inventory()
    conn.close()
    return render_template("outgoing_invoices.html",
        products=pl, beneficiaries=bl, programs=progs,
        inventory=inventory, total_value=total_value)


# ══════════════════════════════════════════
# كشف فواتير الصرف
# ══════════════════════════════════════════
@app.route("/outgoing_invoices_list")
@invoices_view_required
def outgoing_invoices_list():
    org_id = session["org_id"]
    conn = get_connection()
    c = conn.cursor()
    c.execute("""
        SELECT id, invoice_number, invoice_date, beneficiary, notes, created_at,
               COALESCE(is_closed,0) AS is_closed
        FROM outgoing_invoices WHERE org_id=? ORDER BY id DESC
    """, (org_id,))
    invoices_rows = c.fetchall()

    invoices = []
    total_all = 0.0
    for inv in invoices_rows:
        inv_d = dict(inv)
        c.execute("""
            SELECT oi.id, oi.quantity, oi.unit_price, oi.total_price,
                   p.name as product_name, p.unit
            FROM outgoing_invoice_items oi
            JOIN products p ON p.id = oi.product_id
            WHERE oi.invoice_id=?
            ORDER BY oi.id
        """, (inv["id"],))
        items = [dict(r) for r in c.fetchall()]
        inv_total = sum(it["total_price"] for it in items)
        inv_d["lines"] = items
        inv_d["total"] = inv_total
        total_all += inv_total
        invoices.append(inv_d)

    today_str = datetime.now().strftime("%Y-%m-%d")
    conn.close()
    return render_template("outgoing_invoices_list.html",
                           invoices=invoices, total_all=total_all, today_str=today_str)


# ══════════════════════════════════════════
# بند المصروفات التشغيلية
# ══════════════════════════════════════════
import os as _os

@app.route("/expenses/add", methods=["POST"])
@login_required
def expense_add():
    if session.get("role") not in ("admin", "accountant", "data_entry"):
        return redirect(url_for("incoming_invoices_list"))
    if not validate_csrf():
        return redirect(url_for("incoming_invoices_list"))
    org_id       = session["org_id"]
    category     = request.form.get("category", "").strip()
    expense_date = request.form.get("expense_date", "").strip()
    amount       = request.form.get("amount", "0").strip()
    notes        = request.form.get("notes", "").strip()
    if not category:
        flash("يرجى إدخال الصنف", "danger")
        return redirect(url_for("incoming_invoices_list"))
    try:
        amount = float(amount)
    except ValueError:
        amount = 0.0
    conn = get_connection()
    c    = conn.cursor()
    c.execute(
        "INSERT INTO expense_items (org_id, category, expense_date, amount, notes) VALUES (?,?,?,?,?)",
        (org_id, category, expense_date or None, amount, notes)
    )
    new_id = c.lastrowid
    # رفع صورة الإثبات إن وجدت
    file = request.files.get("proof_image")
    if file and file.filename:
        ext  = file.filename.rsplit(".", 1)[-1].lower()
        fname = f"exp_{org_id}_{new_id}.{ext}"
        file.save(_os.path.join(app.config["UPLOAD_FOLDER"], fname))
        c.execute("UPDATE expense_items SET proof_image=? WHERE id=?", (fname, new_id))
    conn.commit()
    conn.close()
    flash("✅ تم إضافة البند", "success")
    return redirect(url_for("incoming_invoices_list") + "#expense-card")


@app.route("/expenses/edit/<int:eid>", methods=["POST"])
@login_required
def expense_edit(eid):
    if session.get("role") not in ("admin", "accountant"):
        return redirect(url_for("incoming_invoices_list"))
    if not validate_csrf():
        return redirect(url_for("incoming_invoices_list"))
    org_id       = session["org_id"]
    category     = request.form.get("category", "").strip()
    expense_date = request.form.get("expense_date", "").strip()
    amount       = request.form.get("amount", "0").strip()
    notes        = request.form.get("notes", "").strip()
    try:
        amount = float(amount)
    except ValueError:
        amount = 0.0
    conn = get_connection()
    c    = conn.cursor()
    # رفع صورة جديدة إن وجدت
    file = request.files.get("proof_image")
    if file and file.filename:
        ext   = file.filename.rsplit(".", 1)[-1].lower()
        fname = f"exp_{org_id}_{eid}.{ext}"
        file.save(_os.path.join(app.config["UPLOAD_FOLDER"], fname))
        c.execute(
            "UPDATE expense_items SET category=?,expense_date=?,amount=?,notes=?,proof_image=? WHERE id=? AND org_id=?",
            (category, expense_date or None, amount, notes, fname, eid, org_id)
        )
    else:
        c.execute(
            "UPDATE expense_items SET category=?,expense_date=?,amount=?,notes=? WHERE id=? AND org_id=?",
            (category, expense_date or None, amount, notes, eid, org_id)
        )
    conn.commit()
    conn.close()
    flash("✅ تم تعديل البند", "success")
    return redirect(url_for("incoming_invoices_list") + "#expense-card")


@app.route("/expenses/delete/<int:eid>", methods=["POST"])
@login_required
def expense_delete(eid):
    if session.get("role") not in ("admin", "accountant"):
        return redirect(url_for("incoming_invoices_list"))
    if not validate_csrf():
        return redirect(url_for("incoming_invoices_list"))
    org_id = session["org_id"]
    conn   = get_connection()
    c      = conn.cursor()
    c.execute("SELECT proof_image FROM expense_items WHERE id=? AND org_id=?", (eid, org_id))
    row = c.fetchone()
    if row and row["proof_image"]:
        try:
            _os.remove(_os.path.join(app.config["UPLOAD_FOLDER"], row["proof_image"]))
        except Exception:
            pass
    c.execute("DELETE FROM expense_items WHERE id=? AND org_id=?", (eid, org_id))
    conn.commit()
    conn.close()
    flash("تم حذف البند", "warning")
    return redirect(url_for("incoming_invoices_list") + "#expense-card")


# ══════════════════════════════════════════
# العاملين
# ══════════════════════════════════════════
@app.route("/workers")
@login_required
def workers_list():
    org_id = session["org_id"]
    conn = get_connection()
    c = conn.cursor()
    c.execute("SELECT * FROM workers WHERE org_id=? ORDER BY id DESC", (org_id,))
    workers = [dict(r) for r in c.fetchall()]
    for w in workers:
        w['monthly_salary'] = w.get('monthly_salary') or 0
        c.execute("SELECT * FROM worker_bonuses WHERE worker_id=? AND org_id=? ORDER BY bonus_date DESC",
                  (w['id'], org_id))
        w['bonuses'] = [dict(b) for b in c.fetchall()]
        w['total_bonuses'] = sum(b['amount'] for b in w['bonuses'])
    c.execute("SELECT name FROM programs WHERE org_id=? AND is_active=1 ORDER BY name", (org_id,))
    programs = [r["name"] for r in c.fetchall()]
    conn.close()
    return render_template("workers.html", workers=workers, programs=programs)


@app.route("/add_worker", methods=["POST"])
@login_required
def add_worker():
    if session.get("role") == "observer":
        flash("مراقب: لا تملك صلاحية الإضافة", "danger")
        return redirect(url_for("workers_list"))
    org_id = session["org_id"]
    full_name      = request.form.get("full_name", "").strip()
    id_number      = request.form.get("id_number", "").strip()
    project_name   = request.form.get("project_name", "").strip()
    job_type       = request.form.get("job_type", "").strip()
    monthly_salary = request.form.get("monthly_salary", "0").strip()
    notes          = request.form.get("notes", "").strip()
    if not full_name:
        flash("يرجى إدخال اسم العامل", "danger")
        return redirect(url_for("workers_list"))
    try:
        salary = float(monthly_salary)
    except ValueError:
        salary = 0.0
    conn = get_connection()
    c = conn.cursor()
    c.execute(
        "INSERT INTO workers (org_id, full_name, id_number, project_name, job_type, monthly_salary, notes) VALUES (?,?,?,?,?,?,?)",
        (org_id, full_name, id_number, project_name, job_type, salary, notes)
    )
    conn.commit()
    conn.close()
    flash("✅ تم إضافة العامل", "success")
    return redirect(url_for("workers_list"))


@app.route("/delete_worker/<int:id>")
@login_required
def delete_worker(id):
    if session.get("role") not in ["admin", "accountant"]:
        flash("ليس لديك صلاحية الحذف", "danger")
        return redirect(url_for("workers_list"))
    conn = get_connection()
    c = conn.cursor()
    c.execute("DELETE FROM workers WHERE id=? AND org_id=?", (id, session["org_id"]))
    conn.commit()
    conn.close()
    flash("تم حذف العامل", "warning")
    return redirect(url_for("workers_list"))


@app.route("/edit_worker/<int:id>", methods=["POST"])
@login_required
def edit_worker(id):
    if session.get("role") not in ["admin", "accountant"]:
        return jsonify({"ok": False, "err": "forbidden"}), 403
    org_id = session["org_id"]
    full_name     = request.form.get("full_name", "").strip()
    id_number     = request.form.get("id_number", "").strip()
    project_name  = request.form.get("project_name", "").strip()
    job_type      = request.form.get("job_type", "").strip()
    monthly_salary= request.form.get("monthly_salary", "0").strip()
    notes         = request.form.get("notes", "").strip()
    if not full_name:
        return jsonify({"ok": False, "err": "name_required"}), 400
    try:
        salary = float(monthly_salary)
    except ValueError:
        salary = 0.0
    conn = get_connection()
    c    = conn.cursor()
    c.execute("""UPDATE workers
                 SET full_name=?, id_number=?, project_name=?, job_type=?, monthly_salary=?, notes=?
                 WHERE id=? AND org_id=?""",
              (full_name, id_number, project_name, job_type, salary, notes, id, org_id))
    conn.commit()
    conn.close()
    return jsonify({"ok": True})


@app.route("/worker/<int:wid>/bonus/add", methods=["POST"])
@login_required
def add_worker_bonus(wid):
    if session.get("role") not in ["admin", "accountant"]:
        return jsonify({"ok": False}), 403
    if not validate_csrf():
        return jsonify({"ok": False}), 403
    org_id     = session["org_id"]
    amount     = request.form.get("amount", "0").strip()
    bonus_date = request.form.get("bonus_date", "").strip()
    notes      = request.form.get("notes", "").strip()
    if not bonus_date:
        return jsonify({"ok": False, "err": "date_required"}), 400
    try:
        amount = float(amount)
    except ValueError:
        amount = 0.0
    from datetime import datetime as _dt
    conn = get_connection()
    c    = conn.cursor()
    c.execute("SELECT id FROM workers WHERE id=? AND org_id=?", (wid, org_id))
    if not c.fetchone():
        conn.close()
        return jsonify({"ok": False}), 404
    c.execute(
        "INSERT INTO worker_bonuses (worker_id, org_id, amount, bonus_date, notes, created_at) VALUES (?,?,?,?,?,?)",
        (wid, org_id, amount, bonus_date, notes, _dt.now().strftime("%Y-%m-%d %H:%M:%S"))
    )
    conn.commit()
    new_id = c.lastrowid
    conn.close()
    return jsonify({"ok": True, "id": new_id, "amount": amount, "date": bonus_date, "notes": notes})


@app.route("/worker/bonus/<int:bid>/delete", methods=["POST"])
@login_required
def delete_worker_bonus(bid):
    if session.get("role") not in ["admin", "accountant"]:
        return jsonify({"ok": False}), 403
    if not validate_csrf():
        return jsonify({"ok": False}), 403
    org_id = session["org_id"]
    conn = get_connection()
    conn.execute("DELETE FROM worker_bonuses WHERE id=? AND org_id=?", (bid, org_id))
    conn.commit()
    conn.close()
    return jsonify({"ok": True})


@app.route("/outgoing_invoice/<int:id>/delete")
@invoices_required
def delete_outgoing_invoice(id):
    """حذف فاتورة صرف مع إعادة كل الكميات للمخزن (أدمن فقط)"""
    if session.get("role") != "admin":
        flash("هذا الإجراء مخصص للأدمن فقط", "danger")
        return redirect(url_for("outgoing_invoices_list"))
    org_id = session["org_id"]
    conn = get_connection()
    c = conn.cursor()
    c.execute("SELECT id FROM outgoing_invoices WHERE id=? AND org_id=?", (id, org_id))
    if not c.fetchone():
        conn.close()
        flash("الفاتورة غير موجودة", "danger")
        return redirect(url_for("outgoing_invoices_list"))

    # إعادة كل الكميات للمخزن
    c.execute("SELECT product_id, quantity, unit_price FROM outgoing_invoice_items WHERE invoice_id=?", (id,))
    items = c.fetchall()
    for item in items:
        c.execute("""
            SELECT id FROM stock_batches
            WHERE product_id=? AND org_id=? ORDER BY id DESC LIMIT 1
        """, (item["product_id"], org_id))
        batch = c.fetchone()
        if batch:
            c.execute("UPDATE stock_batches SET quantity_remaining = quantity_remaining + ? WHERE id=?",
                      (item["quantity"], batch["id"]))
        else:
            c.execute("""
                INSERT INTO stock_batches (org_id, product_id, quantity_remaining, unit_price, entry_date)
                VALUES (?,?,?,?,date('now','localtime'))
            """, (org_id, item["product_id"], item["quantity"], item["unit_price"]))

    c.execute("DELETE FROM outgoing_invoice_items WHERE invoice_id=?", (id,))
    c.execute("DELETE FROM outgoing_invoices WHERE id=? AND org_id=?", (id, org_id))
    conn.commit()
    conn.close()
    flash("✅ تم حذف الفاتورة وإعادة جميع الكميات للمخزن", "success")
    return redirect(url_for("outgoing_invoices_list"))


@app.route("/outgoing_invoice/<int:id>/close")
@invoices_required
def close_outgoing_invoice(id):
    org_id = session["org_id"]
    conn = get_connection()
    c = conn.cursor()
    c.execute("UPDATE outgoing_invoices SET is_closed=1 WHERE id=? AND org_id=?", (id, org_id))
    conn.commit()
    conn.close()
    flash("✅ تم إغلاق فاتورة الصرف", "success")
    return redirect(url_for("outgoing_invoices_list"))


@app.route("/outgoing_invoice/<int:id>/update", methods=["POST"])
@invoices_required
def update_outgoing_invoice(id):
    """تعديل رأس فاتورة الصرف المفتوحة + أصنافها (أدمن فقط)"""
    if session.get("role") != "admin":
        flash("هذا الإجراء مخصص للأدمن فقط", "danger")
        return redirect(url_for("outgoing_invoices_list"))
    org_id = session["org_id"]
    conn = get_connection()
    c = conn.cursor()
    c.execute("SELECT is_closed FROM outgoing_invoices WHERE id=? AND org_id=?", (id, org_id))
    inv = c.fetchone()
    if not inv or inv["is_closed"]:
        conn.close()
        flash("الفاتورة مغلقة أو غير موجودة", "danger")
        return redirect(url_for("outgoing_invoices_list"))

    # تعديل الرأس
    invoice_number = request.form.get("invoice_number", "").strip()
    invoice_date   = request.form.get("invoice_date", "").strip()
    beneficiary    = request.form.get("beneficiary", "").strip()
    notes          = request.form.get("notes", "").strip()
    c.execute("""
        UPDATE outgoing_invoices
        SET invoice_number=?, invoice_date=?, beneficiary=?, notes=?
        WHERE id=? AND org_id=?
    """, (invoice_number, invoice_date, beneficiary, notes, id, org_id))

    # تعديل الأصناف
    item_ids    = request.form.getlist("item_id[]")
    prod_names  = request.form.getlist("prod_name[]")
    unit_prices = request.form.getlist("unit_price[]")

    for iid, pname, uprice in zip(item_ids, prod_names, unit_prices):
        try:
            price_f = float(uprice)
            c.execute("SELECT quantity FROM outgoing_invoice_items WHERE id=? AND invoice_id=?", (iid, id))
            row = c.fetchone()
            if not row:
                continue
            qty = row["quantity"]
            # تحديث اسم الصنف في جدول المنتجات إن وجد
            c.execute("SELECT product_id FROM outgoing_invoice_items WHERE id=?", (iid,))
            pid_row = c.fetchone()
            if pid_row and pname.strip():
                c.execute("UPDATE products SET name=? WHERE id=? AND org_id=?",
                          (pname.strip(), pid_row["product_id"], org_id))
            c.execute("""
                UPDATE outgoing_invoice_items
                SET unit_price=?, total_price=?
                WHERE id=? AND invoice_id=?
            """, (price_f, qty * price_f, iid, id))
        except (ValueError, TypeError):
            pass

    conn.commit()
    conn.close()
    flash("✅ تم حفظ التعديلات", "success")
    return redirect(url_for("outgoing_invoices_list") + f"#inv-{id}")


@app.route("/outgoing_invoice/item/<int:item_id>/delete")
@invoices_required
def delete_outgoing_item(item_id):
    """حذف صنف من فاتورة الصرف المفتوحة مع إعادة كميته للمخزن"""
    if session.get("role") != "admin":
        flash("هذا الإجراء مخصص للأدمن فقط", "danger")
        return redirect(url_for("outgoing_invoices_list"))
    org_id = session["org_id"]
    conn = get_connection()
    c = conn.cursor()

    # جلب بيانات الصنف + التحقق من أن الفاتورة مفتوحة
    c.execute("""
        SELECT oi.id, oi.invoice_id, oi.product_id, oi.quantity, oi.unit_price
        FROM outgoing_invoice_items oi
        JOIN outgoing_invoices inv ON inv.id = oi.invoice_id
        WHERE oi.id=? AND inv.org_id=? AND COALESCE(inv.is_closed,0)=0
    """, (item_id, org_id))
    item = c.fetchone()
    if not item:
        conn.close()
        flash("الصنف غير موجود أو الفاتورة مغلقة", "danger")
        return redirect(url_for("outgoing_invoices_list"))

    invoice_id = item["invoice_id"]
    product_id = item["product_id"]
    qty        = item["quantity"]
    price      = item["unit_price"]

    # إعادة الكمية للمخزن: أضف للدفعة الأخيرة الموجودة أو أنشئ دفعة جديدة
    c.execute("""
        SELECT id FROM stock_batches
        WHERE product_id=? AND org_id=?
        ORDER BY id DESC LIMIT 1
    """, (product_id, org_id))
    batch = c.fetchone()
    if batch:
        c.execute("UPDATE stock_batches SET quantity_remaining = quantity_remaining + ? WHERE id=?",
                  (qty, batch["id"]))
    else:
        c.execute("""
            INSERT INTO stock_batches (org_id, product_id, quantity_remaining, unit_price, entry_date)
            VALUES (?,?,?,?,date('now','localtime'))
        """, (org_id, product_id, qty, price))

    # حذف الصنف
    c.execute("DELETE FROM outgoing_invoice_items WHERE id=?", (item_id,))
    conn.commit()
    conn.close()
    flash("✅ تم حذف الصنف وإعادة الكمية للمخزن", "success")
    return redirect(url_for("outgoing_invoices_list") + f"#inv-{invoice_id}")


# ══════════════════════════════════════════
# الشات الداخلي (أدمن ↔ مراقب)
# ══════════════════════════════════════════
def _chat_allowed():
    return session.get("role") in ("admin", "observer")

@app.route("/messages")
@login_required
def messages_page():
    if not _chat_allowed():
        flash("غير مصرّح", "danger")
        return redirect(url_for("dashboard"))
    org_id  = session["org_id"]
    user_id = session["user_id"]
    conn = get_connection()
    c    = conn.cursor()
    # جلب آخر 200 رسالة
    c.execute("""
        SELECT id, sender_id, sender_name, sender_role, content, attachment,
               COALESCE(edited,0) AS edited, created_at
        FROM messages WHERE org_id=?
        ORDER BY id DESC LIMIT 200
    """, (org_id,))
    msgs = list(reversed(c.fetchall()))
    # تحديث آخر رسالة مقروءة لهذا المستخدم
    if msgs:
        last_id = msgs[-1]["id"]
        c.execute("""
            INSERT INTO message_reads(user_id, last_msg_id) VALUES(?,?)
            ON CONFLICT(user_id) DO UPDATE SET last_msg_id=excluded.last_msg_id
        """, (user_id, last_id))
        conn.commit()
    conn.close()
    return render_template("messages.html", msgs=msgs)


@app.route("/messages/data")
@login_required
def messages_data():
    """إرجاع آخر 200 رسالة كـ JSON للـ floating widget"""
    if not _chat_allowed():
        return jsonify({"ok": False}), 403
    org_id  = session["org_id"]
    user_id = session["user_id"]
    conn = get_connection()
    c    = conn.cursor()
    c.execute("""
        SELECT id, sender_id, sender_name, sender_role, content, attachment,
               COALESCE(edited,0) AS edited, created_at
        FROM messages WHERE org_id=?
        ORDER BY id DESC LIMIT 200
    """, (org_id,))
    rows = list(reversed(c.fetchall()))
    if rows:
        c.execute("""
            INSERT INTO message_reads(user_id, last_msg_id) VALUES(?,?)
            ON CONFLICT(user_id) DO UPDATE SET last_msg_id=excluded.last_msg_id
        """, (user_id, rows[-1]["id"]))
        conn.commit()
    conn.close()
    return jsonify({"ok": True, "messages": [dict(r) for r in rows]})


@app.route("/messages/send", methods=["POST"])
@login_required
def messages_send():
    if not _chat_allowed():
        return jsonify({"ok": False}), 403

    # قبول JSON أو FormData (مع ملف مرفق)
    if request.is_json:
        data    = request.json or {}
        content = data.get("content", "").strip()
        attach  = data.get("attachment", "")
    else:
        content = request.form.get("content", "").strip()
        attach  = ""
        f = request.files.get("attachment")
        if f and f.filename:
            ext = f.filename.rsplit(".", 1)[-1].lower() if "." in f.filename else ""
            if ext in ALLOWED_CHAT_EXT:
                import uuid as _uuid
                safe_name = f"{_uuid.uuid4().hex}.{ext}"
                f.save(os.path.join(CHAT_UPLOAD_FOLDER, safe_name))
                attach = safe_name

    if not content and not attach:
        return jsonify({"ok": False, "err": "empty"}), 400
    if len(content) > 2000:
        return jsonify({"ok": False, "err": "too_long"}), 400

    org_id   = session["org_id"]
    user_id  = session["user_id"]
    name     = session.get("full_name") or session.get("user", "")
    role     = session.get("role", "")
    conn = get_connection()
    c    = conn.cursor()
    c.execute("""
        INSERT INTO messages(org_id, sender_id, sender_name, sender_role, content, attachment)
        VALUES(?,?,?,?,?,?)
    """, (org_id, user_id, name, role, content, attach or None))
    new_id = c.lastrowid
    c.execute("""
        INSERT INTO message_reads(user_id, last_msg_id) VALUES(?,?)
        ON CONFLICT(user_id) DO UPDATE SET last_msg_id=excluded.last_msg_id
    """, (user_id, new_id))
    conn.commit()
    conn.close()
    from datetime import datetime as _dt
    now_str = _dt.now().strftime("%Y-%m-%d %H:%M:%S")
    return jsonify({"ok": True, "msg": {
        "id": new_id, "sender_id": user_id, "sender_name": name,
        "sender_role": role, "content": content,
        "attachment": attach or None, "edited": 0, "created_at": now_str
    }})


@app.route("/messages/poll")
@login_required
def messages_poll():
    if not _chat_allowed():
        return jsonify({"ok": False, "messages": []}), 403
    org_id  = session["org_id"]
    user_id = session["user_id"]
    # يقبل 'last' أو 'since' للتوافق
    since = request.args.get("last", request.args.get("since", 0), type=int)
    conn = get_connection()
    c    = conn.cursor()
    c.execute("""
        SELECT id, sender_id, sender_name, sender_role, content, attachment,
               COALESCE(edited,0) AS edited, created_at
        FROM messages WHERE org_id=? AND id>?
        ORDER BY id ASC LIMIT 50
    """, (org_id, since))
    rows = c.fetchall()
    if rows:
        c.execute("""
            INSERT INTO message_reads(user_id, last_msg_id) VALUES(?,?)
            ON CONFLICT(user_id) DO UPDATE SET last_msg_id=excluded.last_msg_id
        """, (user_id, rows[-1]["id"]))
        conn.commit()
    conn.close()
    return jsonify({"ok": True, "messages": [dict(r) for r in rows]})


@app.route("/messages/unread_count")
@login_required
def messages_unread_count():
    if not _chat_allowed():
        return {"count": 0}
    org_id  = session["org_id"]
    user_id = session["user_id"]
    conn = get_connection()
    c    = conn.cursor()
    c.execute("SELECT last_msg_id FROM message_reads WHERE user_id=?", (user_id,))
    row    = c.fetchone()
    last   = row["last_msg_id"] if row else 0
    c.execute("SELECT COUNT(*) as n FROM messages WHERE org_id=? AND id>?", (org_id, last))
    count  = c.fetchone()["n"]
    conn.close()
    return {"count": count}


# ── رفع مرفق للشات ──
CHAT_UPLOAD_FOLDER = os.path.join("static", "chat_uploads")
os.makedirs(CHAT_UPLOAD_FOLDER, exist_ok=True)

ALLOWED_CHAT_EXT = {
    "jpg","jpeg","png","gif","webp","bmp",   # صور
    "pdf","doc","docx","xls","xlsx",          # مستندات
    "txt","csv","zip","rar"                   # أخرى
}

@app.route("/messages/upload", methods=["POST"])
@login_required
def messages_upload():
    if not _chat_allowed():
        return {"ok": False}, 403
    f = request.files.get("file")
    if not f or not f.filename:
        return {"ok": False, "err": "no_file"}, 400
    ext = f.filename.rsplit(".", 1)[-1].lower() if "." in f.filename else ""
    if ext not in ALLOWED_CHAT_EXT:
        return {"ok": False, "err": "type_not_allowed"}, 400
    import uuid
    safe_name = f"{uuid.uuid4().hex}.{ext}"
    f.save(os.path.join(CHAT_UPLOAD_FOLDER, safe_name))
    return {"ok": True, "filename": safe_name, "original": f.filename}


@app.route("/messages/<int:msg_id>/edit", methods=["POST"])
@login_required
def messages_edit(msg_id):
    if not _chat_allowed():
        return jsonify({"ok": False}), 403
    content = ((request.json or {}).get("content") if request.is_json
               else request.form.get("content", ""))
    content = (content or "").strip()
    if not content or len(content) > 2000:
        return jsonify({"ok": False, "err": "invalid"}), 400
    user_id  = session["user_id"]
    org_id   = session["org_id"]
    conn = get_connection()
    c    = conn.cursor()
    c.execute("SELECT sender_id, org_id FROM messages WHERE id=?", (msg_id,))
    row = c.fetchone()
    if not row or row["org_id"] != org_id or row["sender_id"] != user_id:
        conn.close()
        return {"ok": False, "err": "forbidden"}, 403
    c.execute("UPDATE messages SET content=?, edited=1 WHERE id=?", (content, msg_id))
    conn.commit()
    conn.close()
    return {"ok": True}


@app.route("/messages/<int:msg_id>/delete", methods=["POST"])
@login_required
def messages_delete(msg_id):
    if not _chat_allowed():
        return {"ok": False}, 403
    user_id = session["user_id"]
    org_id  = session["org_id"]
    role    = session.get("role", "")
    conn = get_connection()
    c    = conn.cursor()
    c.execute("SELECT sender_id, org_id, attachment FROM messages WHERE id=?", (msg_id,))
    row = c.fetchone()
    if not row or row["org_id"] != org_id:
        conn.close()
        return {"ok": False, "err": "not_found"}, 404
    # يسمح للمرسل نفسه أو الأدمن بالحذف
    if row["sender_id"] != user_id and role != "admin":
        conn.close()
        return {"ok": False, "err": "forbidden"}, 403
    # حذف الملف المرفق إن وجد
    if row["attachment"]:
        try:
            os.remove(os.path.join(CHAT_UPLOAD_FOLDER, row["attachment"]))
        except Exception:
            pass
    c.execute("DELETE FROM messages WHERE id=?", (msg_id,))
    conn.commit()
    conn.close()
    return {"ok": True}


@app.route("/translate", methods=["POST"])
@login_required
def translate_text():
    """ترجمة نص عبر MyMemory — سيرفر سايد"""
    if not _chat_allowed():
        return jsonify({"ok": False}), 403
    data = request.json or {}
    text = data.get("text", "").strip()
    dst  = data.get("dst", "ar")
    if not text or len(text) > 1000:
        return jsonify({"ok": False, "err": "invalid"}), 400

    import re as _re, urllib.request as _ur, urllib.parse as _up, json as _json2
    # كشف لغة المصدر
    if _re.search(r'[؀-ۿ]', text):
        src = 'ar'
    elif _re.search(r'[çğışöüÇĞİŞÖÜ]', text):
        src = 'tr'
    else:
        src = 'en'

    if src == dst:
        return jsonify({"ok": True, "same": True})

    try:
        url = (
            "https://api.mymemory.translated.net/get?"
            + _up.urlencode({"q": text, "langpair": f"{src}|{dst}"})
        )
        with _ur.urlopen(url, timeout=8) as r:
            d = _json2.loads(r.read())
        t = (d.get("responseData", {}).get("translatedText") or "").strip()
        if not t or t == text or "MYMEMORY" in t.upper():
            return jsonify({"ok": False, "err": "no_result"})
        return jsonify({"ok": True, "translated": t, "src": src})
    except Exception as e:
        return jsonify({"ok": False, "err": str(e)}), 500


@app.route("/messages/clear", methods=["POST"])
@login_required
def messages_clear():
    """يمسح كامل سجل المحادثة — للأدمن والمراقب"""
    if not _chat_allowed():
        return {"ok": False}, 403
    org_id = session["org_id"]
    conn = get_connection()
    c    = conn.cursor()
    # احذف المرفقات أولاً
    c.execute("SELECT attachment FROM messages WHERE org_id=? AND attachment IS NOT NULL", (org_id,))
    for row in c.fetchall():
        try:
            os.remove(os.path.join(CHAT_UPLOAD_FOLDER, row["attachment"]))
        except Exception:
            pass
    c.execute("DELETE FROM messages WHERE org_id=?", (org_id,))
    c.execute("DELETE FROM message_reads WHERE user_id IN (SELECT id FROM users WHERE org_id=?)", (org_id,))
    conn.commit()
    conn.close()
    return {"ok": True}



# ══════════════════════════════════════════
# توحيد المستفيدين بين المؤسسات
# ══════════════════════════════════════════

@app.route("/beneficiaries/check_duplicate")
@login_required
def beneficiary_check_duplicate():
    """تحقق إذا المستفيد (اسم + رقم هوية) موجود في مؤسسة أخرى"""
    org_id   = session["org_id"]
    name     = request.args.get("name", "").strip()
    id_num   = request.args.get("id_number", "").strip()
    if not name or not id_num:
        return jsonify({"found": False})
    conn = get_connection()
    c    = conn.cursor()
    c.execute("""
        SELECT b.id, b.full_name, b.id_number, o.name AS org_name, o.id AS org_id
        FROM beneficiaries b
        JOIN organizations o ON o.id = b.org_id
        WHERE b.org_id != ?
          AND TRIM(LOWER(b.full_name))   = TRIM(LOWER(?))
          AND TRIM(b.id_number) = TRIM(?)
        LIMIT 1
    """, (org_id, name, id_num))
    row = c.fetchone()
    # هل في طلب معلق بالفعل؟
    pending = False
    if row:
        c.execute("""
            SELECT id FROM beneficiary_share_requests
            WHERE requester_org_id=? AND beneficiary_id=? AND status='pending'
        """, (org_id, row["id"]))
        pending = bool(c.fetchone())
    conn.close()
    if row:
        return jsonify({"found": True, "pending": pending,
                        "beneficiary_id": row["id"],
                        "org_name": row["org_name"],
                        "org_id": row["org_id"]})
    return jsonify({"found": False})


@app.route("/beneficiaries/request_share", methods=["POST"])
@login_required
def beneficiary_request_share():
    """إرسال طلب إضافة مستفيد من مؤسسة أخرى"""
    if session.get("role") != "admin":
        return jsonify({"ok": False}), 403
    csrf = request.form.get("_csrf_token") or request.headers.get("X-CSRF-Token", "")
    if csrf != session.get("_csrf"):
        return jsonify({"ok": False}), 403
    org_id         = session["org_id"]
    beneficiary_id = int(request.form.get("beneficiary_id", 0))
    if not beneficiary_id:
        return jsonify({"ok": False, "msg": "معرّف المستفيد مفقود"}), 400
    conn = get_connection()
    c    = conn.cursor()
    # تحقق أن المستفيد موجود في مؤسسة أخرى
    c.execute("SELECT org_id FROM beneficiaries WHERE id=?", (beneficiary_id,))
    row = c.fetchone()
    if not row or row["org_id"] == org_id:
        conn.close()
        return jsonify({"ok": False, "msg": "المستفيد غير موجود أو ينتمي لمؤسستك"}), 400
    owner_org_id = row["org_id"]
    # تحقق إنه ما في طلب معلق بالفعل
    c.execute("""
        SELECT id FROM beneficiary_share_requests
        WHERE requester_org_id=? AND beneficiary_id=? AND status='pending'
    """, (org_id, beneficiary_id))
    if c.fetchone():
        conn.close()
        return jsonify({"ok": False, "msg": "يوجد طلب معلق بالفعل"})
    from datetime import datetime as _dt
    c.execute("""
        INSERT INTO beneficiary_share_requests
            (requester_org_id, owner_org_id, beneficiary_id, status, created_at)
        VALUES (?,?,?,'pending',?)
    """, (org_id, owner_org_id, beneficiary_id, _dt.now().strftime("%Y-%m-%d %H:%M:%S")))
    conn.commit()
    conn.close()
    return jsonify({"ok": True})


@app.route("/notifications")
@login_required
def notifications():
    if session.get("role") != "admin":
        flash("غير مصرّح", "danger")
        return redirect(url_for("dashboard"))
    org_id = session["org_id"]
    conn   = get_connection()
    c      = conn.cursor()
    # طلبات واردة — أنا المالك
    c.execute("""
        SELECT r.id, r.status, r.created_at,
               b.full_name, b.id_number,
               o.name AS requester_name
        FROM beneficiary_share_requests r
        JOIN beneficiaries b ON b.id = r.beneficiary_id
        JOIN organizations o ON o.id = r.requester_org_id
        WHERE r.owner_org_id = ?
        ORDER BY r.id DESC
    """, (org_id,))
    incoming = c.fetchall()
    # طلبات صادرة — أنا الطالب
    c.execute("""
        SELECT r.id, r.status, r.created_at,
               b.full_name, b.id_number,
               o.name AS owner_name
        FROM beneficiary_share_requests r
        JOIN beneficiaries b ON b.id = r.beneficiary_id
        JOIN organizations o ON o.id = r.owner_org_id
        WHERE r.requester_org_id = ?
        ORDER BY r.id DESC
    """, (org_id,))
    outgoing = c.fetchall()
    conn.close()
    return render_template("notifications.html",
                           incoming=incoming, outgoing=outgoing)


@app.route("/notifications/<int:req_id>/approve", methods=["POST"])
@login_required
def notification_approve(req_id):
    if session.get("role") != "admin":
        return jsonify({"ok": False}), 403
    csrf = request.form.get("_csrf_token") or request.headers.get("X-CSRF-Token", "")
    if csrf != session.get("_csrf"):
        return jsonify({"ok": False}), 403
    org_id = session["org_id"]
    conn   = get_connection()
    c      = conn.cursor()
    c.execute("""
        SELECT r.*, b.full_name, b.id_number, b.phone, b.address,
               b.family_size, b.marital_status, b.gender, b.children_count,
               b.beneficiary_type, b.notes
        FROM beneficiary_share_requests r
        JOIN beneficiaries b ON b.id = r.beneficiary_id
        WHERE r.id=? AND r.owner_org_id=? AND r.status='pending'
    """, (req_id, org_id))
    row = c.fetchone()
    if not row:
        conn.close()
        return jsonify({"ok": False, "msg": "الطلب غير موجود"}), 404
    # أضف المستفيد للمؤسسة الطالبة
    from datetime import datetime as _dt
    now = _dt.now().strftime("%Y-%m-%d %H:%M:%S")
    # احسب seq_num التالي للمؤسسة الطالبة
    c.execute("SELECT COALESCE(MAX(seq_num),0)+1 FROM beneficiaries WHERE org_id=?",
              (row["requester_org_id"],))
    next_seq = c.fetchone()[0]
    c.execute("""
        INSERT INTO beneficiaries
            (org_id, seq_num, full_name, id_number, phone, address,
             family_size, marital_status, gender, children_count,
             beneficiary_type, notes, created_at)
        VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?)
    """, (row["requester_org_id"], next_seq,
          row["full_name"], row["id_number"], row["phone"], row["address"],
          row["family_size"], row["marital_status"], row["gender"],
          row["children_count"], row["beneficiary_type"], row["notes"], now))
    # حدّث حالة الطلب
    c.execute("""
        UPDATE beneficiary_share_requests
        SET status='approved', resolved_at=?
        WHERE id=?
    """, (now, req_id))
    conn.commit()
    conn.close()
    return jsonify({"ok": True})


@app.route("/notifications/<int:req_id>/reject", methods=["POST"])
@login_required
def notification_reject(req_id):
    if session.get("role") != "admin":
        return jsonify({"ok": False}), 403
    csrf = request.form.get("_csrf_token") or request.headers.get("X-CSRF-Token", "")
    if csrf != session.get("_csrf"):
        return jsonify({"ok": False}), 403
    org_id = session["org_id"]
    conn   = get_connection()
    c      = conn.cursor()
    from datetime import datetime as _dt
    c.execute("""
        UPDATE beneficiary_share_requests
        SET status='rejected', resolved_at=?
        WHERE id=? AND owner_org_id=? AND status='pending'
    """, (_dt.now().strftime("%Y-%m-%d %H:%M:%S"), req_id, org_id))
    conn.commit()
    conn.close()
    return jsonify({"ok": True})


@app.route("/notifications/unread_count")
@login_required
def notifications_unread_count():
    if session.get("role") != "admin":
        return jsonify({"count": 0})
    org_id = session["org_id"]
    conn   = get_connection()
    c      = conn.cursor()
    c.execute("""
        SELECT COUNT(*) FROM beneficiary_share_requests
        WHERE owner_org_id=? AND status='pending'
    """, (org_id,))
    count = c.fetchone()[0]
    conn.close()
    return jsonify({"count": count})


# ══════════════════════════════════════════
# الشبكة — التواصل بين المؤسسات
# ══════════════════════════════════════════

@app.route("/network")
@login_required
def network():
    if session.get("role") != "admin":
        flash("غير مصرّح", "danger")
        return redirect(url_for("dashboard"))
    org_id = session["org_id"]
    conn   = get_connection()
    c      = conn.cursor()
    # جميع المؤسسات ما عدا المؤسسة الحالية
    c.execute("SELECT id, name FROM organizations WHERE id != ? ORDER BY name", (org_id,))
    orgs = c.fetchall()
    # عدد الرسائل غير المقروءة لكل مؤسسة
    unread = {}
    for org in orgs:
        c.execute("""
            SELECT COUNT(*) FROM org_messages
            WHERE to_org_id=? AND from_org_id=?
              AND id > COALESCE(
                (SELECT last_msg_id FROM org_message_reads
                 WHERE user_id=? AND partner_org=?), 0)
        """, (org_id, org["id"], session["user_id"], org["id"]))
        unread[org["id"]] = c.fetchone()[0]
    conn.close()
    return render_template("network.html", orgs=orgs, unread=unread)


@app.route("/network/chat/<int:partner_id>")
@login_required
def network_chat(partner_id):
    if session.get("role") != "admin":
        flash("غير مصرّح", "danger")
        return redirect(url_for("dashboard"))
    org_id  = session["org_id"]
    user_id = session["user_id"]
    conn    = get_connection()
    c       = conn.cursor()
    c.execute("SELECT id, name FROM organizations WHERE id=?", (partner_id,))
    partner = c.fetchone()
    if not partner:
        conn.close()
        flash("المؤسسة غير موجودة", "danger")
        return redirect(url_for("network"))
    # جلب الرسائل بين المؤسستين
    c.execute("""
        SELECT * FROM org_messages
        WHERE (from_org_id=? AND to_org_id=?)
           OR (from_org_id=? AND to_org_id=?)
        ORDER BY id ASC LIMIT 200
    """, (org_id, partner_id, partner_id, org_id))
    messages = c.fetchall()
    # تحديث آخر قراءة
    if messages:
        last_id = messages[-1]["id"]
        c.execute("""
            INSERT INTO org_message_reads(user_id, partner_org, last_msg_id)
            VALUES(?,?,?)
            ON CONFLICT(user_id, partner_org)
            DO UPDATE SET last_msg_id=excluded.last_msg_id
        """, (user_id, partner_id, last_id))
        conn.commit()
    conn.close()
    return render_template("network_chat.html",
                           partner=partner, messages=messages,
                           my_org_id=org_id)


@app.route("/network/send/<int:partner_id>", methods=["POST"])
@login_required
def network_send(partner_id):
    if session.get("role") != "admin":
        return jsonify({"ok": False}), 403
    org_id  = session["org_id"]
    user_id = session["user_id"]
    name    = session.get("name", "")
    content = request.form.get("content", "").strip()
    if not content:
        return jsonify({"ok": False, "msg": "الرسالة فارغة"}), 400
    csrf = request.form.get("_csrf_token") or request.headers.get("X-CSRF-Token", "")
    if csrf != session.get("_csrf"):
        return jsonify({"ok": False}), 403
    conn = get_connection()
    c    = conn.cursor()
    # تحقق أن المؤسسة موجودة
    c.execute("SELECT id FROM organizations WHERE id=?", (partner_id,))
    if not c.fetchone():
        conn.close()
        return jsonify({"ok": False, "msg": "مؤسسة غير موجودة"}), 404
    from datetime import datetime as _dt
    now = _dt.now().strftime("%Y-%m-%d %H:%M:%S")
    c.execute("""
        INSERT INTO org_messages(from_org_id, to_org_id, sender_id, sender_name, content, created_at)
        VALUES(?,?,?,?,?,?)
    """, (org_id, partner_id, user_id, name, content, now))
    new_id = c.lastrowid
    # تحديث آخر قراءة للمُرسِل
    c.execute("""
        INSERT INTO org_message_reads(user_id, partner_org, last_msg_id)
        VALUES(?,?,?)
        ON CONFLICT(user_id, partner_org)
        DO UPDATE SET last_msg_id=excluded.last_msg_id
    """, (user_id, partner_id, new_id))
    conn.commit()
    conn.close()
    return jsonify({"ok": True, "msg": {
        "id": new_id, "from_org_id": org_id,
        "sender_name": name, "content": content, "created_at": now
    }})


@app.route("/network/delete_msg/<int:msg_id>", methods=["POST"])
@login_required
def network_delete_msg(msg_id):
    if session.get("role") != "admin":
        return jsonify({"ok": False}), 403
    csrf = request.form.get("_csrf_token") or request.headers.get("X-CSRF-Token", "")
    if csrf != session.get("_csrf"):
        return jsonify({"ok": False}), 403
    org_id = session["org_id"]
    conn = get_connection()
    c    = conn.cursor()
    # يسمح فقط لمن أرسل الرسالة (من مؤسسته)
    c.execute("SELECT from_org_id FROM org_messages WHERE id=?", (msg_id,))
    row = c.fetchone()
    if not row or row["from_org_id"] != org_id:
        conn.close()
        return jsonify({"ok": False, "msg": "غير مصرّح"}), 403
    c.execute("DELETE FROM org_messages WHERE id=?", (msg_id,))
    conn.commit()
    conn.close()
    return jsonify({"ok": True})


@app.route("/network/clear/<int:partner_id>", methods=["POST"])
@login_required
def network_clear(partner_id):
    if session.get("role") != "admin":
        return jsonify({"ok": False}), 403
    csrf = request.form.get("_csrf_token") or request.headers.get("X-CSRF-Token", "")
    if csrf != session.get("_csrf"):
        return jsonify({"ok": False}), 403
    org_id = session["org_id"]
    conn = get_connection()
    c    = conn.cursor()
    # يحذف كل رسائل المحادثة من الطرفين
    c.execute("""
        DELETE FROM org_messages
        WHERE (from_org_id=? AND to_org_id=?)
           OR (from_org_id=? AND to_org_id=?)
    """, (org_id, partner_id, partner_id, org_id))
    conn.commit()
    conn.close()
    return jsonify({"ok": True})


@app.route("/network/poll/<int:partner_id>")
@login_required
def network_poll(partner_id):
    if session.get("role") != "admin":
        return jsonify({"ok": False}), 403
    org_id  = session["org_id"]
    user_id = session["user_id"]
    since   = int(request.args.get("since", 0))
    conn    = get_connection()
    c       = conn.cursor()
    c.execute("""
        SELECT * FROM org_messages
        WHERE ((from_org_id=? AND to_org_id=?) OR (from_org_id=? AND to_org_id=?))
          AND id > ?
        ORDER BY id ASC
    """, (org_id, partner_id, partner_id, org_id, since))
    rows = c.fetchall()
    if rows:
        c.execute("""
            INSERT INTO org_message_reads(user_id, partner_org, last_msg_id)
            VALUES(?,?,?)
            ON CONFLICT(user_id, partner_org)
            DO UPDATE SET last_msg_id=excluded.last_msg_id
        """, (user_id, partner_id, rows[-1]["id"]))
        conn.commit()
    conn.close()
    return jsonify({"ok": True, "messages": [dict(r) for r in rows]})


@app.route("/network/unread")
@login_required
def network_unread():
    if session.get("role") != "admin":
        return jsonify({"count": 0})
    org_id  = session["org_id"]
    user_id = session["user_id"]
    conn    = get_connection()
    c       = conn.cursor()
    c.execute("""
        SELECT COUNT(*) FROM org_messages m
        WHERE m.to_org_id = ?
          AND m.id > COALESCE(
            (SELECT last_msg_id FROM org_message_reads
             WHERE user_id=? AND partner_org=m.from_org_id), 0)
    """, (org_id, user_id))
    count = c.fetchone()[0]
    conn.close()
    return jsonify({"count": count})


# ══════════════════════════════════════════
# التقارير
# ══════════════════════════════════════════

@app.route("/ai_reports")
@login_required
def ai_reports():
    if session.get("role") not in ("admin", "observer"):
        flash("غير مصرّح", "danger")
        return redirect(url_for("dashboard"))
    return render_template("ai_reports.html")


@app.route("/ai_reports/generate", methods=["POST"])
@login_required
def ai_reports_generate():
    if session.get("role") not in ("admin", "observer"):
        return jsonify({"ok": False}), 403

    import io
    from datetime import datetime as _dt

    rtype  = request.form.get("rtype", "").strip()
    fmt    = request.form.get("fmt", "docx").strip()
    org_id = session["org_id"]
    lang   = session.get("lang", "ar")

    valid_types = ("purchases","outgoing","beneficiaries","workers","programs","stock","summary")
    if rtype not in valid_types:
        return jsonify({"ok": False, "msg": "نوع التقرير غير صالح"}), 400

    conn = get_connection()
    c    = conn.cursor()

    c.execute("SELECT name FROM organizations WHERE id=?", (org_id,))
    row_org  = c.fetchone()
    org_name = row_org["name"] if row_org else ""

    data = {}

    if rtype in ("purchases", "summary"):
        c.execute("""
            SELECT ii.seq_num, ii.invoice_number, ii.supplier, ii.invoice_date,
                   ii.is_closed, ii.is_paid,
                   COALESCE(SUM(iii.total_price),0) AS total
            FROM incoming_invoices ii
            LEFT JOIN incoming_invoice_items iii ON iii.invoice_id=ii.id
            WHERE ii.org_id=? GROUP BY ii.id ORDER BY ii.id
        """, (org_id,))
        data["purchases"] = c.fetchall()

    if rtype in ("outgoing", "summary"):
        c.execute("""
            SELECT oi.invoice_number, oi.invoice_date, oi.beneficiary, oi.is_closed,
                   COALESCE(SUM(oii.total_price),0) AS total
            FROM outgoing_invoices oi
            LEFT JOIN outgoing_invoice_items oii ON oii.invoice_id=oi.id
            WHERE oi.org_id=? GROUP BY oi.id ORDER BY oi.id
        """, (org_id,))
        data["outgoing"] = c.fetchall()

    if rtype in ("beneficiaries", "summary"):
        c.execute("""
            SELECT full_name, gender, id_number, phone,
                   marital_status, family_size, children_count,
                   beneficiary_type, notes
            FROM beneficiaries WHERE org_id=? ORDER BY id
        """, (org_id,))
        data["beneficiaries"] = c.fetchall()

    if rtype in ("workers", "summary"):
        c.execute("""
            SELECT full_name, id_number, project_name, job_type,
                   monthly_salary, notes
            FROM workers WHERE org_id=? ORDER BY id
        """, (org_id,))
        data["workers"] = c.fetchall()

    if rtype in ("programs", "summary"):
        c.execute("SELECT name FROM programs WHERE org_id=? ORDER BY id", (org_id,))
        data["programs"] = c.fetchall()

    if rtype in ("stock", "summary"):
        c.execute("""
            SELECT p.name, p.unit,
                   COALESCE(SUM(sb.quantity_remaining),0) AS qty
            FROM products p
            LEFT JOIN stock_batches sb ON sb.product_id=p.id AND sb.org_id=p.org_id
            WHERE p.org_id=? GROUP BY p.id ORDER BY p.name
        """, (org_id,))
        data["stock"] = c.fetchall()

    conn.close()

    T = {
        "ar": dict(title_purchases="تقرير المشتريات", title_outgoing="تقرير الصرف",
                   title_beneficiaries="تقرير المستفيدين", title_workers="تقرير العاملين",
                   title_programs="تقرير البرامج", title_stock="تقرير المخزون",
                   title_summary="التقرير الشامل",
                   no_data="لا توجد بيانات", total="الإجمالي", count="العدد الكلي",
                   open_="مفتوحة", closed="مغلقة", paid="مدفوعة", unpaid="غير مدفوعة",
                   h_purchases="فواتير المشتريات", h_outgoing="فواتير الصرف",
                   h_beneficiaries="المستفيدون", h_workers="العاملون",
                   h_programs="البرامج والمشاريع", h_stock="المخزون"),
        "tr": dict(title_purchases="Satın Alma Raporu", title_outgoing="Çıkış Faturaları Raporu",
                   title_beneficiaries="Yararlanıcılar Raporu", title_workers="Çalışanlar Raporu",
                   title_programs="Programlar Raporu", title_stock="Stok Raporu",
                   title_summary="Genel Özet Raporu",
                   no_data="Veri yok", total="Toplam", count="Toplam Sayı",
                   open_="Açık", closed="Kapalı", paid="Ödendi", unpaid="Ödenmedi",
                   h_purchases="Satın Alma Faturaları", h_outgoing="Çıkış Faturaları",
                   h_beneficiaries="Yararlanıcılar", h_workers="Çalışanlar",
                   h_programs="Programlar ve Projeler", h_stock="Stok"),
        "en": dict(title_purchases="Purchase Invoices Report", title_outgoing="Outgoing Report",
                   title_beneficiaries="Beneficiaries Report", title_workers="Workers Report",
                   title_programs="Programs Report", title_stock="Stock Report",
                   title_summary="Full Summary Report",
                   no_data="No data", total="Total", count="Total Count",
                   open_="Open", closed="Closed", paid="Paid", unpaid="Unpaid",
                   h_purchases="Purchase Invoices", h_outgoing="Outgoing Invoices",
                   h_beneficiaries="Beneficiaries", h_workers="Workers",
                   h_programs="Programs & Projects", h_stock="Stock"),
    }
    tx       = T.get(lang, T["ar"])
    title    = tx.get(f"title_{rtype}", "Report")
    now_str  = _dt.now().strftime("%Y-%m-%d %H:%M")
    is_rtl   = (lang == "ar")

    # ─────────────────────────────────────────
    # PDF — HTML يُفتح في نافذة ويُطبع
    # ─────────────────────────────────────────
    if fmt == "pdf":
        dir_attr = "rtl" if is_rtl else "ltr"
        th_align = "right" if is_rtl else "left"

        def _tbl(headers, rows, fn):
            if not rows:
                return f"<p class='nd'>{tx['no_data']}</p>"
            ths = "".join(f"<th>{h}</th>" for h in headers)
            trs = ""
            for r in rows:
                vals = fn(r)
                trs += "<tr>" + "".join(
                    f"<td>{str(v) if v is not None else ''}</td>" for v in vals
                ) + "</tr>"
            return f"<table><thead><tr>{ths}</tr></thead><tbody>{trs}</tbody></table>"

        body = ""

        if "purchases" in data:
            h = {"ar":["#","رقم الفاتورة","المورد","التاريخ","الإجمالي","الحالة","الدفع"],
                 "tr":["#","Fatura No","Tedarikçi","Tarih","Toplam","Durum","Ödeme"],
                 "en":["#","Invoice No","Supplier","Date","Total","Status","Payment"]}.get(lang,[])
            rows = data["purchases"]
            body += f"<h2>{tx['h_purchases']}</h2>"
            body += _tbl(h, rows, lambda r:[
                r["seq_num"] or"", r["invoice_number"] or"", r["supplier"] or"",
                r["invoice_date"] or"", f"{r['total']:.2f}",
                tx["closed"] if r["is_closed"] else tx["open_"],
                tx["paid"] if r["is_paid"] else tx["unpaid"]])
            body += f"<p class='sm'>{tx['total']}: {sum(r['total'] for r in rows):.2f}</p>"

        if "outgoing" in data:
            h = {"ar":["رقم الفاتورة","التاريخ","الجهة","الإجمالي","الحالة"],
                 "tr":["Fatura No","Tarih","Kurum","Toplam","Durum"],
                 "en":["Invoice No","Date","Beneficiary","Total","Status"]}.get(lang,[])
            rows = data["outgoing"]
            body += f"<h2>{tx['h_outgoing']}</h2>"
            body += _tbl(h, rows, lambda r:[
                r["invoice_number"] or"", r["invoice_date"] or"",
                r["beneficiary"] or"", f"{r['total']:.2f}",
                tx["closed"] if r["is_closed"] else tx["open_"]])
            body += f"<p class='sm'>{tx['total']}: {sum(r['total'] for r in rows):.2f}</p>"

        if "beneficiaries" in data:
            h = {"ar":["الاسم","الجنس","رقم الهوية","الجوال","الحالة","عدد الأفراد","النوع"],
                 "tr":["Ad Soyad","Cinsiyet","Kimlik","Telefon","Durum","Üye Sayısı","Tür"],
                 "en":["Name","Gender","ID","Phone","Status","Family Size","Type"]}.get(lang,[])
            rows = data["beneficiaries"]
            body += f"<h2>{tx['h_beneficiaries']}</h2>"
            body += _tbl(h, rows, lambda r:[
                r["full_name"] or"", r["gender"] or"", r["id_number"] or"",
                r["phone"] or"", r["marital_status"] or"",
                str(r["family_size"] or""), r["beneficiary_type"] or""])
            body += f"<p class='sm'>{tx['count']}: {len(rows)}</p>"

        if "workers" in data:
            h = {"ar":["الاسم","رقم الهوية","المشروع","طبيعة العمل","المكافأة","ملاحظات"],
                 "tr":["Ad","Kimlik","Proje","Görev","Maaş","Notlar"],
                 "en":["Name","ID","Project","Role","Salary","Notes"]}.get(lang,[])
            rows = data["workers"]
            body += f"<h2>{tx['h_workers']}</h2>"
            body += _tbl(h, rows, lambda r:[
                r["full_name"] or"", r["id_number"] or"", r["project_name"] or"",
                r["job_type"] or"",
                f"{r['monthly_salary']:.2f}" if r["monthly_salary"] else "0.00",
                r["notes"] or""])
            body += f"<p class='sm'>{tx['total']}: {sum((r['monthly_salary'] or 0) for r in rows):.2f}</p>"

        if "programs" in data:
            rows = data["programs"]
            body += f"<h2>{tx['h_programs']}</h2><ol>"
            for r in rows:
                body += f"<li>{r['name']}</li>"
            body += "</ol>"

        if "stock" in data:
            h = {"ar":["الصنف","الوحدة","الكمية المتاحة"],
                 "tr":["Ürün","Birim","Mevcut Miktar"],
                 "en":["Item","Unit","Available Qty"]}.get(lang,[])
            rows = data["stock"]
            body += f"<h2>{tx['h_stock']}</h2>"
            body += _tbl(h, rows, lambda r:[r["name"] or"", r["unit"] or"", f"{r['qty']:.2f}"])

        # ── بناء HTML للـ PDF ──
        dir_attr = "rtl" if is_rtl else "ltr"
        th_align = "right" if is_rtl else "left"
        sum_align = "left" if is_rtl else "right"
        ol_pad   = "right" if is_rtl else "left"

        def _tbl(headers, rows, fn):
            if not rows:
                return "<p class='nd'>" + tx["no_data"] + "</p>"
            ths = "".join("<th>" + str(h) + "</th>" for h in headers)
            trs = ""
            for r in rows:
                vals = fn(r)
                trs += "<tr>" + "".join(
                    "<td>" + (str(v) if v is not None else "") + "</td>" for v in vals
                ) + "</tr>"
            return "<table><thead><tr>" + ths + "</tr></thead><tbody>" + trs + "</tbody></table>"

        body = ""

        if "purchases" in data:
            h = {"ar":["#","رقم الفاتورة","المورد","التاريخ","الإجمالي","الحالة","الدفع"],
                 "tr":["#","Fatura No","Tedarikçi","Tarih","Toplam","Durum","Ödeme"],
                 "en":["#","Invoice No","Supplier","Date","Total","Status","Payment"]}.get(lang,[])
            rows = data["purchases"]
            body += "<h2>" + tx["h_purchases"] + "</h2>"
            body += _tbl(h, rows, lambda r:[
                r["seq_num"] or "", r["invoice_number"] or "", r["supplier"] or "",
                r["invoice_date"] or "", "%.2f" % r["total"],
                tx["closed"] if r["is_closed"] else tx["open_"],
                tx["paid"] if r["is_paid"] else tx["unpaid"]])
            body += "<p class='sm'>" + tx["total"] + ": " + ("%.2f" % sum(r["total"] for r in rows)) + "</p>"

        if "outgoing" in data:
            h = {"ar":["رقم الفاتورة","التاريخ","الجهة","الإجمالي","الحالة"],
                 "tr":["Fatura No","Tarih","Kurum","Toplam","Durum"],
                 "en":["Invoice No","Date","Beneficiary","Total","Status"]}.get(lang,[])
            rows = data["outgoing"]
            body += "<h2>" + tx["h_outgoing"] + "</h2>"
            body += _tbl(h, rows, lambda r:[
                r["invoice_number"] or "", r["invoice_date"] or "",
                r["beneficiary"] or "", "%.2f" % r["total"],
                tx["closed"] if r["is_closed"] else tx["open_"]])
            body += "<p class='sm'>" + tx["total"] + ": " + ("%.2f" % sum(r["total"] for r in rows)) + "</p>"

        if "beneficiaries" in data:
            h = {"ar":["الاسم","الجنس","رقم الهوية","الجوال","الحالة","عدد الأفراد","النوع"],
                 "tr":["Ad Soyad","Cinsiyet","Kimlik","Telefon","Durum","Üye Sayısı","Tür"],
                 "en":["Name","Gender","ID","Phone","Status","Family Size","Type"]}.get(lang,[])
            rows = data["beneficiaries"]
            body += "<h2>" + tx["h_beneficiaries"] + "</h2>"
            body += _tbl(h, rows, lambda r:[
                r["full_name"] or "", r["gender"] or "", r["id_number"] or "",
                r["phone"] or "", r["marital_status"] or "",
                str(r["family_size"] or ""), r["beneficiary_type"] or ""])
            body += "<p class='sm'>" + tx["count"] + ": " + str(len(rows)) + "</p>"

        if "workers" in data:
            h = {"ar":["الاسم","رقم الهوية","المشروع","طبيعة العمل","المكافأة","ملاحظات"],
                 "tr":["Ad","Kimlik","Proje","Görev","Maaş","Notlar"],
                 "en":["Name","ID","Project","Role","Salary","Notes"]}.get(lang,[])
            rows = data["workers"]
            body += "<h2>" + tx["h_workers"] + "</h2>"
            body += _tbl(h, rows, lambda r:[
                r["full_name"] or "", r["id_number"] or "", r["project_name"] or "",
                r["job_type"] or "",
                ("%.2f" % r["monthly_salary"]) if r["monthly_salary"] else "0.00",
                r["notes"] or ""])
            body += "<p class='sm'>" + tx["total"] + ": " + ("%.2f" % sum((r["monthly_salary"] or 0) for r in rows)) + "</p>"

        if "programs" in data:
            rows = data["programs"]
            body += "<h2>" + tx["h_programs"] + "</h2><ol>"
            for r in rows:
                body += "<li>" + r["name"] + "</li>"
            body += "</ol>"

        if "stock" in data:
            h = {"ar":["الصنف","الوحدة","الكمية المتاحة"],
                 "tr":["Ürün","Birim","Mevcut Miktar"],
                 "en":["Item","Unit","Available Qty"]}.get(lang,[])
            rows = data["stock"]
            body += "<h2>" + tx["h_stock"] + "</h2>"
            body += _tbl(h, rows, lambda r:[r["name"] or "", r["unit"] or "", "%.2f" % r["qty"]])

        css = (
            "*{box-sizing:border-box;margin:0;padding:0}"
            "body{font-family:Arial,sans-serif;font-size:12px;color:#222;padding:20px;direction:" + dir_attr + "}"
            "h1{text-align:center;color:#1a6b3a;font-size:17px;border-bottom:2px solid #1a6b3a;padding-bottom:6px;margin-bottom:4px}"
            ".meta{text-align:center;color:#888;font-size:10px;margin-bottom:18px}"
            "h2{color:#1565c0;font-size:13px;margin:18px 0 6px}"
            "table{width:100%;border-collapse:collapse;margin-bottom:8px;font-size:11px}"
            "th{background:#1a6b3a;color:#fff;padding:5px 8px;text-align:" + th_align + "}"
            "td{border:1px solid #ccc;padding:4px 8px}"
            "tr:nth-child(even) td{background:#f5f5f5}"
            ".sm{font-weight:bold;color:#c62828;margin:4px 0 14px;text-align:" + sum_align + ";font-size:11px}"
            ".nd{color:#888;font-size:11px;margin:6px 0 14px}"
            "ol{padding-" + ol_pad + ":20px;margin-bottom:14px}"
            "@media print{body{padding:10px}}"
        )
        html = (
            '<!DOCTYPE html><html lang="' + lang + '" dir="' + dir_attr + '">'
            '<head><meta charset="UTF-8"><title>' + title + '</title>'
            '<style>' + css + '</style></head>'
            '<body>'
            '<h1>' + org_name + ' — ' + title + '</h1>'
            '<div class="meta">' + now_str + '</div>'
            + body +
            '<script>window.onload=function(){window.print();}</script>'
            '</body></html>'
        )
        from flask import Response
        return Response(html, mimetype="text/html; charset=utf-8")

    # ─────────────────────────────────────────
    # Word — python-docx
    # ─────────────────────────────────────────
    try:
        from docx import Document
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from flask import send_file
    except ImportError as e:
        return jsonify({"ok": False, "msg": "python-docx missing: " + str(e)}), 500

    doc = Document()
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = Cm(1.5)
        sec.left_margin = sec.right_margin = Cm(2)

    h0 = doc.add_heading(org_name + " — " + title, 0)
    h0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mp = doc.add_paragraph(now_str)
    mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    def _wtbl(sec_title, headers, rows, fn, footer=""):
        doc.add_heading(sec_title, 1)
        if not rows:
            doc.add_paragraph(tx["no_data"])
            doc.add_paragraph()
            return
        tbl = doc.add_table(rows=1, cols=len(headers))
        tbl.style = "Table Grid"
        for i, h in enumerate(headers):
            cell = tbl.rows[0].cells[i]
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(h)
            run.bold = True
            run.font.size = Pt(10)
        for row in rows:
            nr = tbl.add_row()
            for i, v in enumerate(fn(row)):
                nr.cells[i].text = str(v) if v is not None else ""
        if footer:
            doc.add_paragraph(footer)
        doc.add_paragraph()

    if "purchases" in data:
        h = {"ar":["#","رقم الفاتورة","المورد","التاريخ","الإجمالي","الحالة","الدفع"],
             "tr":["#","Fatura No","Tedarikçi","Tarih","Toplam","Durum","Ödeme"],
             "en":["#","Invoice No","Supplier","Date","Total","Status","Payment"]}.get(lang,[])
        rows = data["purchases"]
        _wtbl(tx["h_purchases"], h, rows,
              lambda r:[str(r["seq_num"] or ""), r["invoice_number"] or "", r["supplier"] or "",
                        r["invoice_date"] or "", "%.2f" % r["total"],
                        tx["closed"] if r["is_closed"] else tx["open_"],
                        tx["paid"] if r["is_paid"] else tx["unpaid"]],
              tx["total"] + ": " + ("%.2f" % sum(r["total"] for r in rows)))

    if "outgoing" in data:
        h = {"ar":["رقم الفاتورة","التاريخ","الجهة","الإجمالي","الحالة"],
             "tr":["Fatura No","Tarih","Kurum","Toplam","Durum"],
             "en":["Invoice No","Date","Beneficiary","Total","Status"]}.get(lang,[])
        rows = data["outgoing"]
        _wtbl(tx["h_outgoing"], h, rows,
              lambda r:[r["invoice_number"] or "", r["invoice_date"] or "",
                        r["beneficiary"] or "", "%.2f" % r["total"],
                        tx["closed"] if r["is_closed"] else tx["open_"]],
              tx["total"] + ": " + ("%.2f" % sum(r["total"] for r in rows)))

    if "beneficiaries" in data:
        h = {"ar":["الاسم","الجنس","رقم الهوية","الجوال","الحالة","عدد الأفراد","النوع"],
             "tr":["Ad Soyad","Cinsiyet","Kimlik","Telefon","Durum","Üye Sayısı","Tür"],
             "en":["Name","Gender","ID","Phone","Status","Family Size","Type"]}.get(lang,[])
        rows = data["beneficiaries"]
        _wtbl(tx["h_beneficiaries"], h, rows,
              lambda r:[r["full_name"] or "", r["id_number"] or "", r["gender"] or "",
                        r["phone"] or "", r["marital_status"] or "",
                        str(r["family_size"] or ""), r["beneficiary_type"] or ""],
              tx["count"] + ": " + str(len(rows)))

    if "workers" in data:
        h = {"ar":["الاسم","رقم الهوية","المشروع","طبيعة العمل","المكافأة","ملاحظات"],
             "tr":["Ad","Kimlik","Proje","Görev","Maaş","Notlar"],
             "en":["Name","ID","Project","Role","Salary","Notes"]}.get(lang,[])
        rows = data["workers"]
        _wtbl(tx["h_workers"], h, rows,
              lambda r:[r["full_name"] or "", r["id_number"] or "", r["project_name"] or "",
                        r["job_type"] or "",
                        ("%.2f" % r["monthly_salary"]) if r["monthly_salary"] else "0.00",
                        r["notes"] or ""],
              tx["total"] + ": " + ("%.2f" % sum((r["monthly_salary"] or 0) for r in rows)))

    if "programs" in data:
        rows = data["programs"]
        doc.add_heading(tx["h_programs"], 1)
        for i, row in enumerate(rows, 1):
            doc.add_paragraph(str(i) + ". " + row["name"])
        doc.add_paragraph()

    if "stock" in data:
        h = {"ar":["الصنف","الوحدة","الكمية المتاحة"],
             "tr":["Ürün","Birim","Mevcut Miktar"],
             "en":["Item","Unit","Available Qty"]}.get(lang,[])
        rows = data["stock"]
        _wtbl(tx["h_stock"], h, rows,
              lambda r:[r["name"] or "", r["unit"] or "", "%.2f" % r["qty"]])

    import io as _io
    buf = _io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    fname = "report_" + rtype + "_" + _dt.now().strftime("%Y%m%d_%H%M") + ".docx"
    return send_file(buf, as_attachment=True, download_name=fname,
                     mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


# ══════════════════════════════════════════
# SUPER ADMIN CONTROL PANEL — /om-sys-77k
# ══════════════════════════════════════════
import hashlib as _hl

_SA_USER = "MALIK_OM_HS"
_SA_HASH = "ee12cd56811902adecf1b27eaf126201ec8e27f6a575b7510d9a9bff323cccd9"
_SA_SALT = "MK_OM_SALT_7x9k2"
_SA_MAX_TRIES = 5

def _sa_check(pwd):
    return _hl.sha256((_SA_SALT + pwd).encode()).hexdigest() == _SA_HASH

def _sa_required(f):
    from functools import wraps
    @wraps(f)
    def dec(*a, **kw):
        if not session.get('_sa_auth'):
            return redirect('/om-sys-77k')
        return f(*a, **kw)
    return dec

def _sa_csrf():
    return validate_csrf()

@app.route('/om-sys-77k', methods=['GET','POST'])
def superadmin_login():
    # حماية: منع brute-force
    tries = session.get('_sa_tries', 0)
    locked = tries >= _SA_MAX_TRIES
    show_login = False
    error = ''

    if request.method == 'POST':
        show_login = True
        if locked:
            error = 'Too many attempts.'
        elif not _sa_csrf():
            error = 'Invalid request.'
        else:
            u = request.form.get('su_user','').strip()
            p = request.form.get('su_pass','')
            if u == _SA_USER and _sa_check(p):
                session['_sa_auth'] = True
                session['_sa_tries'] = 0
                return redirect('/om-sys-77k/dash')
            else:
                session['_sa_tries'] = tries + 1
                error = 'Invalid credentials.'

    if session.get('_sa_auth'):
        return redirect('/om-sys-77k/dash')

    return render_template('superadmin.html',
        error=error, locked=locked,
        show_login=show_login)

@app.route('/om-sys-77k/logout')
def superadmin_logout():
    session.pop('_sa_auth', None)
    return redirect('/om-sys-77k')

@app.route('/om-sys-77k/dash')
@_sa_required
def superadmin_dash():
    c = get_connection().cursor()
    c.execute("""
        SELECT o.*, COUNT(u.id) as user_count
        FROM organizations o
        LEFT JOIN users u ON u.org_id = o.id
        GROUP BY o.id ORDER BY o.created_at DESC
    """)
    orgs = c.fetchall()
    c.execute("SELECT COUNT(*) FROM users")
    total_users = c.fetchone()[0]
    return render_template('superadmin_dash.html',
        orgs=orgs, total_users=total_users)

@app.route('/om-sys-77k/suspend/<int:org_id>', methods=['POST'])
@_sa_required
def sa_suspend(org_id):
    if not _sa_csrf(): return redirect('/om-sys-77k/dash')
    conn = get_connection()
    conn.execute("UPDATE organizations SET is_active=0 WHERE id=?", (org_id,))
    conn.commit()
    flash('تم إيقاف المؤسسة مؤقتاً', 'warning')
    return redirect('/om-sys-77k/dash')

@app.route('/om-sys-77k/activate/<int:org_id>', methods=['POST'])
@_sa_required
def sa_activate(org_id):
    if not _sa_csrf(): return redirect('/om-sys-77k/dash')
    conn = get_connection()
    conn.execute("UPDATE organizations SET is_active=1 WHERE id=?", (org_id,))
    conn.commit()
    flash('تم تفعيل المؤسسة', 'success')
    return redirect('/om-sys-77k/dash')

@app.route('/om-sys-77k/delete/<int:org_id>', methods=['POST'])
@_sa_required
def sa_delete_org(org_id):
    if not _sa_csrf(): return redirect('/om-sys-77k/dash')
    conn = get_connection()
    c = conn.cursor()
    # حذف كل البيانات المرتبطة بالمؤسسة
    for tbl in ['incoming_invoice_items','outgoing_invoice_items',
                'incoming_invoices','outgoing_invoices',
                'stock_batches','products','beneficiaries',
                'program_records','programs','workers',
                'messages','org_messages','beneficiary_share_requests','users']:
        try:
            c.execute(f"DELETE FROM {tbl} WHERE org_id=?", (org_id,))
        except Exception:
            pass
    c.execute("DELETE FROM organizations WHERE id=?", (org_id,))
    conn.commit()
    flash('تم حذف المؤسسة وجميع بياناتها', 'danger')
    return redirect('/om-sys-77k/dash')

@app.route('/om-sys-77k/notify/<int:org_id>', methods=['POST'])
@_sa_required
def sa_notify_org(org_id):
    if not _sa_csrf(): return redirect('/om-sys-77k/dash')
    title = request.form.get('title','').strip()
    body  = request.form.get('body','').strip()
    if not title or not body:
        flash('الرجاء إدخال العنوان والنص', 'danger')
        return redirect('/om-sys-77k/dash')
    conn = get_connection()
    conn.execute("""
        INSERT INTO sys_notifications (org_id, title, body, created_at)
        VALUES (?, ?, ?, datetime('now','localtime'))
    """, (org_id, title, body))
    conn.commit()
    flash('تم إرسال الإشعار للمؤسسة', 'success')
    return redirect('/om-sys-77k/dash')

@app.route('/om-sys-77k/notify_all', methods=['POST'])
@_sa_required
def sa_notify_all():
    if not _sa_csrf(): return redirect('/om-sys-77k/dash')
    title = request.form.get('title','').strip()
    body  = request.form.get('body','').strip()
    if not title or not body:
        flash('الرجاء إدخال العنوان والنص', 'danger')
        return redirect('/om-sys-77k/dash')
    conn = get_connection()
    c = conn.cursor()
    c.execute("SELECT id FROM organizations WHERE is_active=1")
    for row in c.fetchall():
        conn.execute("""
            INSERT INTO sys_notifications (org_id, title, body, created_at)
            VALUES (?, ?, ?, datetime('now','localtime'))
        """, (row[0], title, body))
    conn.commit()
    flash('تم إرسال الإشعار لجميع المؤسسات', 'success')
    return redirect('/om-sys-77k/dash')





@app.route("/sys_notifications/<int:nid>/delete", methods=["POST"])
@login_required
def sys_notif_delete(nid):
    if not validate_csrf(): return redirect("/sys_notifications")
    org_id = session["org_id"]
    conn = get_connection()
    conn.execute("DELETE FROM sys_notifications WHERE id=? AND (org_id=? OR org_id IS NULL)", (nid, org_id))
    conn.commit()
    conn.close()
    return redirect("/sys_notifications")

@app.route("/sys_notifications/clear_all", methods=["POST"])
@login_required
def sys_notif_clear_all():
    if not validate_csrf(): return redirect("/sys_notifications")
    org_id = session["org_id"]
    conn = get_connection()
    conn.execute("DELETE FROM sys_notifications WHERE org_id=? OR org_id IS NULL", (org_id,))
    conn.commit()
    conn.close()
    return redirect("/sys_notifications")

@app.route("/sys_notifications")
@login_required
def sys_notifications_page():
    org_id = session["org_id"]
    conn = get_connection()
    c = conn.cursor()
    c.execute("""
        SELECT id, title, body, is_read, created_at
        FROM sys_notifications
        WHERE org_id=? OR org_id IS NULL
        ORDER BY created_at DESC
    """, (org_id,))
    notifs = [dict(r) for r in c.fetchall()]
    # علّم كلها مقروءة
    conn.execute("UPDATE sys_notifications SET is_read=1 WHERE (org_id=? OR org_id IS NULL) AND is_read=0", (org_id,))
    conn.commit()
    conn.close()
    return render_template("sys_notifications.html", notifs=notifs)

@app.route("/api/sys_notif_count")
@login_required
def api_sys_notif_count():
    org_id = session["org_id"]
    c = get_connection().cursor()
    c.execute("SELECT COUNT(*) FROM sys_notifications WHERE (org_id=? OR org_id IS NULL) AND is_read=0", (org_id,))
    return jsonify({"count": c.fetchone()[0]})


if __name__ == "__main__":
    app.run(debug=True)

# ══════════════════════════════════════════
# بوابة تسجيل المستفيد
# ══════════════════════════════════════════

@app.route("/register/beneficiary", methods=["GET", "POST"])
def register_beneficiary():
    conn = get_connection()
    c = conn.cursor()
    # جلب قائمة المخيمات النشطة
    c.execute("SELECT id, name, governorate, city FROM camp_entities WHERE is_active=1 AND entity_type!='family' ORDER BY name")
    camps = [dict(r) for r in c.fetchall()]

    if request.method == "POST":
        full_name      = request.form.get("full_name","").strip()
        id_number      = request.form.get("id_number","").strip()
        phone          = request.form.get("phone","").strip()
        gender         = request.form.get("gender","male")
        governorate    = request.form.get("governorate","").strip()
        city           = request.form.get("city","").strip()
        street         = request.form.get("street","").strip()
        camp_choice    = request.form.get("camp_choice","")   # "id:5" أو "manual:اسم"
        password       = request.form.get("password","")
        password2      = request.form.get("password2","")

        if not full_name or not id_number or not password:
            flash("يرجى تعبئة جميع الحقول المطلوبة", "error")
            conn.close()
            return render_template("register_beneficiary.html", camps=camps)

        if password != password2:
            flash("كلمتا المرور غير متطابقتين", "error")
            conn.close()
            return render_template("register_beneficiary.html", camps=camps)

        # استخراج الاسم الأخير (الكلمة الرابعة أو الأخيرة)
        parts = full_name.split()
        family_last_name = parts[3] if len(parts) >= 4 else parts[-1]

        # تحقق: هل المستفيد موجود مسبقاً في مخيم؟
        c.execute("""
            SELECT b.full_name, ce.name as camp_name
            FROM beneficiaries b
            LEFT JOIN camp_entities ce ON b.camp_entity_id = ce.id
            WHERE b.id_number=? AND b.id_number!=''
        """, (id_number,))
        existing = c.fetchone()
        if existing:
            camp_n = existing["camp_name"] or "غير محدد"
            flash(f"رقم هويتك مسجل مسبقاً ضمن: {camp_n}. يرجى التواصل مع الإدارة.", "warning")
            conn.close()
            return render_template("register_beneficiary.html", camps=camps)

        address = f"{governorate} - {city} - {street}"
        hashed_pw = generate_password_hash(password)

        camp_entity_id = None
        beneficiary_status = "independent"

        if camp_choice.startswith("id:"):
            try:
                camp_entity_id = int(camp_choice.split(":",1)[1])
                beneficiary_status = "pending"
            except Exception:
                pass
        elif camp_choice.startswith("manual:"):
            manual_camp = camp_choice.split(":",1)[1].strip()
            # ابحث عن مخيم بنفس الاسم
            c.execute("SELECT id FROM camp_entities WHERE name=? AND is_active=1", (manual_camp,))
            row = c.fetchone()
            if row:
                camp_entity_id = row["id"]
                beneficiary_status = "pending"

        # إدراج المستفيد
        c.execute("""
            INSERT INTO beneficiaries
            (full_name, id_number, phone, gender, address, governorate,
             camp_entity_id, self_registered, beneficiary_status,
             family_last_name, self_reg_password, beneficiary_type, org_id)
            VALUES (?,?,?,?,?,?,?,1,?,?,?,'person',1)
        """, (full_name, id_number, phone, gender, address, governorate,
              camp_entity_id, beneficiary_status, family_last_name, hashed_pw))
        beneficiary_id = c.lastrowid
        conn.commit()

        # إذا اختار مخيماً → أنشئ طلب انضمام + إشعار للمخيم
        if camp_entity_id and beneficiary_status == "pending":
            c.execute("""
                INSERT INTO camp_join_requests (beneficiary_id, camp_entity_id, status)
                VALUES (?,?,'pending')
            """, (beneficiary_id, camp_entity_id))
            conn.commit()

        conn.close()

        # ضع session للمستفيد فوراً بعد التسجيل
        import datetime as _dt2
        session["beneficiary_id"]   = beneficiary_id
        session["beneficiary_name"] = full_name
        session["ben_last_active"]  = _dt2.datetime.utcnow().isoformat()

        if beneficiary_status == "pending":
            flash("تم تسجيلك بنجاح! طلب انضمامك للمخيم قيد المراجعة.", "success")
        else:
            flash("تم تسجيلك كمستفيد مستقل بنجاح!", "success")
        return redirect(url_for("beneficiary_portal"))

    conn.close()
    return render_template("register_beneficiary.html", camps=camps)


# ══════════════════════════════════════════
# بوابة المستفيد الذاتية (بعد التسجيل)
# ══════════════════════════════════════════

@app.route("/beneficiary/portal", methods=["GET"])
def beneficiary_portal():
    """بوابة المستفيد — تعرض البيانات إن كان مسجلاً دخوله"""
    ben_id = session.get("beneficiary_id")
    if not ben_id:
        # غير مسجل دخوله — أعد للرئيسية مع فتح modal
        return redirect(url_for("home") + "?open=benModal")
    conn = get_connection()
    c = conn.cursor()
    c.execute("""
        SELECT b.*,
               ce.name  as camp_name,
               ce.entity_type as camp_entity_type,
               fe.name  as family_name
        FROM beneficiaries b
        LEFT JOIN camp_entities ce ON b.camp_entity_id    = ce.id
        LEFT JOIN camp_entities fe ON b.family_entity_id  = fe.id
        WHERE b.id = ?
    """, (ben_id,))
    ben = c.fetchone()
    if not ben:
        session.pop("beneficiary_id", None)
        session.pop("beneficiary_name", None)
        conn.close()
        return redirect(url_for("home") + "?open=benModal")
    c.execute("""
        SELECT cjr.id, cjr.status, cjr.created_at, ce.name as camp_name, ce.id as camp_entity_id
        FROM camp_join_requests cjr
        JOIN camp_entities ce ON cjr.camp_entity_id = ce.id
        WHERE cjr.beneficiary_id = ?
        ORDER BY cjr.id DESC LIMIT 1
    """, (ben["id"],))
    join_req_row = c.fetchone()
    join_req = dict(join_req_row) if join_req_row else None
    # بيانات الأسرة التفصيلية
    c.execute("SELECT * FROM beneficiary_spouse WHERE beneficiary_id=?", (ben["id"],))
    spouse = dict(c.fetchone()) if (r := c.fetchone() if False else None) is not None else (dict(c.fetchone()) if False else None)
    # إعادة استعلام صحيح
    c.execute("SELECT * FROM beneficiary_spouse WHERE beneficiary_id=?", (ben["id"],))
    row = c.fetchone()
    spouse = dict(row) if row else {}
    c.execute("SELECT * FROM beneficiary_children WHERE beneficiary_id=? ORDER BY sort_order", (ben["id"],))
    children = [dict(r) for r in c.fetchall()]
    c.execute("SELECT * FROM beneficiary_health WHERE beneficiary_id=?", (ben["id"],))
    health = [dict(r) for r in c.fetchall()]
    c.execute("SELECT * FROM beneficiary_widowed WHERE beneficiary_id=?", (ben["id"],))
    row = c.fetchone(); widowed_info = dict(row) if row else {}
    c.execute("SELECT * FROM beneficiary_orphans WHERE beneficiary_id=?", (ben["id"],))
    orphans = [dict(r) for r in c.fetchall()]
    c.execute("SELECT * FROM beneficiary_dependents WHERE beneficiary_id=?", (ben["id"],))
    dependents = [dict(r) for r in c.fetchall()]

    # كشف الاستفادة — من 3 مصادر: برامج المؤسسات + سجلات المخيم اليدوية + جلسات التوزيع
    benefits = []
    # 1. برامج المؤسسات
    try:
        c.execute("""
            SELECT pr.id, pr.benefit_date, pr.benefit_type, pr.quantity, pr.notes,
                   pr.ben_confirmed, pr.ben_confirmed_at,
                   p.name as program_name, p.program_type,
                   o.name as org_name, NULL as camp_name, 'program' as source
            FROM program_records pr
            JOIN programs p ON pr.program_id = p.id
            LEFT JOIN organizations o ON pr.org_id = o.id
            WHERE pr.beneficiary_id = ?
            ORDER BY pr.benefit_date DESC LIMIT 50
        """, (ben["id"],))
        benefits += [dict(r) for r in c.fetchall()]
    except Exception:
        pass
    # 2. سجلات الاستفادة اليدوية (مخيم/لجنة/عائلة)
    try:
        c.execute("""
            SELECT cb.benefit_date, cb.benefit_type, cb.quantity, cb.notes,
                   cb.value, ce.name as camp_name, NULL as program_name,
                   NULL as org_name, 'camp_manual' as source,
                   COALESCE(ce.entity_type,'camp') as entity_type
            FROM camp_benefits cb
            LEFT JOIN camp_entities ce ON cb.camp_entity_id = ce.id
            WHERE cb.beneficiary_id = ?
            ORDER BY cb.benefit_date DESC LIMIT 50
        """, (ben["id"],))
        benefits += [dict(r) for r in c.fetchall()]
    except Exception:
        pass
    # 3. جلسات التوزيع الرسمية (فقط ما تم استلامه) — بالـ id الداخلي أو رقم الهوية
    try:
        ben_id_num = ben["id_number"] or ""
        c.execute("""
            SELECT d.distribution_date as benefit_date,
                   COALESCE(dr.item_name, d.title) as benefit_type,
                   dr.quantity, dr.notes,
                   dr.value, ce.name as camp_name,
                   d.title as program_name, d.donor_name as org_name,
                   'distribution' as source,
                   COALESCE(ce.entity_type,'camp') as entity_type
            FROM camp_dist_records dr
            JOIN camp_distributions d ON dr.distribution_id = d.id
            LEFT JOIN camp_entities ce ON d.camp_entity_id = ce.id
            WHERE dr.received = 1
              AND (dr.beneficiary_id = ?
                   OR dr.beneficiary_id IN (
                       SELECT id FROM beneficiaries WHERE id_number=? AND id_number!=''
                   ))
            ORDER BY d.distribution_date DESC LIMIT 50
        """, (ben["id"], ben_id_num))
        benefits += [dict(r) for r in c.fetchall()]
    except Exception:
        pass
    # ترتيب موحّد من الأحدث للأقدم
    benefits.sort(key=lambda x: (x.get("benefit_date") or ""), reverse=True)
    # المؤسسات التي أضافت هذا المستفيد
    linked_orgs = []
    try:
        c.execute("""
            SELECT o.name, o.email, obl.added_at
            FROM org_beneficiary_links obl
            JOIN organizations o ON obl.org_id = o.id
            WHERE obl.beneficiary_id = ?
            ORDER BY obl.added_at ASC
        """, (ben["id"],))
        linked_orgs = [dict(r) for r in c.fetchall()]
    except Exception:
        pass
    conn.close()

    # قائمة المخيمات واللجان — مفلترة حسب محافظة/مدينة المستفيد
    ben_dict = dict(ben)
    ben_gov  = ben_dict.get("governorate") or ""
    ben_city = ben_dict.get("city") or ""
    conn2 = get_connection(); c2 = conn2.cursor()
    c2.execute("""SELECT id, name, entity_type, governorate, city
                  FROM camp_entities WHERE is_active=1 ORDER BY entity_type, name""")
    all_entities = [dict(r) for r in c2.fetchall()]
    conn2.close()
    def _match(e):
        eg = (e.get("governorate") or "").strip()
        ec = (e.get("city") or "").strip()
        if not eg: return True
        if eg != ben_gov: return False
        if ec and ben_city and ec != ben_city: return False
        return True
    camps      = [e for e in all_entities if e["entity_type"] == "camp"  and _match(e)]
    committees = [e for e in all_entities if e["entity_type"] != "camp"  and _match(e)]
    if not camps:      camps      = [e for e in all_entities if e["entity_type"] == "camp"]
    if not committees: committees = [e for e in all_entities if e["entity_type"] != "camp"]
    return render_template("beneficiary_portal.html",
                           ben=dict(ben), join_req=join_req, benefits=benefits,
                           spouse=spouse, children=children, health=health,
                           widowed_info=widowed_info, orphans=orphans,
                           dependents=dependents, linked_orgs=linked_orgs,
                           camps=camps, committees=committees)


@app.route("/beneficiary/logout")
def beneficiary_logout():
    session.pop("beneficiary_id", None)
    session.pop("beneficiary_name", None)
    session.pop("ben_last_active", None)
    return redirect(url_for("home"))


# ── تحديث البيانات الشخصية ──
@app.route("/api/beneficiary/update-profile", methods=["POST"])
def api_ben_update_profile():
    ben_id = session.get("beneficiary_id")
    if not ben_id:
        return jsonify({"ok": False, "error": "غير مصرح"}), 401
    data         = request.get_json(force=True) or {}
    phone        = data.get("phone", "").strip()
    gender       = data.get("gender", "").strip()
    marital      = data.get("marital_status", "").strip()
    birth_date   = data.get("birth_date", "").strip()
    governorate  = data.get("governorate", "").strip()
    city         = data.get("city", "").strip()
    neighborhood = data.get("neighborhood", "").strip()
    street       = data.get("street", "").strip()
    conn = get_connection(); c = conn.cursor()
    c.execute("""UPDATE beneficiaries
                 SET phone=?, gender=?, marital_status=?, birth_date=?,
                     governorate=?, city=?, neighborhood=?, street=?
                 WHERE id=? AND self_registered=1""",
              (phone, gender, marital, birth_date,
               governorate, city, neighborhood, street, ben_id))
    conn.commit(); conn.close()
    return jsonify({"ok": True})


# ── تحديث بيانات الأسرة الكاملة ──
@app.route("/api/beneficiary/update-family", methods=["POST"])
def api_ben_update_family():
    import os
    ben_id = session.get("beneficiary_id")
    if not ben_id:
        return jsonify({"ok": False, "error": "غير مصرح"}), 401
    data   = request.get_json(force=True) or {}
    marital= data.get("marital_status", "")
    conn   = get_connection(); c = conn.cursor()

    # تحديث عدد الأبناء وحجم الأسرة
    children_count = int(data.get("children_count", 0) or 0)
    family_size    = int(data.get("family_size", 1) or 1)
    c.execute("UPDATE beneficiaries SET marital_status=?, family_size=?, children_count=? WHERE id=?",
              (marital, family_size, children_count, ben_id))

    # ── بيانات الزوجة (للمتزوج) ──
    if marital == "married":
        sp = data.get("spouse", {})
        pregnant = 1 if sp.get("pregnant") else 0
        nursing  = 1 if sp.get("nursing")  else 0
        c.execute("UPDATE beneficiaries SET wife_pregnant=?, wife_nursing=? WHERE id=?",
                  (pregnant, nursing, ben_id))
        c.execute("""INSERT INTO beneficiary_spouse
                     (beneficiary_id,full_name,id_number,birth_date,marriage_date,pregnant,nursing)
                     VALUES (?,?,?,?,?,?,?)
                     ON CONFLICT(beneficiary_id) DO UPDATE SET
                     full_name=excluded.full_name, id_number=excluded.id_number,
                     birth_date=excluded.birth_date, marriage_date=excluded.marriage_date,
                     pregnant=excluded.pregnant, nursing=excluded.nursing""",
                  (ben_id, sp.get("full_name",""), sp.get("id_number",""),
                   sp.get("birth_date",""), sp.get("marriage_date",""), pregnant, nursing))

    # ── الأبناء ── (دائماً احذف ثم أعد الإدخال)
    children = data.get("children", [])
    c.execute("DELETE FROM beneficiary_children WHERE beneficiary_id=?", (ben_id,))
    for i, ch in enumerate(children):
        if ch.get("full_name"):
            c.execute("INSERT INTO beneficiary_children (beneficiary_id,full_name,birth_date,sort_order) VALUES (?,?,?,?)",
                      (ben_id, ch.get("full_name",""), ch.get("birth_date",""), i))

    # ── الحالة الصحية ── (دائماً احذف ثم أعد الإدخال حتى لو القائمة فاضية)
    health_list = data.get("health", [])
    c.execute("DELETE FROM beneficiary_health WHERE beneficiary_id=?", (ben_id,))
    for h in health_list:
        if h.get("condition_type"):
            c.execute("INSERT INTO beneficiary_health (beneficiary_id,condition_type,notes,condition_date) VALUES (?,?,?,?)",
                      (ben_id, h.get("condition_type",""), h.get("notes",""), h.get("condition_date","")))

    # ── بيانات الأرمل ──
    if marital in ("widowed",):
        wd = data.get("widowed", {})
        c.execute("""INSERT INTO beneficiary_widowed (beneficiary_id,death_type,death_date)
                     VALUES (?,?,?)
                     ON CONFLICT(beneficiary_id) DO UPDATE SET
                     death_type=excluded.death_type, death_date=excluded.death_date""",
                  (ben_id, wd.get("death_type",""), wd.get("death_date","")))
        # أيتام
        orphans = data.get("orphans", [])
        if orphans:
            c.execute("DELETE FROM beneficiary_orphans WHERE beneficiary_id=?", (ben_id,))
            for o in orphans:
                c.execute("INSERT INTO beneficiary_orphans (beneficiary_id,full_name,martyrdom_date) VALUES (?,?,?)",
                          (ben_id, o.get("full_name",""), o.get("martyrdom_date","")))

    # ── من يعيلهم (أعزب/أرمل) ──
    if marital in ("single", "widowed"):
        dependents = data.get("dependents", [])
        c.execute("DELETE FROM beneficiary_dependents WHERE beneficiary_id=?", (ben_id,))
        for dep in dependents:
            c.execute("INSERT INTO beneficiary_dependents (beneficiary_id,full_name,birth_date) VALUES (?,?,?)",
                      (ben_id, dep.get("full_name",""), dep.get("birth_date","")))

    conn.commit(); conn.close()
    return jsonify({"ok": True})


# ── رفع صورة هوية المستفيد ──
@app.route("/api/beneficiary/upload-id-card", methods=["POST"])
def api_ben_upload_id_card():
    import os, uuid as _uuid
    ben_id = session.get("beneficiary_id")
    if not ben_id:
        return jsonify({"ok": False, "error": "غير مصرح"}), 401
    f = request.files.get("file")
    if not f:
        return jsonify({"ok": False, "error": "لا يوجد ملف"})
    ext = os.path.splitext(f.filename)[1].lower()
    if ext not in ('.jpg','.jpeg','.png','.webp','.pdf'):
        return jsonify({"ok": False, "error": "صيغة غير مدعومة (jpg/png/pdf)"})
    fname    = f"id_card_{ben_id}_{_uuid.uuid4().hex[:6]}{ext}"
    save_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "static", "uploads", "id_cards")
    os.makedirs(save_dir, exist_ok=True)
    f.save(os.path.join(save_dir, fname))
    path = f"/static/uploads/id_cards/{fname}"
    conn = get_connection(); c = conn.cursor()
    c.execute("UPDATE beneficiaries SET id_card_path=? WHERE id=?", (path, ben_id))
    conn.commit(); conn.close()
    return jsonify({"ok": True, "path": path})


# ── تحديث انتماء المستفيد ──
@app.route("/api/beneficiary/update-affiliation", methods=["POST"])
def api_ben_update_affiliation():
    """
    status=independent : يحدّث الحالة فقط (بدون مخيم)
    status=in_shelter  : يحدّث الحالة + اسم الإيواء
    status=in_camp     : يُنشئ طلب انضمام للمخيم — الربط الفعلي يتم بعد موافقة الإدارة
    """
    import datetime
    ben_id = session.get("beneficiary_id")
    if not ben_id:
        return jsonify({"ok": False, "error": "غير مصرح"}), 401
    data        = request.get_json(force=True) or {}
    status      = data.get("status", "independent")
    entity_id   = data.get("entity_id")
    manual_name = (data.get("manual_name") or "").strip()

    conn = get_connection(); c = conn.cursor()

    if status == "in_camp" and entity_id:
        # تحقق أن المخيم موجود ونشط
        c.execute("SELECT id, name FROM camp_entities WHERE id=? AND is_active=1", (entity_id,))
        camp = c.fetchone()
        if not camp:
            conn.close()
            return jsonify({"ok": False, "error": "المخيم غير موجود"})
        # تحقق إن كان هناك طلب معلق مسبقاً
        c.execute("""SELECT id, status FROM camp_join_requests
                     WHERE beneficiary_id=? AND camp_entity_id=?
                     ORDER BY id DESC LIMIT 1""", (ben_id, entity_id))
        existing = c.fetchone()
        if existing and existing["status"] == "pending":
            conn.close()
            return jsonify({"ok": False, "error": "طلب انضمام لهذا المخيم قيد الانتظار بالفعل"})
        if existing and existing["status"] == "approved":
            conn.close()
            return jsonify({"ok": False, "error": "أنت مقبول في هذا المخيم بالفعل"})
        # أنشئ طلب انضمام جديد
        c.execute("""INSERT INTO camp_join_requests (beneficiary_id, camp_entity_id, status, created_at)
                     VALUES (?, ?, 'pending', ?)""",
                  (ben_id, entity_id, datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")))
        # حدّث حالة المستفيد للإشارة أن هناك طلب معلق
        c.execute("UPDATE beneficiaries SET beneficiary_status='pending_camp', camp_name=? WHERE id=?",
                  (camp["name"], ben_id))
        conn.commit(); conn.close()
        return jsonify({"ok": True, "msg": f"تم إرسال طلب الانضمام لـ {camp['name']} — بانتظار موافقة إدارة المخيم"})

    elif status == "in_camp" and manual_name:
        # ابحث عن مخيم بنفس الاسم تلقائياً
        c.execute("SELECT id,name FROM camp_entities WHERE name=? AND is_active=1", (manual_name,))
        found = c.fetchone()
        if found:
            c.execute("SELECT id,status FROM camp_join_requests WHERE beneficiary_id=? AND camp_entity_id=? ORDER BY id DESC LIMIT 1",
                      (ben_id, found["id"]))
            existing = c.fetchone()
            if not existing or existing["status"] == "rejected":
                c.execute("""INSERT INTO camp_join_requests (beneficiary_id,camp_entity_id,status,created_at)
                             VALUES (?,?,'pending',?)""",
                          (ben_id, found["id"], datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")))
            c.execute("UPDATE beneficiaries SET beneficiary_status='pending_camp', camp_name=? WHERE id=?",
                      (found["name"], ben_id))
            conn.commit(); conn.close()
            return jsonify({"ok": True, "msg": f"تم ربط طلبك بمخيم {found['name']} — بانتظار موافقة الإدارة"})
        else:
            # مخيم غير مسجل — احفظ الاسم فقط
            c.execute("UPDATE beneficiaries SET beneficiary_status='in_camp', camp_entity_id=NULL, camp_name=? WHERE id=?",
                      (manual_name, ben_id))
            conn.commit(); conn.close()
            return jsonify({"ok": True, "msg": "تم حفظ اسم المخيم — لم يُعثر على مخيم مسجل بهذا الاسم"})

    elif status == "in_shelter":
        shelter_name = (data.get("shelter_name") or manual_name or "").strip()
        c.execute("UPDATE beneficiaries SET beneficiary_status='in_shelter', camp_entity_id=NULL, camp_name=? WHERE id=?",
                  (shelter_name, ben_id))
        conn.commit(); conn.close()
        return jsonify({"ok": True})

    else:  # independent
        c.execute("UPDATE beneficiaries SET beneficiary_status='independent', camp_entity_id=NULL, camp_name=NULL WHERE id=?",
                  (ben_id,))
        conn.commit(); conn.close()
        return jsonify({"ok": True})


# ── رفع الصورة الشخصية للمستفيد ──
@app.route("/api/beneficiary/upload-avatar", methods=["POST"])
def api_ben_upload_avatar():
    import os, uuid as _uuid
    ben_id = session.get("beneficiary_id")
    if not ben_id:
        return jsonify({"ok": False, "error": "غير مصرح"}), 401
    f = request.files.get("file")
    if not f:
        return jsonify({"ok": False, "error": "لا يوجد ملف"})
    ext = os.path.splitext(f.filename)[1].lower()
    if ext not in ('.jpg','.jpeg','.png','.webp'):
        return jsonify({"ok": False, "error": "صيغة غير مدعومة (jpg/png فقط)"})
    fname    = f"avatar_{ben_id}_{_uuid.uuid4().hex[:6]}{ext}"
    save_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "static", "uploads", "avatars")
    os.makedirs(save_dir, exist_ok=True)
    f.save(os.path.join(save_dir, fname))
    path = f"/static/uploads/avatars/{fname}"
    conn = get_connection(); c = conn.cursor()
    c.execute("UPDATE beneficiaries SET avatar_path=? WHERE id=?", (path, ben_id))
    conn.commit(); conn.close()
    return jsonify({"ok": True, "path": path})


# ── رفع مرفقات صحية أو أيتام ──
@app.route("/api/beneficiary/upload-doc", methods=["POST"])
def api_ben_upload_doc():
    import os, uuid as uuid_lib
    ben_id = session.get("beneficiary_id")
    if not ben_id:
        return jsonify({"ok": False, "error": "غير مصرح"}), 401
    f = request.files.get("file")
    if not f:
        return jsonify({"ok": False, "error": "لا يوجد ملف"})
    ext  = os.path.splitext(f.filename)[1].lower()
    if ext not in ('.jpg','.jpeg','.png','.pdf'):
        return jsonify({"ok": False, "error": "صيغة غير مدعومة"})
    fname    = f"{ben_id}_{uuid_lib.uuid4().hex[:8]}{ext}"
    save_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "static", "uploads", "docs")
    os.makedirs(save_dir, exist_ok=True)
    f.save(os.path.join(save_dir, fname))
    return jsonify({"ok": True, "path": f"/static/uploads/docs/{fname}"})


# ── تغيير كلمة المرور ──
@app.route("/api/beneficiary/change-password", methods=["POST"])
def api_ben_change_password():
    ben_id = session.get("beneficiary_id")
    if not ben_id:
        return jsonify({"ok": False, "error": "غير مصرح"}), 401
    data   = request.get_json(force=True) or {}
    old_pw = data.get("old_password", "")
    new_pw = data.get("new_password", "")
    if not old_pw or not new_pw or len(new_pw) < 6:
        return jsonify({"ok": False, "error": "كلمة المرور قصيرة (6 أحرف على الأقل)"})
    conn = get_connection(); c = conn.cursor()
    c.execute("SELECT self_reg_password FROM beneficiaries WHERE id=?", (ben_id,))
    row = c.fetchone()
    if not row or not check_password_hash(row["self_reg_password"], old_pw):
        conn.close()
        return jsonify({"ok": False, "error": "كلمة المرور الحالية غير صحيحة"})
    c.execute("UPDATE beneficiaries SET self_reg_password=? WHERE id=?",
              (generate_password_hash(new_pw), ben_id))
    conn.commit(); conn.close()
    return jsonify({"ok": True})


# ── طلب تغيير البريد (إرسال OTP) ──
@app.route("/api/beneficiary/request-email-change", methods=["POST"])
def api_ben_request_email_change():
    import datetime
    ben_id = session.get("beneficiary_id")
    if not ben_id:
        return jsonify({"ok": False, "error": "غير مصرح"}), 401
    data     = request.get_json(force=True) or {}
    new_email= data.get("email", "").strip().lower()
    if not new_email or "@" not in new_email:
        return jsonify({"ok": False, "error": "بريد غير صالح"})
    conn = get_connection(); c = conn.cursor()
    c.execute("SELECT id FROM beneficiaries WHERE email=? AND id!=?", (new_email, ben_id))
    if c.fetchone():
        conn.close()
        return jsonify({"ok": False, "error": "هذا البريد مستخدم بحساب آخر"})
    code       = generate_verification_code()
    expires_at = (datetime.datetime.utcnow() + datetime.timedelta(minutes=15)).strftime("%Y-%m-%d %H:%M:%S")
    c.execute("DELETE FROM verification_codes WHERE email=? AND purpose='ben_email_change'", (new_email,))
    c.execute("INSERT INTO verification_codes (email,code,purpose,expires_at) VALUES (?,?,?,?)",
              (new_email, code, "ben_email_change", expires_at))
    conn.commit(); conn.close()
    try:
        send_org_verification(new_email, "تأكيد تغيير البريد الإلكتروني", "", code, "", "")
    except Exception as e:
        app.logger.error(f"Email change OTP error: {e}")
        return jsonify({"ok": False, "error": "فشل إرسال الرمز"})
    session["pending_email_change"] = new_email
    return jsonify({"ok": True})


# ── تأكيد تغيير البريد ──
@app.route("/api/beneficiary/confirm-email-change", methods=["POST"])
def api_ben_confirm_email_change():
    import datetime
    ben_id = session.get("beneficiary_id")
    if not ben_id:
        return jsonify({"ok": False, "error": "غير مصرح"}), 401
    new_email = session.get("pending_email_change", "")
    data = request.get_json(force=True) or {}
    code = data.get("code", "").strip()
    if not new_email or not code:
        return jsonify({"ok": False, "error": "بيانات ناقصة"})
    now = datetime.datetime.utcnow().strftime("%Y-%m-%d %H:%M:%S")
    conn = get_connection(); c = conn.cursor()
    c.execute("""SELECT * FROM verification_codes
                 WHERE email=? AND purpose='ben_email_change' AND used=0 AND expires_at>?
                 ORDER BY id DESC LIMIT 1""", (new_email, now))
    row = c.fetchone()
    if not row or row["code"] != code:
        conn.close()
        return jsonify({"ok": False, "error": "الرمز غير صحيح أو منتهي الصلاحية"})
    c.execute("UPDATE verification_codes SET used=1 WHERE id=?", (row["id"],))
    c.execute("UPDATE beneficiaries SET email=? WHERE id=?", (new_email, ben_id))
    conn.commit(); conn.close()
    session.pop("pending_email_change", None)
    return jsonify({"ok": True})


# ══════════════════════════════════════════
# API: مخيمات حسب المنطقة
# ══════════════════════════════════════════

@app.route("/api/camps-by-area")
def api_camps_by_area():
    area = request.args.get("area","").strip()
    conn = get_connection()
    c = conn.cursor()
    if area:
        c.execute("""
            SELECT id, name, entity_type, governorate, city
            FROM camp_entities
            WHERE is_active=1 AND entity_type!='family'
              AND (governorate LIKE ? OR city LIKE ? OR street LIKE ?)
            ORDER BY name
        """, (f"%{area}%", f"%{area}%", f"%{area}%"))
    else:
        c.execute("""
            SELECT id, name, entity_type, governorate, city
            FROM camp_entities WHERE is_active=1 AND entity_type!='family'
            ORDER BY name
        """)
    camps = [{"id": r["id"], "name": r["name"],
              "type": r["entity_type"],
              "location": f"{r['governorate'] or ''} - {r['city'] or ''}".strip(" -")}
             for r in c.fetchall()]
    conn.close()
    return jsonify(camps)


# ══════════════════════════════════════════
# تسجيل المخيم / اللجنة / العائلة
# ══════════════════════════════════════════

@app.route("/register/camp", methods=["GET", "POST"])
def register_camp():
    if request.method == "POST":
        action = request.form.get("action","")

        if action == "send_code":
            email = request.form.get("email","").strip()
            if not email:
                flash("يرجى إدخال البريد الإلكتروني", "error")
                return render_template("register_camp.html")
            # تحقق: هل البريد مسجل مسبقاً؟
            conn = get_connection()
            c = conn.cursor()
            c.execute("SELECT id FROM camp_entities WHERE email=?", (email,))
            if c.fetchone():
                conn.close()
                flash("هذا البريد الإلكتروني مسجل مسبقاً", "error")
                return render_template("register_camp.html")
            # أنشئ كود تحقق
            code = generate_verification_code()
            expires = (datetime.now() + timedelta(minutes=30)).strftime("%Y-%m-%d %H:%M:%S")
            c.execute("""
                INSERT INTO verification_codes (email, code, purpose, expires_at)
                VALUES (?,?,'camp_register',?)
            """, (email, code, expires))
            conn.commit()
            conn.close()
            try:
                send_org_verification(email, code)
                flash(f"تم إرسال رمز التحقق إلى {email}", "success")
            except Exception:
                flash(f"كود التحقق: {code} (تعذر إرسال البريد)", "warning")
            session["camp_reg_email"] = email
            return render_template("register_camp.html", step="verify", email=email)

        elif action == "verify_code":
            email = request.form.get("email","").strip()
            code  = request.form.get("code","").strip()
            conn  = get_connection()
            c     = conn.cursor()
            c.execute("""
                SELECT id FROM verification_codes
                WHERE email=? AND code=? AND purpose='camp_register'
                  AND used=0 AND expires_at > datetime('now','localtime')
                ORDER BY id DESC LIMIT 1
            """, (email, code))
            row = c.fetchone()
            if not row:
                conn.close()
                flash("الرمز غير صحيح أو منتهي الصلاحية", "error")
                return render_template("register_camp.html", step="verify", email=email)
            c.execute("UPDATE verification_codes SET used=1 WHERE id=?", (row["id"],))
            conn.commit()
            conn.close()
            session["camp_reg_email"]    = email
            session["camp_reg_verified"] = True
            return render_template("register_camp.html", step="form", email=email)

        elif action == "submit":
            if not session.get("camp_reg_verified"):
                flash("يرجى التحقق من بريدك الإلكتروني أولاً", "error")
                return render_template("register_camp.html")

            email       = session.get("camp_reg_email","")
            entity_type = request.form.get("entity_type","camp")
            name        = request.form.get("name","").strip()
            manager     = request.form.get("manager_name","").strip()
            id_number   = request.form.get("id_number","").strip()
            mobile      = request.form.get("mobile","").strip()
            whatsapp    = request.form.get("whatsapp","").strip()
            governorate = request.form.get("governorate","").strip()
            city        = request.form.get("city","").strip()
            street      = request.form.get("street","").strip()
            families    = request.form.get("registered_families",0)
            password    = request.form.get("password","")
            password2   = request.form.get("password2","")

            if not name or not manager or not password:
                flash("يرجى تعبئة جميع الحقول المطلوبة", "error")
                return render_template("register_camp.html", step="form", email=email)
            if password != password2:
                flash("كلمتا المرور غير متطابقتين", "error")
                return render_template("register_camp.html", step="form", email=email)
            if len(id_number) != 9 or not id_number.isdigit():
                flash("رقم الهوية يجب أن يكون 9 أرقام", "error")
                return render_template("register_camp.html", step="form", email=email)

            hashed_pw = generate_password_hash(password)
            conn = get_connection()
            c    = conn.cursor()
            try:
                c.execute("""
                    INSERT INTO camp_entities
                    (entity_type, name, manager_name, id_number, mobile, whatsapp,
                     email, email_verified, governorate, city, street,
                     registered_families, password, is_active)
                    VALUES (?,?,?,?,?,?,?,1,?,?,?,?,?,1)
                """, (entity_type, name, manager, id_number, mobile, whatsapp,
                      email, governorate, city, street, families or 0, hashed_pw))
                new_id = c.lastrowid
                conn.commit()
            except Exception as e:
                conn.close()
                flash(f"خطأ في التسجيل: {e}", "error")
                return render_template("register_camp.html", step="form", email=email)

            # إشعار للمستفيدين الذين كتبوا نفس اسم المخيم يدوياً
            c.execute("""
                SELECT id FROM beneficiaries
                WHERE camp_name=? AND camp_entity_id IS NULL AND self_registered=1
            """, (name,))
            pending_bens = c.fetchall()
            for pb in pending_bens:
                c.execute("""
                    UPDATE beneficiaries
                    SET camp_entity_id=?, beneficiary_status='pending'
                    WHERE id=?
                """, (new_id, pb["id"]))
                c.execute("""
                    INSERT OR IGNORE INTO camp_join_requests (beneficiary_id, camp_entity_id, status)
                    VALUES (?,?,'pending')
                """, (pb["id"], new_id))
            conn.commit()
            conn.close()

            session.pop("camp_reg_email", None)
            session.pop("camp_reg_verified", None)
            flash("تم تسجيل حسابك بنجاح! يمكنك الآن تسجيل الدخول.", "success")
            return redirect(url_for("camp_login"))

    return render_template("register_camp.html", step="email")


# ══════════════════════════════════════════
# دخول المخيم / اللجنة / العائلة
# ══════════════════════════════════════════

@app.route("/camp/login", methods=["GET", "POST"])
def camp_login():
    import datetime as _dt
    try:
      if request.method == "POST":
        email    = (request.form.get("email") or "").strip().lower()
        password = (request.form.get("password") or "").strip()
        ip       = request.remote_addr or "0.0.0.0"
        now      = _dt.datetime.now()
        conn = get_connection(); c = conn.cursor()

        # ── Rate limiting على IP ──
        c.execute("""SELECT COUNT(*) as cnt FROM login_attempts
                     WHERE ip_address=? AND attempt_type='camp'
                     AND success=0 AND created_at >= datetime('now','-15 minutes','localtime')""", (ip,))
        ip_count = c.fetchone()["cnt"]
        if ip_count >= 10:
            conn.close()
            return jsonify({"ok": False, "error": "تم تجاوز الحد المسموح — حاول بعد 15 دقيقة"})                    if request.is_json else (flash("تم تجاوز الحد المسموح — حاول بعد 15 دقيقة", "error") or
                   render_template("camp_login.html"))

        c.execute("SELECT * FROM camp_entities WHERE LOWER(email)=? AND is_active=1", (email,))
        _row = c.fetchone()
        entity = dict(_row) if _row else None

        # ── قفل الحساب ──
        if entity and entity.get("locked_until"):
            try:
                lock_dt = _dt.datetime.fromisoformat(entity["locked_until"])
                if now < lock_dt:
                    remaining = int((lock_dt - now).total_seconds() // 60) + 1
                    conn.close()
                    flash(f"الحساب مقفل — حاول بعد {remaining} دقيقة", "error")
                    return render_template("camp_login.html")
                else:
                    c.execute("UPDATE camp_entities SET failed_attempts=0, locked_until=NULL WHERE id=?", (entity["id"],))
                    conn.commit()
            except Exception:
                pass

        # ── التحقق من كلمة المرور ──
        if entity and entity["password"] and check_password_hash(entity["password"], password):
            c.execute("UPDATE camp_entities SET failed_attempts=0, locked_until=NULL WHERE id=?", (entity["id"],))
            c.execute("INSERT INTO login_attempts (identifier,attempt_type,ip_address,success) VALUES (?,?,?,1)",
                      (email, "camp", ip))
            conn.commit(); conn.close()
            session["camp_id"]   = entity["id"]
            session["camp_name"] = entity["name"]
            session["camp_type"] = entity["entity_type"]
            session["camp_last_active"] = _dt.datetime.utcnow().isoformat()
            return redirect(url_for("camp_dashboard"))

        # ── فشل الدخول ──
        c.execute("INSERT INTO login_attempts (identifier,attempt_type,ip_address,success) VALUES (?,?,?,0)",
                  (email, "camp", ip))
        if entity:
            fails = (entity.get("failed_attempts") or 0) + 1
            locked = None
            if fails >= 5:
                locked = (now + _dt.timedelta(minutes=30)).isoformat()
                fails  = 0
            c.execute("UPDATE camp_entities SET failed_attempts=?, locked_until=? WHERE id=?",
                      (fails, locked, entity["id"]))
            remaining_tries = 5 - fails if not locked else 0
            conn.commit(); conn.close()
            if locked:
                flash("تم قفل الحساب 30 دقيقة بسبب المحاولات المتكررة", "error")
            else:
                flash(f"كلمة المرور غير صحيحة — تبقى {remaining_tries} محاولة", "error")
        else:
            conn.commit(); conn.close()
            flash("البريد الإلكتروني غير مسجل", "error")
      return render_template("camp_login.html")
    except Exception as _e:
        import traceback
        return f"<pre style='direction:ltr'>{traceback.format_exc()}</pre>", 500


# ── نسيت كلمة مرور المخيم ──
@app.route("/api/camp/forgot-password", methods=["POST"])
def api_camp_forgot_password():
    data  = request.get_json(force=True) or {}
    email = (data.get("email") or "").strip().lower()
    if not email:
        return jsonify({"ok": False, "error": "أدخل البريد الإلكتروني"})
    conn = get_connection(); c = conn.cursor()
    c.execute("SELECT id, name FROM camp_entities WHERE LOWER(email)=? AND is_active=1", (email,))
    entity = c.fetchone()
    if not entity:
        conn.close()
        return jsonify({"ok": False, "error": "البريد الإلكتروني غير مسجل"})
    code = generate_verification_code()
    import datetime as _dt
    expires = (_dt.datetime.now() + _dt.timedelta(minutes=15)).isoformat()
    c.execute("DELETE FROM verification_codes WHERE email=? AND purpose='camp_reset'", (email,))
    c.execute("INSERT INTO verification_codes (email,code,purpose,expires_at) VALUES (?,?,?,?)",
              (email, code, "camp_reset", expires))
    conn.commit(); conn.close()
    try:
        send_org_verification(
            email,
            entity["name"],   # org_name
            entity["name"],   # admin_name (نفس الاسم كبديل)
            code,
            "",               # org_code (غير مطلوب لإعادة التعيين)
            email             # username
        )
        return jsonify({"ok": True})
    except Exception as e:
        return jsonify({"ok": False, "error": f"خطأ في الإرسال: {str(e)}"})


@app.route("/api/camp/reset-password", methods=["POST"])
def api_camp_reset_password():
    data     = request.get_json(force=True) or {}
    email    = (data.get("email") or "").strip().lower()
    code     = (data.get("code") or "").strip()
    new_pass = (data.get("password") or "").strip()
    if len(new_pass) < 6:
        return jsonify({"ok": False, "error": "كلمة المرور قصيرة جداً"})
    conn = get_connection(); c = conn.cursor()
    import datetime as _dt
    c.execute("""SELECT * FROM verification_codes
                 WHERE email=? AND code=? AND purpose='camp_reset'
                 AND used=0 AND expires_at > datetime('now','localtime')""", (email, code))
    row = c.fetchone()
    if not row:
        conn.close()
        return jsonify({"ok": False, "error": "الرمز غير صحيح أو منتهي الصلاحية"})
    c.execute("UPDATE camp_entities SET password=?, failed_attempts=0, locked_until=NULL WHERE email=?",
              (generate_password_hash(new_pass), email))
    c.execute("UPDATE verification_codes SET used=1 WHERE email=? AND purpose='camp_reset'", (email,))
    conn.commit(); conn.close()
    return jsonify({"ok": True})


def camp_login_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if "camp_id" not in session:
            return redirect(url_for("camp_login"))
        return f(*args, **kwargs)
    return decorated


@app.route("/camp/logout")
def camp_logout():
    session.pop("camp_id", None)
    session.pop("camp_name", None)
    session.pop("camp_type", None)
    return redirect(url_for("camp_login"))


@app.route("/camp/dashboard")
@camp_login_required
def camp_dashboard():
    camp_id = session["camp_id"]
    conn = get_connection()
    c    = conn.cursor()
    c.execute("SELECT * FROM camp_entities WHERE id=?", (camp_id,))
    entity = dict(c.fetchone())

    # طلبات الانضمام
    c.execute("""
        SELECT cjr.id, cjr.created_at, cjr.status,
               b.full_name, b.id_number, b.phone, b.address, b.gender, b.id as ben_id
        FROM camp_join_requests cjr
        JOIN beneficiaries b ON cjr.beneficiary_id=b.id
        WHERE cjr.camp_entity_id=?
        ORDER BY cjr.id DESC
    """, (camp_id,))
    requests_list = [dict(r) for r in c.fetchall()]

    # المستفيدون المقبولون
    c.execute("""
        SELECT * FROM beneficiaries
        WHERE camp_entity_id=? AND beneficiary_status='in_camp'
        ORDER BY full_name
    """, (camp_id,))
    members = [dict(r) for r in c.fetchall()]

    # ── إحصاءات تفصيلية ──
    orphans_count  = sum(1 for m in members if m.get("has_orphans"))
    pregnant_count = sum(1 for m in members if m.get("wife_pregnant"))
    nursing_count  = sum(1 for m in members if m.get("wife_nursing"))

    # مرضى: تصنيف حسب نوع الحالة
    sick_breakdown = []
    sick_total = 0
    if members:
        ids = [m["id"] for m in members]
        ph  = ",".join("?" * len(ids))
        c.execute(f"""SELECT COALESCE(condition_type,'غير محدد') as condition_type,
                             COUNT(DISTINCT beneficiary_id) as n
                      FROM beneficiary_health
                      WHERE beneficiary_id IN ({ph})
                      GROUP BY condition_type ORDER BY n DESC""", ids)
        sick_breakdown = [dict(r) for r in c.fetchall()]
        sick_total = sum(r["n"] for r in sick_breakdown)

    # آخر التوزيعات — مع الجهة المانحة (من المصدرين)
    c.execute("""
        SELECT b.full_name, cb.benefit_type, cb.value, cb.quantity,
               cb.benefit_date, NULL as donor_name, 'يدوي' as src
        FROM camp_benefits cb
        JOIN beneficiaries b ON cb.beneficiary_id=b.id
        WHERE cb.camp_entity_id=?
        UNION ALL
        SELECT b.full_name,
               COALESCE(dr.item_name, d.title) as benefit_type,
               dr.value, dr.quantity,
               d.distribution_date as benefit_date,
               d.donor_name, 'توزيع' as src
        FROM camp_dist_records dr
        JOIN camp_distributions d ON dr.distribution_id = d.id
        JOIN beneficiaries b ON dr.beneficiary_id = b.id
        WHERE d.camp_entity_id=? AND dr.received=1
        ORDER BY benefit_date DESC LIMIT 8
    """, (camp_id, camp_id))
    recent_benefits = [dict(r) for r in c.fetchall()]

    conn.close()
    return render_template("camp_dashboard.html",
        entity=entity,
        requests_list=requests_list,
        members=members,
        recent_benefits=recent_benefits,
        orphans_count=orphans_count,
        pregnant_count=pregnant_count,
        nursing_count=nursing_count,
        sick_total=sick_total,
        sick_breakdown=sick_breakdown)


@app.route("/camp/join-request/<int:req_id>/<action>")
@camp_login_required
def camp_handle_request(req_id, action):
    if action not in ("approve", "reject"):
        return redirect(url_for("camp_dashboard"))
    camp_id = session["camp_id"]
    conn = get_connection()
    c    = conn.cursor()
    c.execute("SELECT * FROM camp_join_requests WHERE id=? AND camp_entity_id=?", (req_id, camp_id))
    req = c.fetchone()
    if req:
        now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        if action == "approve":
            # جلب نوع الكيان
            c.execute("SELECT entity_type, name FROM camp_entities WHERE id=?", (camp_id,))
            _ent = c.fetchone()
            entity_type = _ent["entity_type"] if _ent else "camp"
            entity_label = {"camp": "مخيم", "committee": "لجنة", "family": "عائلة"}.get(entity_type, "كيان")

            # ── فحص التكرار حسب نوع الكيان ──
            if entity_type == "family":
                # عائلة: يُسمح بمخيم/لجنة منفصل — فقط منع تكرار العائلة
                c.execute("""SELECT b.full_name, ce.name as fam_name
                             FROM beneficiaries b
                             LEFT JOIN camp_entities ce ON b.family_entity_id = ce.id
                             WHERE b.id = ? AND b.family_entity_id IS NOT NULL""",
                          (req["beneficiary_id"],))
                dup = c.fetchone()
                if dup:
                    flash(f"⚠️ '{dup['full_name']}' مسجل بالفعل في عائلة '{dup['fam_name']}' — لا يمكن الانضمام لعائلتين", "warning")
                    conn.close()
                    return redirect(url_for("camp_dashboard"))
            else:
                # مخيم أو لجنة: يُسمح بعائلة منفصلة — فقط منع تكرار المخيم/اللجنة
                c.execute("""SELECT b.full_name, ce.name as entity_name
                             FROM beneficiaries b
                             LEFT JOIN camp_entities ce ON b.camp_entity_id = ce.id
                             WHERE b.id = ? AND b.beneficiary_status = 'in_camp'""",
                          (req["beneficiary_id"],))
                dup = c.fetchone()
                if dup and dup["entity_name"]:
                    if dup["entity_name"] == session.get("camp_name"):
                        flash(f"⚠️ '{dup['full_name']}' مسجل بالفعل في {entity_label}ك — لا يمكن القبول مرتين", "warning")
                    else:
                        flash(f"⚠️ '{dup['full_name']}' منضم لـ '{dup['entity_name']}' — لا يمكن القبول في {entity_label}ين بالوقت نفسه", "warning")
                    conn.close()
                    return redirect(url_for("camp_dashboard"))

            # ── موافقة سليمة ──
            c.execute("UPDATE camp_join_requests SET status='approved', resolved_at=? WHERE id=?", (now, req_id))
            if entity_type == "family":
                # العائلة: عمود منفصل — لا تغيّر beneficiary_status
                c.execute("UPDATE beneficiaries SET family_entity_id=? WHERE id=?",
                          (camp_id, req["beneficiary_id"]))
                flash(f"تم قبول المستفيد في {entity_label}ك", "success")
            else:
                # مخيم أو لجنة
                c.execute("""UPDATE beneficiaries
                             SET beneficiary_status='in_camp', camp_entity_id=?
                             WHERE id=?""", (camp_id, req["beneficiary_id"]))
                flash(f"تم قبول المستفيد في {entity_label}ك", "success")
        else:
            c.execute("UPDATE camp_join_requests SET status='rejected', resolved_at=? WHERE id=?", (now, req_id))
            c.execute("UPDATE beneficiaries SET beneficiary_status='independent', camp_entity_id=NULL WHERE id=?", (req["beneficiary_id"],))
            flash("تم رفض الطلب، سيُسجل المستفيد كمستفيد مستقل", "info")
        conn.commit()
    conn.close()
    return redirect(url_for("camp_dashboard"))


@app.route("/camp/add-beneficiary", methods=["GET","POST"])
@camp_login_required
def camp_add_beneficiary():
    camp_id = session["camp_id"]
    if request.method == "POST":
        full_name   = request.form.get("full_name","").strip()
        id_number   = request.form.get("id_number","").strip()
        phone       = request.form.get("phone","").strip()
        gender      = request.form.get("gender","male")
        family_size = request.form.get("family_size",1)
        address     = request.form.get("address","").strip()
        notes       = request.form.get("notes","").strip()
        parts       = full_name.split()
        family_last_name = parts[3] if len(parts) >= 4 else (parts[-1] if parts else "")
        conn = get_connection()
        c    = conn.cursor()
        # تحقق من رقم الهوية المكرر
        if id_number:
            c.execute("SELECT id FROM beneficiaries WHERE TRIM(id_number)=TRIM(?)", (id_number,))
            existing = c.fetchone()
            if existing:
                # ربط المستفيد بالمخيم إن لم يكن مرتبطاً
                c.execute("UPDATE beneficiaries SET camp_entity_id=?, beneficiary_status='in_camp' WHERE id=?",
                          (camp_id, existing["id"]))
                conn.commit()
                conn.close()
                flash("هذا المستفيد مسجل مسبقاً — تم ربطه بالمخيم تلقائياً", "warning")
                return redirect(url_for("camp_dashboard"))
        c.execute("""
            INSERT INTO beneficiaries
            (full_name, id_number, phone, gender, family_size, address,
             camp_entity_id, beneficiary_status, family_last_name,
             self_registered, beneficiary_type, org_id, notes)
            VALUES (?,?,?,?,?,?,?,'in_camp',?,0,'person',1,?)
        """, (full_name, id_number, phone, gender, family_size, address,
              camp_id, family_last_name, notes))
        conn.commit()
        conn.close()
        flash("تم إضافة المستفيد بنجاح", "success")
        return redirect(url_for("camp_dashboard"))
    return render_template("camp_add_beneficiary.html")

# ══════════════════════════════════════════
# لوحة تحكم المخيم — ميزات إضافية
# ══════════════════════════════════════════

@app.route("/camp/beneficiary/<int:ben_id>")
@camp_login_required
def camp_beneficiary_profile(ben_id):
    camp_id = session["camp_id"]
    conn = get_connection(); c = conn.cursor()
    c.execute("""SELECT * FROM beneficiaries WHERE id=? AND camp_entity_id=?""", (ben_id, camp_id))
    ben = c.fetchone()
    if not ben:
        conn.close()
        flash("المستفيد غير موجود", "error")
        return redirect(url_for("camp_dashboard"))
    ben = dict(ben)
    c.execute("""SELECT * FROM camp_benefits WHERE beneficiary_id=? AND camp_entity_id=?
                 ORDER BY benefit_date DESC""", (ben_id, camp_id))
    benefits = [dict(r) for r in c.fetchall()]
    # جلسات التوزيع الرسمية المرتبطة بهذا المستفيد
    c.execute("""SELECT dr.*, d.title as dist_title, d.distribution_date, d.donor_name
                 FROM camp_dist_records dr
                 JOIN camp_distributions d ON dr.distribution_id=d.id
                 WHERE dr.beneficiary_id=? AND dr.camp_entity_id=?
                 ORDER BY d.distribution_date DESC""", (ben_id, camp_id))
    dist_records = [dict(r) for r in c.fetchall()]
    c.execute("SELECT * FROM camp_entities WHERE id=?", (camp_id,))
    _ent = c.fetchone()
    entity = dict(_ent) if _ent else {"name": session.get("camp_name","المخيم")}
    conn.close()
    return render_template("camp_beneficiary.html", ben=ben, benefits=benefits,
                           dist_records=dist_records, entity=entity)

@app.route("/camp/benefit/add", methods=["POST"])
@camp_login_required
def camp_add_benefit():
    camp_id    = session["camp_id"]
    ben_id     = request.form.get("beneficiary_id","")
    btype      = request.form.get("benefit_type","").strip()
    value      = request.form.get("value","").strip()
    quantity   = request.form.get("quantity","").strip()
    notes      = request.form.get("notes","").strip()
    bdate      = request.form.get("benefit_date","").strip()
    if not ben_id or not btype:
        flash("يرجى تعبئة الحقول المطلوبة", "error")
        return redirect(url_for("camp_dashboard"))
    from datetime import date as _date
    if not bdate:
        bdate = _date.today().isoformat()
    conn = get_connection(); c = conn.cursor()
    c.execute("""INSERT INTO camp_benefits
        (camp_entity_id, beneficiary_id, benefit_type, value, quantity, notes, benefit_date)
        VALUES (?,?,?,?,?,?,?)""",
        (camp_id, ben_id, btype, value, quantity, notes, bdate))
    conn.commit(); conn.close()
    flash("تم تسجيل الاستفادة بنجاح", "success")
    return redirect(request.referrer or url_for("camp_dashboard"))

@app.route("/camp/benefit/delete/<int:bid>", methods=["POST"])
@camp_login_required
def camp_delete_benefit(bid):
    camp_id = session["camp_id"]
    conn = get_connection(); c = conn.cursor()
    c.execute("DELETE FROM camp_benefits WHERE id=? AND camp_entity_id=?", (bid, camp_id))
    conn.commit(); conn.close()
    flash("تم حذف السجل", "success")
    return redirect(request.referrer or url_for("camp_dashboard"))

@app.route("/camp/upload-logo", methods=["POST"])
@camp_login_required
def camp_upload_logo():
    """رفع شعار/صورة المخيم"""
    import os, uuid
    camp_id = session["camp_id"]
    f = request.files.get("logo")
    if not f or not f.filename:
        return jsonify({"ok": False, "error": "لم يُرفع ملف"})
    ext = f.filename.rsplit(".", 1)[-1].lower()
    if ext not in ("jpg", "jpeg", "png", "gif", "webp"):
        return jsonify({"ok": False, "error": "صيغة غير مدعومة"})
    uploads_dir = os.path.join(os.path.dirname(__file__), "static", "uploads")
    os.makedirs(uploads_dir, exist_ok=True)
    fname = f"camp_{camp_id}_{uuid.uuid4().hex[:8]}.{ext}"
    f.save(os.path.join(uploads_dir, fname))
    conn = get_connection(); c = conn.cursor()
    c.execute("UPDATE camp_entities SET logo_image=? WHERE id=?", (fname, camp_id))
    conn.commit(); conn.close()
    return jsonify({"ok": True, "url": f"/static/uploads/{fname}"})


@app.route("/camp/settings", methods=["GET","POST"])
@camp_login_required
def camp_settings():
    camp_id = session["camp_id"]
    conn = get_connection(); c = conn.cursor()
    if request.method == "POST":
        action = request.form.get("action","")
        if action == "update_info":
            name        = request.form.get("name","").strip()
            manager     = request.form.get("manager_name","").strip()
            mobile      = request.form.get("mobile","").strip()
            whatsapp    = request.form.get("whatsapp","").strip()
            governorate = request.form.get("governorate","").strip()
            city        = request.form.get("city","").strip()
            street      = request.form.get("street","").strip()
            families    = request.form.get("registered_families",0)
            c.execute("""UPDATE camp_entities SET name=?,manager_name=?,mobile=?,whatsapp=?,
                         governorate=?,city=?,street=?,registered_families=? WHERE id=?""",
                      (name,manager,mobile,whatsapp,governorate,city,street,families or 0,camp_id))
            conn.commit()
            session["camp_name"] = name
            flash("تم حفظ البيانات", "success")
        elif action == "change_password":
            old_pw  = request.form.get("old_password","")
            new_pw  = request.form.get("new_password","")
            new_pw2 = request.form.get("new_password2","")
            c.execute("SELECT password FROM camp_entities WHERE id=?", (camp_id,))
            row = c.fetchone()
            if not row or not check_password_hash(row["password"], old_pw):
                flash("كلمة المرور الحالية غير صحيحة", "error")
            elif new_pw != new_pw2:
                flash("كلمتا المرور الجديدتان غير متطابقتين", "error")
            elif len(new_pw) < 6:
                flash("كلمة المرور قصيرة جداً (6 أحرف على الأقل)", "error")
            else:
                c.execute("UPDATE camp_entities SET password=? WHERE id=?",
                          (generate_password_hash(new_pw), camp_id))
                conn.commit()
                flash("تم تغيير كلمة المرور", "success")
        conn.close()
        return redirect(url_for("camp_settings"))
    c.execute("SELECT * FROM camp_entities WHERE id=?", (camp_id,))
    entity = dict(c.fetchone())
    conn.close()
    return render_template("camp_settings.html", entity=entity)

@app.route("/camp/reports")
@camp_login_required
def camp_reports():
    camp_id = session["camp_id"]
    conn = get_connection(); c = conn.cursor()
    c.execute("SELECT * FROM camp_entities WHERE id=?", (camp_id,))
    entity = dict(c.fetchone())
    # إجماليات الاستفادة حسب النوع
    c.execute("""SELECT benefit_type, COUNT(*) as cnt
                 FROM camp_benefits WHERE camp_entity_id=?
                 GROUP BY benefit_type ORDER BY cnt DESC""", (camp_id,))
    by_type = [dict(r) for r in c.fetchall()]
    # آخر 30 توزيع
    c.execute("""SELECT cb.*, b.full_name, b.id_number
                 FROM camp_benefits cb
                 JOIN beneficiaries b ON cb.beneficiary_id=b.id
                 WHERE cb.camp_entity_id=?
                 ORDER BY cb.created_at DESC LIMIT 30""", (camp_id,))
    recent = [dict(r) for r in c.fetchall()]
    # إجمالي المستفيدين
    c.execute("SELECT COUNT(*) as cnt FROM beneficiaries WHERE camp_entity_id=? AND beneficiary_status='in_camp'", (camp_id,))
    total_members = c.fetchone()["cnt"]
    # إجمالي سجلات الاستفادة
    c.execute("SELECT COUNT(*) as cnt FROM camp_benefits WHERE camp_entity_id=?", (camp_id,))
    total_benefits = c.fetchone()["cnt"]
    conn.close()
    return render_template("camp_reports.html",
        entity=entity, by_type=by_type, recent=recent,
        total_members=total_members, total_benefits=total_benefits)


# ══════════════════════════════════════════
# Auto-Deploy Webhook
# ══════════════════════════════════════════
import subprocess, hashlib, hmac

@app.route("/deploy/webhook", methods=["POST"])
def deploy_webhook():
    import os
    # التحقق من التوكن
    token = os.environ.get("DEPLOY_TOKEN", "")
    if token:
        sent = request.headers.get("X-Deploy-Token", "")
        if sent != token:
            return jsonify({"error": "Unauthorized"}), 401
    project_dir = os.path.dirname(os.path.abspath(__file__))
    result = subprocess.run(
        ["git", "pull"],
        cwd=project_dir,
        capture_output=True, text=True, timeout=60
    )
    return jsonify({
        "status": "ok",
        "stdout": result.stdout,
        "returncode": result.returncode
    })


# ══════════════════════════════════════════
# API: بوابة المستفيد — دخول وتسجيل بالبريد
# ══════════════════════════════════════════

# ── Session timeout: 5 دقائق ──
@app.before_request
def check_beneficiary_session_timeout():
    try:
        import datetime
        if session.get("beneficiary_id"):
            last = session.get("ben_last_active")
            if last:
                diff = (datetime.datetime.utcnow() - datetime.datetime.fromisoformat(last)).total_seconds()
                if diff > 1800:  # 30 دقيقة
                    session.pop("beneficiary_id", None)
                    session.pop("beneficiary_name", None)
                    session.pop("ben_last_active", None)
                    return
            session["ben_last_active"] = datetime.datetime.utcnow().isoformat()
    except Exception:
        pass


@app.route("/api/beneficiary/ping", methods=["POST"])
def api_beneficiary_ping():
    """تجديد الـ session — يُستدعى كل 2 دقيقة من الصفحة"""
    import datetime
    if not session.get("beneficiary_id"):
        return jsonify({"ok": False}), 401
    session["ben_last_active"] = datetime.datetime.utcnow().isoformat()
    return jsonify({"ok": True})


@app.route("/api/beneficiary/login", methods=["POST"])
def api_beneficiary_login():
    """دخول المستفيد: قفل بعد 5 محاولات + rate limiting على IP"""
    import datetime
    data      = request.get_json(force=True) or {}
    id_number = data.get("id_number", "").strip()
    password  = data.get("password", "").strip()
    ip        = request.headers.get("X-Forwarded-For", request.remote_addr or "").split(",")[0].strip()

    if not id_number or not password:
        return jsonify({"ok": False, "error": "يرجى إدخال رقم الهوية وكلمة المرور"})

    now     = datetime.datetime.utcnow()
    win15   = (now - datetime.timedelta(minutes=15)).strftime("%Y-%m-%d %H:%M:%S")
    now_str = now.strftime("%Y-%m-%d %H:%M:%S")

    conn = get_connection()
    c    = conn.cursor()

    # ── Rate limiting: IP ──
    c.execute("""
        SELECT COUNT(*) FROM login_attempts
        WHERE ip_address=? AND success=0 AND attempt_type='beneficiary'
        AND created_at > ?
    """, (ip, win15))
    ip_fails = c.fetchone()[0]
    if ip_fails >= 10:
        conn.close()
        return jsonify({"ok": False, "error": "محاولات كثيرة من هذا الجهاز، انتظر 15 دقيقة"})

    # ── جلب الحساب ──
    c.execute("""
        SELECT b.*, ce.name as camp_name
        FROM beneficiaries b
        LEFT JOIN camp_entities ce ON b.camp_entity_id = ce.id
        WHERE b.id_number = ? AND b.self_registered = 1
    """, (id_number,))
    ben = c.fetchone()

    if not ben:
        c.execute("INSERT INTO login_attempts (identifier,ip_address,success,attempt_type) VALUES (?,?,0,'beneficiary')",
                  (id_number, ip))
        conn.commit()
        conn.close()
        return jsonify({"ok": False, "error": "رقم الهوية أو كلمة المرور غير صحيحة"})

    # ── قفل الحساب ──
    locked_until = ben["locked_until"] if ben["locked_until"] else None
    if locked_until and locked_until > now_str:
        conn.close()
        return jsonify({"ok": False, "error": "الحساب مقفل مؤقتاً بسبب محاولات متعددة، حاول بعد 30 دقيقة"})

    # ── التحقق من كلمة المرور ──
    if ben["self_reg_password"] and check_password_hash(ben["self_reg_password"], password):
        # نجح — إعادة تعيين المحاولات
        c.execute("UPDATE beneficiaries SET failed_attempts=0, locked_until=NULL WHERE id=?", (ben["id"],))
        c.execute("INSERT INTO login_attempts (identifier,ip_address,success,attempt_type) VALUES (?,?,1,'beneficiary')",
                  (id_number, ip))
        conn.commit()
        conn.close()
        session["beneficiary_id"]   = ben["id"]
        session["beneficiary_name"] = ben["full_name"]
        session["ben_last_active"]  = now.isoformat()
        return jsonify({"ok": True})

    # ── فشل ──
    fails = (ben["failed_attempts"] or 0) + 1
    lock  = None
    if fails >= 5:
        lock  = (now + datetime.timedelta(minutes=30)).strftime("%Y-%m-%d %H:%M:%S")
        fails = 0
    c.execute("UPDATE beneficiaries SET failed_attempts=?, locked_until=? WHERE id=?",
              (fails, lock, ben["id"]))
    c.execute("INSERT INTO login_attempts (identifier,ip_address,success,attempt_type) VALUES (?,?,0,'beneficiary')",
              (id_number, ip))
    conn.commit()
    conn.close()

    remaining = 5 - fails
    if lock:
        return jsonify({"ok": False, "error": "تم قفل الحساب 30 دقيقة بسبب محاولات متعددة"})
    return jsonify({"ok": False, "error": f"كلمة المرور غير صحيحة — متبقي {remaining} محاولة قبل القفل"})


@app.route("/api/beneficiary/forgot-password", methods=["POST"])
def api_beneficiary_forgot_password():
    """إرسال رمز إعادة تعيين كلمة المرور عبر البريد الإلكتروني"""
    import datetime
    data  = request.get_json(force=True) or {}
    email = data.get("email", "").strip().lower()
    if not email:
        return jsonify({"ok": False, "error": "أدخل بريدك الإلكتروني"})
    conn = get_connection()
    c    = conn.cursor()
    c.execute("SELECT id, email, full_name FROM beneficiaries WHERE email=? AND self_registered=1", (email,))
    ben = c.fetchone()
    if not ben:
        conn.close()
        return jsonify({"ok": False, "error": "لا يوجد حساب مرتبط بهذا البريد"})
    code       = generate_verification_code()
    expires_at = (datetime.datetime.utcnow() + datetime.timedelta(minutes=15)).strftime("%Y-%m-%d %H:%M:%S")
    c.execute("DELETE FROM verification_codes WHERE email=? AND purpose='ben_reset'", (ben["email"],))
    c.execute("INSERT INTO verification_codes (email,code,purpose,expires_at) VALUES (?,?,?,?)",
              (ben["email"], code, "ben_reset", expires_at))
    conn.commit()
    conn.close()
    try:
        send_org_verification(ben["email"], "إعادة تعيين كلمة المرور", ben["full_name"], code, "", "")
    except Exception as e:
        app.logger.error(f"Reset email error: {e}")
        return jsonify({"ok": False, "error": "فشل إرسال البريد"})
    return jsonify({"ok": True})


@app.route("/api/beneficiary/reset-password", methods=["POST"])
def api_beneficiary_reset_password():
    """تعيين كلمة مرور جديدة بعد التحقق من OTP"""
    import datetime
    data   = request.get_json(force=True) or {}
    email  = data.get("email", "").strip().lower()
    code   = data.get("code", "").strip()
    new_pw = data.get("password", "")
    if not email or not code or not new_pw or len(new_pw) < 6:
        return jsonify({"ok": False, "error": "بيانات ناقصة أو كلمة المرور قصيرة"})
    conn = get_connection()
    c    = conn.cursor()
    c.execute("SELECT id, email FROM beneficiaries WHERE email=? AND self_registered=1", (email,))
    ben = c.fetchone()
    if not ben:
        conn.close()
        return jsonify({"ok": False, "error": "حساب غير موجود"})
    now = datetime.datetime.utcnow().strftime("%Y-%m-%d %H:%M:%S")
    c.execute("""
        SELECT * FROM verification_codes
        WHERE email=? AND purpose='ben_reset' AND used=0 AND expires_at > ?
        ORDER BY id DESC LIMIT 1
    """, (email, now))
    row = c.fetchone()
    if not row or row["code"] != code:
        conn.close()
        return jsonify({"ok": False, "error": "الرمز غير صحيح أو منتهي الصلاحية"})
    c.execute("UPDATE verification_codes SET used=1 WHERE id=?", (row["id"],))
    c.execute("UPDATE beneficiaries SET self_reg_password=?, failed_attempts=0, locked_until=NULL WHERE id=?",
              (generate_password_hash(new_pw), ben["id"]))
    conn.commit()
    conn.close()
    return jsonify({"ok": True})


@app.route("/api/beneficiary/send-otp", methods=["POST"])
def api_beneficiary_send_otp():
    """إرسال رمز تحقق إلى البريد الإلكتروني للمستفيد الجديد"""
    data      = request.get_json(force=True) or {}
    email     = data.get("email", "").strip().lower()
    id_number = data.get("id_number", "").strip()
    if not email or not id_number:
        return jsonify({"ok": False, "error": "يرجى إدخال البريد ورقم الهوية"})
    # تحقق من عدم وجود حساب مسبق بنفس رقم الهوية
    conn = get_connection()
    c = conn.cursor()
    c.execute("SELECT id FROM beneficiaries WHERE id_number=? AND self_registered=1", (id_number,))
    if c.fetchone():
        conn.close()
        return jsonify({"ok": False, "error": "رقم الهوية مسجل مسبقاً، يرجى تسجيل الدخول"})
    # توليد كود 6 أرقام وحفظه
    import datetime
    code       = generate_verification_code()
    expires_at = (datetime.datetime.utcnow() + datetime.timedelta(minutes=15)).strftime("%Y-%m-%d %H:%M:%S")
    c.execute("DELETE FROM verification_codes WHERE email=? AND purpose='ben_register'", (email,))
    c.execute(
        "INSERT INTO verification_codes (email, code, purpose, expires_at) VALUES (?,?,?,?)",
        (email, code, "ben_register", expires_at)
    )
    conn.commit()
    conn.close()
    # حفظ البريد ورقم الهوية في الجلسة مؤقتاً
    session["ben_reg_email"]     = email
    session["ben_reg_id_number"] = id_number
    # إرسال البريد
    try:
        send_org_verification(email, "تسجيل مستفيد جديد", "مستفيد عزيز", code, "", "")
    except Exception as e:
        app.logger.error(f"OTP email error: {e}")
        return jsonify({"ok": False, "error": "فشل إرسال البريد، تحقق من العنوان وأعد المحاولة"})
    return jsonify({"ok": True})


@app.route("/api/beneficiary/verify-otp", methods=["POST"])
def api_beneficiary_verify_otp():
    """التحقق من رمز OTP المرسل للبريد"""
    data  = request.get_json(force=True) or {}
    email = data.get("email", "").strip().lower()
    code  = data.get("code", "").strip()
    if not email or not code:
        return jsonify({"ok": False, "error": "بيانات ناقصة"})
    import datetime
    now  = datetime.datetime.utcnow().strftime("%Y-%m-%d %H:%M:%S")
    conn = get_connection()
    c    = conn.cursor()
    c.execute("""
        SELECT * FROM verification_codes
        WHERE email=? AND purpose='ben_register' AND used=0 AND expires_at > ?
        ORDER BY id DESC LIMIT 1
    """, (email, now))
    row = c.fetchone()
    if not row or row["code"] != code:
        conn.close()
        return jsonify({"ok": False, "error": "الرمز غير صحيح أو منتهي الصلاحية"})
    c.execute("UPDATE verification_codes SET used=1 WHERE id=?", (row["id"],))
    conn.commit()
    conn.close()
    session["ben_reg_verified"] = True
    return jsonify({"ok": True})


@app.route("/api/beneficiary/register", methods=["POST"])
def api_beneficiary_register():
    """إتمام تسجيل المستفيد بعد التحقق من البريد"""
    if not session.get("ben_reg_verified"):
        return jsonify({"ok": False, "error": "لم يتم التحقق من البريد"})
    data      = request.get_json(force=True) or {}
    email     = session.get("ben_reg_email", "").strip()
    id_number = session.get("ben_reg_id_number", "").strip()
    full_name = data.get("full_name", "").strip()
    phone     = data.get("phone", "").strip()
    password  = data.get("password", "")
    if not full_name or not password or len(password) < 6:
        return jsonify({"ok": False, "error": "يرجى إدخال الاسم وكلمة المرور (6 أحرف على الأقل)"})
    parts            = full_name.split()
    family_last_name = parts[3] if len(parts) >= 4 else (parts[-1] if parts else "")
    hashed_pw        = generate_password_hash(password)
    conn = get_connection()
    c    = conn.cursor()
    # تحقق مجدداً من عدم التكرار
    c.execute("SELECT id FROM beneficiaries WHERE id_number=? AND self_registered=1", (id_number,))
    if c.fetchone():
        conn.close()
        return jsonify({"ok": False, "error": "رقم الهوية مسجل مسبقاً"})
    c.execute("""
        INSERT INTO beneficiaries
            (full_name, id_number, phone, email, family_last_name,
             self_registered, beneficiary_status, self_reg_password, beneficiary_type, org_id)
        VALUES (?,?,?,?,?, 1,'independent',?,'person',1)
    """, (full_name, id_number, phone, email, family_last_name, hashed_pw))
    new_id = c.lastrowid
    conn.commit()
    conn.close()
    # تنظيف الجلسة
    session.pop("ben_reg_email", None)
    session.pop("ben_reg_id_number", None)
    session.pop("ben_reg_verified", None)
    # تسجيل الدخول مباشرة
    session["beneficiary_id"]   = new_id
    session["beneficiary_name"] = full_name
    return jsonify({"ok": True})

# ══════════════════════════════════════════════════════════════════
# نظام إدارة المخيم الكامل — مستودع / توزيع / أرشيف / تنبيهات
# ══════════════════════════════════════════════════════════════════

# ── مساعد: جلب المستفيدين مع تصنيفاتهم ──
def get_classified_members(camp_id, conn):
    c = conn.cursor()
    c.execute("""SELECT b.*,
        (SELECT COUNT(*) FROM camp_dist_records dr WHERE dr.beneficiary_id=b.id AND dr.received=1) as total_received,
        (SELECT MAX(d.distribution_date) FROM camp_dist_records dr2
         JOIN camp_distributions d ON dr2.distribution_id=d.id
         WHERE dr2.beneficiary_id=b.id AND dr2.received=1) as last_received_date
        FROM beneficiaries b
        WHERE b.camp_entity_id=? AND b.beneficiary_status='in_camp'
        ORDER BY b.full_name""", (camp_id,))
    return [dict(r) for r in c.fetchall()]

# ── مستودع الاستلام ──
@app.route("/camp/inventory", methods=["GET","POST"])
@camp_login_required
def camp_inventory():
    camp_id = session["camp_id"]
    conn = get_connection(); c = conn.cursor()
    if request.method == "POST":
        action = request.form.get("action","add")
        if action == "add":
            donor    = request.form.get("donor_name","").strip()
            dtype    = request.form.get("donor_type","organization")
            item     = request.form.get("item_name","").strip()
            qty      = request.form.get("quantity","0")
            unit     = request.form.get("unit","وحدة").strip()
            rdate    = request.form.get("receive_date","")
            notes    = request.form.get("notes","").strip()
            proof    = None
            if 'proof_image' in request.files:
                f = request.files['proof_image']
                if f and f.filename:
                    import os, uuid
                    ext = f.filename.rsplit('.',1)[-1].lower()
                    fname = f"inv_{uuid.uuid4().hex[:8]}.{ext}"
                    fpath = os.path.join(os.path.dirname(__file__),'static','uploads','inventory')
                    os.makedirs(fpath, exist_ok=True)
                    f.save(os.path.join(fpath, fname))
                    proof = f"inventory/{fname}"
            if donor and item and rdate:
                c.execute("""INSERT INTO camp_inventory
                    (camp_entity_id,donor_name,donor_type,item_name,quantity,unit,receive_date,notes,proof_image)
                    VALUES(?,?,?,?,?,?,?,?,?)""",
                    (camp_id,donor,dtype,item,float(qty or 0),unit,rdate,notes,proof))
                conn.commit()
                flash("تم تسجيل الاستلام","success")
        elif action == "delete":
            iid = request.form.get("item_id")
            c.execute("DELETE FROM camp_inventory WHERE id=? AND camp_entity_id=?",(iid,camp_id))
            conn.commit()
            flash("تم الحذف","success")
        conn.close()
        return redirect(url_for("camp_inventory"))
    c.execute("SELECT * FROM camp_entities WHERE id=?",(camp_id,))
    entity = dict(c.fetchone())
    c.execute("""SELECT * FROM camp_inventory WHERE camp_entity_id=? ORDER BY receive_date DESC""",(camp_id,))
    items = [dict(r) for r in c.fetchall()]
    c.execute("SELECT COUNT(*) as n, SUM(quantity) as s FROM camp_inventory WHERE camp_entity_id=?",(camp_id,))
    stats = dict(c.fetchone())
    conn.close()
    return render_template("camp_inventory.html", entity=entity, items=items, stats=stats)

# ── التوزيعات — قائمة ──
@app.route("/camp/distributions")
@camp_login_required
def camp_distributions():
    camp_id = session["camp_id"]
    conn = get_connection(); c = conn.cursor()
    c.execute("SELECT * FROM camp_entities WHERE id=?",(camp_id,))
    entity = dict(c.fetchone())
    c.execute("""SELECT d.*,
        (SELECT COUNT(*) FROM camp_dist_records dr WHERE dr.distribution_id=d.id) as total,
        (SELECT COUNT(*) FROM camp_dist_records dr WHERE dr.distribution_id=d.id AND dr.received=1) as done
        FROM camp_distributions d WHERE d.camp_entity_id=? ORDER BY d.distribution_date DESC""",(camp_id,))
    dists = [dict(r) for r in c.fetchall()]
    conn.close()
    return render_template("camp_distributions.html", entity=entity, dists=dists)

# ── توزيع جديد ──
@app.route("/camp/distribution/new", methods=["GET","POST"])
@camp_login_required
def camp_distribution_new():
    camp_id = session["camp_id"]
    conn = get_connection(); c = conn.cursor()
    if request.method == "POST":
        title   = request.form.get("title","").strip()
        ddate   = request.form.get("distribution_date","")
        donor   = request.form.get("donor_name","").strip()
        ftype   = request.form.get("filter_type","all")
        fval    = request.form.get("filter_value","").strip()
        item    = request.form.get("item_name","").strip()
        qty     = request.form.get("quantity","").strip()
        value   = request.form.get("value","").strip()
        notes   = request.form.get("notes","").strip()
        if not title or not ddate:
            flash("العنوان والتاريخ مطلوبان","error")
        else:
            c.execute("""INSERT INTO camp_distributions
                (camp_entity_id,title,distribution_date,donor_name,filter_type,filter_value,status,notes)
                VALUES(?,?,?,?,?,?,'active',?)""",
                (camp_id,title,ddate,donor,ftype,fval,notes))
            dist_id = c.lastrowid
            # جلب المستفيدين حسب التصنيف
            members = get_classified_members(camp_id, conn)
            filtered = _apply_filter(members, ftype, fval)
            for m in filtered:
                c.execute("""INSERT OR IGNORE INTO camp_dist_records
                    (distribution_id,camp_entity_id,beneficiary_id,item_name,quantity,value)
                    VALUES(?,?,?,?,?,?)""",
                    (dist_id,camp_id,m["id"],item,qty,value))
                # إضافة id_number إذا موجود (للربط برقم الهوية مستقبلاً)
                try:
                    c.execute("UPDATE camp_dist_records SET notes=notes WHERE distribution_id=? AND beneficiary_id=?",
                              (dist_id, m["id"]))  # placeholder للتوافقية
                except: pass
            conn.commit()
            conn.close()
            flash(f"تم إنشاء جلسة التوزيع — {len(filtered)} مستفيد","success")
            return redirect(url_for("camp_distribution_detail", dist_id=dist_id))
    c.execute("SELECT * FROM camp_entities WHERE id=?",(camp_id,))
    entity = dict(c.fetchone())
    members = get_classified_members(camp_id, conn)
    conn.close()
    return render_template("camp_distribution_new.html", entity=entity, members=members, total=len(members))

def _apply_filter(members, ftype, fval):
    if ftype == "all":
        return members
    elif ftype == "count":
        try: n = int(fval)
        except: n = len(members)
        # الأقل استفادة أولاً
        return sorted(members, key=lambda m: (m.get("total_received") or 0, m.get("last_received_date") or ""))[:n]
    elif ftype == "orphans":
        return [m for m in members if m.get("has_orphans") or m.get("is_orphan")]
    elif ftype == "pregnant":
        return [m for m in members if m.get("wife_pregnant") or m.get("wife_nursing")]
    elif ftype == "sick":
        return [m for m in members if m.get("has_chronic_disease") or m.get("has_disability")]
    elif ftype == "large_family":
        try: n = int(fval or 5)
        except: n = 5
        return [m for m in members if (m.get("family_size") or 1) >= n]
    elif ftype == "least_received":
        # من لم يستلم أو استلم الأقل
        try: n = int(fval or len(members))
        except: n = len(members)
        return sorted(members, key=lambda m: (m.get("total_received") or 0))[:n]
    return members

# ── تفاصيل التوزيع ──
@app.route("/camp/distribution/<int:dist_id>", methods=["GET","POST"])
@camp_login_required
def camp_distribution_detail(dist_id):
    camp_id = session["camp_id"]
    conn = get_connection(); c = conn.cursor()
    c.execute("SELECT * FROM camp_distributions WHERE id=? AND camp_entity_id=?",(dist_id,camp_id))
    dist = c.fetchone()
    if not dist:
        conn.close(); flash("الجلسة غير موجودة","error"); return redirect(url_for("camp_distributions"))
    dist = dict(dist)
    if request.method == "POST":
        action = request.form.get("action","")
        if action == "toggle":
            rec_id = request.form.get("record_id")
            from datetime import datetime as _dtm
            c.execute("SELECT received FROM camp_dist_records WHERE id=? AND camp_entity_id=?",(rec_id,camp_id))
            row = c.fetchone()
            if row:
                new_val = 0 if row["received"] else 1
                ts = _dtm.now().strftime("%Y-%m-%d %H:%M") if new_val else None
                c.execute("UPDATE camp_dist_records SET received=?,received_at=? WHERE id=?",(new_val,ts,rec_id))
                conn.commit()
        elif action == "complete":
            c.execute("UPDATE camp_distributions SET status='completed' WHERE id=?",(dist_id,))
            conn.commit()
            flash("تم إغلاق جلسة التوزيع","success")
        conn.close()
        return redirect(url_for("camp_distribution_detail", dist_id=dist_id))
    c.execute("""SELECT dr.*,b.full_name,b.id_number,b.phone,b.family_size,b.gender,
                        b.has_orphans,b.wife_pregnant,b.wife_nursing
                 FROM camp_dist_records dr
                 JOIN beneficiaries b ON dr.beneficiary_id=b.id
                 WHERE dr.distribution_id=? ORDER BY b.full_name""",(dist_id,))
    records = [dict(r) for r in c.fetchall()]
    c.execute("SELECT * FROM camp_entities WHERE id=?",(camp_id,))
    entity = dict(c.fetchone())
    done = sum(1 for r in records if r["received"])
    conn.close()
    return render_template("camp_distribution_detail.html",
        entity=entity, dist=dist, records=records,
        done=done, total=len(records))

# ── طباعة كشف التوزيع ──
@app.route("/camp/distribution/<int:dist_id>/print")
@camp_login_required
def camp_distribution_print(dist_id):
    camp_id = session["camp_id"]
    conn = get_connection(); c = conn.cursor()
    c.execute("SELECT * FROM camp_distributions WHERE id=? AND camp_entity_id=?",(dist_id,camp_id))
    dist = dict(c.fetchone())
    c.execute("""SELECT dr.*,b.full_name,b.id_number,b.phone,b.family_size,b.address,b.gender
                 FROM camp_dist_records dr JOIN beneficiaries b ON dr.beneficiary_id=b.id
                 WHERE dr.distribution_id=? ORDER BY b.full_name""",(dist_id,))
    records = [dict(r) for r in c.fetchall()]
    c.execute("SELECT * FROM camp_entities WHERE id=?",(camp_id,))
    entity = dict(c.fetchone())
    conn.close()
    return render_template("camp_distribution_print.html", entity=entity, dist=dist, records=records)

# ── أرشيف الأنشطة ──
@app.route("/camp/activities", methods=["GET","POST"])
@camp_login_required
def camp_activities_route():
    camp_id = session["camp_id"]
    conn = get_connection(); c = conn.cursor()
    if request.method == "POST":
        action = request.form.get("action","add")
        if action == "add":
            title  = request.form.get("title","").strip()
            adate  = request.form.get("activity_date","")
            atype  = request.form.get("activity_type","general")
            desc   = request.form.get("description","").strip()
            if title and adate:
                c.execute("""INSERT INTO camp_activities(camp_entity_id,activity_date,title,description,activity_type)
                    VALUES(?,?,?,?,?)""",(camp_id,adate,title,desc,atype))
                act_id = c.lastrowid
                # رفع مرفقات
                import os, uuid
                for f in request.files.getlist("attachments"):
                    if f and f.filename:
                        ext = f.filename.rsplit('.',1)[-1].lower()
                        fname = f"act_{uuid.uuid4().hex[:8]}.{ext}"
                        fdir = os.path.join(os.path.dirname(__file__),'static','uploads','activities')
                        os.makedirs(fdir, exist_ok=True)
                        f.save(os.path.join(fdir, fname))
                        c.execute("INSERT INTO camp_activity_attachments(activity_id,file_path) VALUES(?,?)",
                                  (act_id, f"activities/{fname}"))
                conn.commit()
                flash("تم إضافة النشاط","success")
        elif action == "delete":
            aid = request.form.get("activity_id")
            c.execute("DELETE FROM camp_activities WHERE id=? AND camp_entity_id=?",(aid,camp_id))
            c.execute("DELETE FROM camp_activity_attachments WHERE activity_id=?",(aid,))
            conn.commit()
            flash("تم الحذف","success")
        conn.close()
        return redirect(url_for("camp_activities_route"))
    c.execute("SELECT * FROM camp_entities WHERE id=?",(camp_id,))
    entity = dict(c.fetchone())
    c.execute("""SELECT a.*,
        (SELECT COUNT(*) FROM camp_activity_attachments att WHERE att.activity_id=a.id) as attach_count
        FROM camp_activities a WHERE a.camp_entity_id=? ORDER BY a.activity_date DESC""",(camp_id,))
    activities = [dict(r) for r in c.fetchall()]
    conn.close()
    return render_template("camp_activities.html", entity=entity, activities=activities)

# ── تفاصيل نشاط ──
@app.route("/camp/activity/<int:act_id>")
@camp_login_required
def camp_activity_detail(act_id):
    camp_id = session["camp_id"]
    conn = get_connection(); c = conn.cursor()
    c.execute("SELECT * FROM camp_activities WHERE id=? AND camp_entity_id=?",(act_id,camp_id))
    act = c.fetchone()
    if not act:
        conn.close(); return redirect(url_for("camp_activities_route"))
    c.execute("SELECT * FROM camp_activity_attachments WHERE activity_id=?",(act_id,))
    attachments = [dict(r) for r in c.fetchall()]
    c.execute("SELECT * FROM camp_entities WHERE id=?",(camp_id,))
    entity = dict(c.fetchone())
    conn.close()
    return render_template("camp_activity_detail.html", entity=entity, act=dict(act), attachments=attachments)


# ── التنبيهات ──
@app.route("/camp/alerts", methods=["GET", "POST"])
@camp_login_required
def camp_alerts_route():
    camp_id = session["camp_id"]
    conn = get_connection(); c = conn.cursor()
    c.execute("SELECT * FROM camp_entities WHERE id=?", (camp_id,))
    entity = dict(c.fetchone())
    c.execute("""SELECT b.* FROM beneficiaries b
                 JOIN camp_join_requests jr ON jr.beneficiary_id=b.id
                 WHERE jr.camp_entity_id=? AND jr.status='approved'
                 ORDER BY b.full_name""", (camp_id,))
    beneficiaries = [dict(r) for r in c.fetchall()]
    conn.close()
    return render_template("camp_alerts.html", entity=entity, beneficiaries=beneficiaries)


# ── التصنيف الذكي ──
@app.route("/camp/smart-classify")
@camp_login_required
def camp_smart_classify():
    camp_id = session["camp_id"]
    conn = get_connection(); c = conn.cursor()
    c.execute("SELECT * FROM camp_entities WHERE id=?", (camp_id,))
    entity = dict(c.fetchone())
    c.execute("""SELECT b.* FROM beneficiaries b
                 JOIN camp_join_requests jr ON jr.beneficiary_id=b.id
                 WHERE jr.camp_entity_id=? AND jr.status='approved'
                 ORDER BY b.full_name""", (camp_id,))
    beneficiaries = [dict(r) for r in c.fetchall()]
    conn.close()
    orphans  = [b for b in beneficiaries if b.get("has_orphans")]
    pregnant = [b for b in beneficiaries if b.get("wife_pregnant")]
    nursing  = [b for b in beneficiaries if b.get("wife_nursing")]
    return render_template("camp_smart_classify.html",
                           entity=entity, beneficiaries=beneficiaries,
                           orphans=orphans, pregnant=pregnant, nursing=nursing)


# ══════════════════════════════════════════════════════════
# تأكيد الاستلام من طرف المؤسسة (بدون انتظار المستفيد)
# ══════════════════════════════════════════════════════════
@app.route("/api/org/confirm-benefit", methods=["POST"])
def api_org_confirm_benefit():
    import datetime
    if "user_id" not in session:
        return jsonify({"ok": False, "error": "غير مصرح"}), 401
    org_id = session["org_id"]
    data      = request.get_json(force=True) or {}
    record_id = data.get("record_id")
    if not record_id:
        return jsonify({"ok": False, "error": "معرّف السجل مطلوب"})
    conn = get_connection(); c = conn.cursor()
    c.execute("SELECT id FROM program_records WHERE id=? AND org_id=?", (record_id, org_id))
    if not c.fetchone():
        conn.close()
        return jsonify({"ok": False, "error": "السجل غير موجود"})
    now_str = datetime.datetime.utcnow().strftime("%Y-%m-%d %H:%M:%S")
    c.execute("UPDATE program_records SET ben_confirmed=1, ben_confirmed_at=? WHERE id=?", (now_str, record_id))
    conn.commit(); conn.close()
    return jsonify({"ok": True})


# ══════════════════════════════════════════════════════════
# تأكيد استلام المستفيد لسجل برنامج
# ══════════════════════════════════════════════════════════
@app.route("/api/beneficiary/confirm-benefit", methods=["POST"])
def api_ben_confirm_benefit():
    import datetime
    ben_id = session.get("beneficiary_id")
    if not ben_id:
        return jsonify({"ok": False, "error": "غير مصرح"}), 401
    data      = request.get_json(force=True) or {}
    record_id = data.get("record_id")
    if not record_id:
        return jsonify({"ok": False, "error": "معرّف السجل مطلوب"})
    conn = get_connection(); c = conn.cursor()
    # تأكد أن السجل يخص هذا المستفيد
    c.execute("SELECT id FROM program_records WHERE id=? AND beneficiary_id=?", (record_id, ben_id))
    row = c.fetchone()
    if not row:
        conn.close()
        return jsonify({"ok": False, "error": "السجل غير موجود"})
    now_str = datetime.datetime.utcnow().strftime("%Y-%m-%d %H:%M:%S")
    c.execute("UPDATE program_records SET ben_confirmed=1, ben_confirmed_at=? WHERE id=?", (now_str, record_id))
    conn.commit(); conn.close()
    return jsonify({"ok": True})


# ══════════════════════════════════════════════════════════
# تسليمات المؤسسة للمخيمات
# ══════════════════════════════════════════════════════════
@app.route("/org/camp-deliveries")
@login_required
def org_camp_deliveries():
    org_id = session["org_id"]
    conn = get_connection(); c = conn.cursor()
    # جلب المخيمات
    c.execute("SELECT id, name, entity_type FROM camp_entities WHERE is_active=1 ORDER BY name")
    camps = [dict(r) for r in c.fetchall()]
    # جلب سجلات التسليم
    c.execute("""
        SELECT ocd.*, ce.name as camp_name, ce.entity_type
        FROM org_camp_deliveries ocd
        JOIN camp_entities ce ON ocd.camp_entity_id = ce.id
        WHERE ocd.org_id=?
        ORDER BY ocd.delivery_date DESC, ocd.id DESC
    """, (org_id,))
    deliveries = [dict(r) for r in c.fetchall()]
    conn.close()
    import datetime as _dt_ocd
    now_date = _dt_ocd.date.today().isoformat()
    return render_template("org_camp_deliveries.html", camps=camps, deliveries=deliveries, now_date=now_date)


@app.route("/org/camp-deliveries/add", methods=["POST"])
@login_required
def org_camp_delivery_add():
    org_id = session["org_id"]
    camp_entity_id = request.form.get("camp_entity_id", "").strip()
    delivery_date  = request.form.get("delivery_date", "").strip()
    title          = request.form.get("title", "").strip()
    items          = request.form.get("items", "").strip()
    quantity       = request.form.get("quantity", "").strip()
    notes          = request.form.get("notes", "").strip()
    if not camp_entity_id or not delivery_date or not title:
        flash("يرجى تعبئة الحقول المطلوبة", "danger")
        return redirect(url_for("org_camp_deliveries"))
    conn = get_connection(); c = conn.cursor()
    c.execute("""
        INSERT INTO org_camp_deliveries
            (org_id, camp_entity_id, delivery_date, title, items, quantity, notes, created_by)
        VALUES (?,?,?,?,?,?,?,?)
    """, (org_id, camp_entity_id, delivery_date, title, items, quantity, notes, session.get("user", "")))
    conn.commit(); conn.close()
    flash("تم تسجيل التسليم بنجاح — في انتظار تأكيد المخيم", "success")
    return redirect(url_for("org_camp_deliveries"))


@app.route("/org/camp-deliveries/delete/<int:did>", methods=["POST"])
@login_required
def org_camp_delivery_delete(did):
    org_id = session["org_id"]
    conn = get_connection(); c = conn.cursor()
    c.execute("DELETE FROM org_camp_deliveries WHERE id=? AND org_id=? AND confirmed=0", (did, org_id))
    conn.commit(); conn.close()
    flash("تم حذف التسليم", "info")
    return redirect(url_for("org_camp_deliveries"))


# ══════════════════════════════════════════════════════════
# تأكيد استلام المخيم للتسليمات الواردة من المؤسسات
# ══════════════════════════════════════════════════════════
@app.route("/camp/deliveries")
@camp_login_required
def camp_deliveries():
    camp_id = session["camp_id"]
    conn = get_connection(); c = conn.cursor()
    c.execute("""
        SELECT ocd.*, o.name as org_name
        FROM org_camp_deliveries ocd
        JOIN organizations o ON ocd.org_id = o.id
        WHERE ocd.camp_entity_id=?
        ORDER BY ocd.confirmed ASC, ocd.delivery_date DESC
    """, (camp_id,))
    deliveries = [dict(r) for r in c.fetchall()]
    conn.close()
    return render_template("camp_deliveries.html", deliveries=deliveries)


@app.route("/api/camp/confirm-delivery", methods=["POST"])
@camp_login_required
def api_camp_confirm_delivery():
    import datetime
    camp_id = session["camp_id"]
    data        = request.get_json(force=True) or {}
    delivery_id = data.get("delivery_id")
    if not delivery_id:
        return jsonify({"ok": False, "error": "معرّف التسليم مطلوب"})
    conn = get_connection(); c = conn.cursor()
    c.execute("SELECT id FROM org_camp_deliveries WHERE id=? AND camp_entity_id=?", (delivery_id, camp_id))
    row = c.fetchone()
    if not row:
        conn.close()
        return jsonify({"ok": False, "error": "التسليم غير موجود"})
    now_str = datetime.datetime.utcnow().strftime("%Y-%m-%d %H:%M:%S")
    confirmed_by = session.get("camp_name", "")
    c.execute("""
        UPDATE org_camp_deliveries
        SET confirmed=1, confirmed_at=?, confirmed_by=?
        WHERE id=?
    """, (now_str, confirmed_by, delivery_id))
    conn.commit(); conn.close()
    return jsonify({"ok": True})


if __name__ == "__main__":
    init_db()
    app.run(debug=True)
